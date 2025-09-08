#!/usr/bin/env python
"""
Script untuk menganalisis isi file DOCX yang di-generate
"""
import os
import django

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'core.settings')
django.setup()

from docx import Document

def analyze_docx_content():
    """Menganalisis isi file DOCX yang di-generate"""
    print("=== Analyzing Generated DOCX Content ===\n")

    # Path ke file DOCX yang baru di-generate (gunakan file terbaru)
    import glob
    import os

    # Cari file DOCX terbaru untuk staff.konsuler1 bulan Agustus 2025
    pattern = '/app/media/overtime_summary_docx/rekap_lembur_STK001_2025-08*.docx'
    files = glob.glob(pattern)

    if files:
        # Ambil file terbaru berdasarkan waktu modifikasi
        docx_path = max(files, key=os.path.getmtime)
    else:
        docx_path = '/app/media/overtime_summary_docx/rekap_lembur_STK001_2025-08.docx'

    if not os.path.exists(docx_path):
        print(f"❌ File tidak ditemukan: {docx_path}")
        return

    print(f"✓ File ditemukan: {docx_path}")
    print(f"  Ukuran file: {os.path.getsize(docx_path)} bytes")

    try:
        # Load document
        doc = Document(docx_path)
        print(f"✓ Document berhasil dimuat dengan {len(doc.paragraphs)} paragraf")

        # Analisis paragraf
        print("\n--- Analisis Paragraf ---")
        for i, para in enumerate(doc.paragraphs[:10]):  # Tampilkan 10 paragraf pertama
            text = para.text.strip()
            if text:
                print(f"Paragraf {i+1}: {text[:100]}{'...' if len(text) > 100 else ''}")

        # Analisis tabel
        print(f"\n--- Analisis Tabel ({len(doc.tables)} tabel ditemukan) ---")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTabel {table_idx + 1}:")
            print(f"  Jumlah baris: {len(table.rows)}")
            print(f"  Jumlah kolom: {len(table.rows[0].cells) if table.rows else 0}")

            # Tampilkan header jika ada
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                print(f"  Header: {headers}")

                # Tampilkan beberapa baris data
                for row_idx in range(min(5, len(table.rows))):
                    row_data = [cell.text.strip() for cell in table.rows[row_idx].cells]
                    print(f"  Baris {row_idx + 1}: {row_data}")

        # Cari placeholder yang belum terganti
        print("\n--- Pencarian Placeholder ---")
        placeholders_found = []
        for para in doc.paragraphs:
            text = para.text
            if '{{' in text and '}}' in text:
                import re
                matches = re.findall(r'\{\{[^}]+\}\}', text)
                placeholders_found.extend(matches)

        # Cari di tabel juga
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if '{{' in text and '}}' in text:
                            import re
                            matches = re.findall(r'\{\{[^}]+\}\}', text)
                            placeholders_found.extend(matches)

        if placeholders_found:
            unique_placeholders = list(set(placeholders_found))
            print(f"⚠ Ditemukan {len(unique_placeholders)} placeholder yang belum terganti:")
            for placeholder in unique_placeholders:
                print(f"  - {placeholder}")
        else:
            print("✓ Semua placeholder sudah terganti")

        # Cari teks khusus terkait tabel overtime
        print("\n--- Pencarian Teks Tabel Overtime ---")
        overtime_keywords = ['overtime', 'lembur', 'tabel', 'detail', 'pengajuan']
        found_keywords = []

        for para in doc.paragraphs:
            text = para.text.lower()
            for keyword in overtime_keywords:
                if keyword in text:
                    found_keywords.append(f"Paragraf: '{para.text[:50]}...'")
                    break

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.lower()
                        for keyword in overtime_keywords:
                            if keyword in text:
                                found_keywords.append(f"Tabel: '{para.text[:50]}...'")
                                break

        if found_keywords:
            print("✓ Ditemukan teks terkait overtime:")
            for item in found_keywords[:5]:  # Tampilkan 5 pertama
                print(f"  - {item}")
        else:
            print("⚠ Tidak ditemukan teks terkait overtime")

        print("\n✓ Analisis selesai")

    except Exception as e:
        print(f"❌ Error saat menganalisis document: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    analyze_docx_content()
