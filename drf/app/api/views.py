from django.http import JsonResponse
from django.contrib.auth.models import User
from django.views.decorators.http import require_GET
from drf_spectacular.utils import extend_schema, OpenApiParameter
from rest_framework.decorators import api_view, permission_classes, authentication_classes, parser_classes
from rest_framework.permissions import IsAuthenticated, AllowAny, IsAdminUser
from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework import serializers
from .models import Division, Position, Employee, WorkSettings, Holiday, Attendance, AttendanceCorrection, OvertimeRequest, OvertimeSummaryRequest, GroupPermission, GroupPermissionTemplate
from .serializers import (
    DivisionSerializer,
    PositionSerializer,
    EmployeeSerializer,
    WorkSettingsSerializer,
    HolidaySerializer,
    AttendanceSerializer,
    AttendanceCorrectionSerializer,
    # Role-based serializers
    EmployeeAdminSerializer,
    EmployeeSupervisorSerializer,
    EmployeeEmployeeSerializer,
    WorkSettingsAdminSerializer,
    WorkSettingsSupervisorSerializer,
    AttendanceAdminSerializer,
    AttendanceSupervisorSerializer,
    AttendanceEmployeeSerializer,
    AttendanceCorrectionAdminSerializer,
    AttendanceCorrectionSupervisorSerializer,
    AttendanceCorrectionEmployeeSerializer,
    HolidayAdminSerializer,
    HolidayPublicSerializer,
    # Overtime request serializers
    OvertimeRequestAdminSerializer,
    OvertimeRequestSupervisorSerializer,
    OvertimeRequestEmployeeSerializer,
    OvertimeRequestCreateSerializer,
    # Monthly summary request serializers
    OvertimeSummaryRequestAdminSerializer,
    OvertimeSummaryRequestSupervisorSerializer,
    OvertimeSummaryRequestEmployeeSerializer,
    OvertimeSummaryRequestCreateSerializer,
    GroupSerializer,
    GroupCreateSerializer,
    GroupUpdateSerializer,
    GroupDetailSerializer,
    GroupAdminSerializer,
    GroupPermissionSerializer,
    GroupPermissionCreateSerializer,
    GroupPermissionUpdateSerializer,
    GroupPermissionDetailSerializer,
    GroupPermissionTemplateSerializer,
    GroupPermissionTemplateCreateSerializer,
    GroupPermissionTemplateUpdateSerializer,
    GroupWithPermissionsSerializer,
    BulkPermissionUpdateSerializer,
)
from .pagination import DefaultPagination
from .utils import evaluate_lateness_as_dict, haversine_meters
from .permissions import (
    IsAdmin,
    IsSupervisor,
    IsEmployee,
    IsAdminOrSupervisor,
    IsAdminOrSupervisorWithApproval,
    IsAdminOrSupervisorOvertimeApproval,
    IsAdminOrSupervisorReadOnly,
    IsAdminOrReadOnly,
    IsOwnerOrAdmin,
    IsDivisionMemberOrAdmin,
    IsOvertimeRequestOwnerOrSupervisor,
    IsOvertimeRequestApprover,
)
from zoneinfo import ZoneInfo
from django.utils import timezone as dj_timezone
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from django.db import models
from django.http import HttpResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from io import BytesIO
from datetime import datetime, timedelta
from django.db import IntegrityError
from django.db import transaction
from django.middleware.csrf import get_token
from django.views.decorators.csrf import ensure_csrf_cookie, csrf_exempt
from django.utils.decorators import method_decorator
from django.db.models import Q

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def format_work_hours(minutes, use_indonesian=True):
    """Format work hours from minutes to human-readable format"""
    if not minutes or minutes <= 0:
        return '-'
    
    hours = int(minutes // 60)
    mins = round(minutes % 60)
    
    # Handle edge case where minutes rounds up to 60
    if mins == 60:
        hours += 1
        mins = 0
    
    # Format based on whether we have hours and/or minutes
    hour_symbol = 'j' if use_indonesian else 'h'
    
    if hours > 0 and mins > 0:
        return f"{hours}{hour_symbol} {mins}m"
    elif hours > 0:
        return f"{hours}{hour_symbol}"
    elif mins > 0:
        return f"{mins}m"
    else:
        return '0m'

def health(request):
    return JsonResponse({"status": "ok"})


@extend_schema(
    parameters=[
        OpenApiParameter(name='username', type=str, required=True, description='Username to check'),
    ],
    responses={200: dict},
)
@require_GET
def check_user(request):
    username = request.GET.get('username')
    if not username:
        return JsonResponse({"detail": "username is required"}, status=400)
    exists = User.objects.filter(username=username).exists()
    return JsonResponse({"username": username, "exists": exists})


@extend_schema(
    responses={200: dict},
    auth=[{"BearerAuth": []}],
)
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def me(request):
    user = request.user
    data = {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "first_name": user.first_name,
        "last_name": user.last_name,
        "is_superuser": user.is_superuser,
        "groups": list(user.groups.values_list('name', flat=True)),
    }
    
    # Add position data for position-based approval system
    try:
        employee = user.employee
        if employee and employee.position:
            data["position"] = {
                "id": employee.position.id,
                "name": employee.position.name,
                "approval_level": employee.position.approval_level,
                "can_approve_overtime_org_wide": employee.position.can_approve_overtime_org_wide,
            }
        else:
            data["position"] = None
    except:
        data["position"] = None
    
    return JsonResponse(data)


@extend_schema(
    responses={200: dict},
    auth=[{"BearerAuth": []}],
)
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def employee_me(request):
    """Get current user's employee profile with position and approval level."""
    user = request.user
    try:
        employee = user.employee
        if not employee:
            return JsonResponse({"detail": "employee_record_not_found"}, status=404)
        
        data = {
            "id": employee.id,
            "user_id": employee.user_id,
            "fullname": employee.fullname,
            "division_id": employee.division_id,
            "position_id": employee.position_id,
            "user": {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "first_name": user.first_name,
                "last_name": user.last_name,
                "is_superuser": user.is_superuser,
                "groups": list(user.groups.values_list('name', flat=True)),
            }
        }
        
        # Add position details if available
        if employee.position:
            data["position"] = {
                "id": employee.position.id,
                "name": employee.position.name,
                "approval_level": employee.position.approval_level,
                "can_approve_overtime_org_wide": employee.position.can_approve_overtime_org_wide,
            }
        
        # Add division details if available
        if employee.division:
            data["division"] = {
                "id": employee.division.id,
                "name": employee.division.name,
            }
            
        return JsonResponse(data)
        
    except Exception as e:
        return JsonResponse({"detail": "error_fetching_profile", "error": str(e)}, status=500)


@extend_schema(
    responses={200: {"type": "object", "properties": {"message": {"type": "string"}}}},
    description="Logout user (JWT tokens are stateless, so this is mainly for client-side cleanup)",
)
@api_view(['POST'])
@permission_classes([AllowAny])  # Allow both authenticated and unauthenticated users
def logout(request):
    """
    Logout endpoint for JWT authentication.
    Since JWT tokens are stateless, this endpoint mainly serves
    to provide a consistent API interface for client-side cleanup.
    """
    response = JsonResponse({"message": "Logout successful"}, status=200)
    # Clear auth cookies used by frontend
    response.delete_cookie('access_token', path='/')
    response.delete_cookie('refresh_token', path='/')
    return response


@extend_schema(
    responses={200: list},
    auth=[{"BearerAuth": []}],
)
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def users_list(request):
    # Only allow admin or superuser to view the list of users
    user = request.user
    is_admin = bool(user.is_superuser or user.groups.filter(name='admin').exists())
    if not is_admin:
        return JsonResponse({"detail": "forbidden"}, status=403)

    User = get_user_model()
    users = User.objects.all().prefetch_related('groups')
    data = [
        {
            "id": u.id,
            "username": u.username,
            "email": u.email,
            "is_superuser": u.is_superuser,
            "groups": list(u.groups.values_list('name', flat=True)),
        }
        for u in users
    ]
    return JsonResponse(data, safe=False)


@extend_schema(
    request={
        'application/json': {
            'type': 'object',
            'properties': {
                'username': {'type': 'string'},
                'password': {'type': 'string'},
                'email': {'type': 'string'},
                'group': {'type': 'string', 'enum': ['admin', 'supervisor', 'pegawai']},
            },
            'required': ['username', 'group'],
        }
    },
    responses={200: dict},
    description='Provision user baru dan assign ke grup supervisor/pegawai (mendukung legacy HS/LS). Dev-only.',
)
@api_view(['POST'])
@permission_classes([AllowAny])
@authentication_classes([])
@csrf_exempt
def provision_user(request):
    data = request.data if hasattr(request, 'data') else {}
    username = data.get('username')
    password = data.get('password') or '1'
    email = data.get('email') or ''
    group_name = (data.get('group') or '').strip()
    alias = {'hs': 'supervisor', 'ls': 'pegawai', 'localstaff': 'pegawai'}
    mapped = alias.get(group_name.lower(), group_name)
    if not username or mapped not in {'admin', 'supervisor', 'pegawai'}:
        return JsonResponse({'detail': 'username dan group (admin/supervisor/pegawai) wajib.'}, status=400)

    User = get_user_model()
    user, created = User.objects.get_or_create(username=username, defaults={'email': email})
    if not created and email and user.email != email:
        user.email = email
    user.is_active = True
    user.set_password(password)
    user.save()

    group, _ = Group.objects.get_or_create(name=mapped)
    user.groups.add(group)

    return JsonResponse({
        'id': user.id,
        'username': user.username,
        'email': user.email,
        'groups': list(user.groups.values_list('name', flat=True)),
        'created': created,
    })


class DivisionViewSet(viewsets.ModelViewSet):
    queryset = Division.objects.all()
    serializer_class = DivisionSerializer
    permission_classes = [IsAdminOrReadOnly]
    pagination_class = DefaultPagination


class PositionViewSet(viewsets.ModelViewSet):
    queryset = Position.objects.all()
    serializer_class = PositionSerializer
    permission_classes = [IsAdminOrReadOnly]
    pagination_class = DefaultPagination


class EmployeeViewSet(viewsets.ModelViewSet):
    queryset = Employee.objects.select_related('user', 'division', 'position').all()
    serializer_class = EmployeeSerializer
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination


class WorkSettingsViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminOrSupervisorReadOnly]

    def list(self, request):
        obj, _ = WorkSettings.objects.get_or_create()
        return JsonResponse(WorkSettingsSerializer(obj).data, safe=False)

    def update(self, request, pk=None):
        obj, _ = WorkSettings.objects.get_or_create()
        serializer = WorkSettingsSerializer(obj, data=request.data, partial=False)
        if serializer.is_valid():
            # Basic validation: start < end (no overnight for now)
            start = serializer.validated_data.get("start_time", obj.start_time)
            end = serializer.validated_data.get("end_time", obj.end_time)
            if start >= end:
                return JsonResponse({"detail": "start_time must be earlier than end_time"}, status=400)
            
            # Friday-specific validation
            friday_start = serializer.validated_data.get("friday_start_time", obj.friday_start_time)
            friday_end = serializer.validated_data.get("friday_end_time", obj.friday_end_time)
            if friday_start >= friday_end:
                return JsonResponse({"detail": "friday_start_time must be earlier than friday_end_time"}, status=400)
            
            # Validation for earliest check-in time
            earliest_check_in = serializer.validated_data.get("earliest_check_in_time", obj.earliest_check_in_time)
            if earliest_check_in >= start:
                return JsonResponse({
                    "detail": "earliest_check_in_time must be earlier than start_time"
                }, status=400)
            
            workdays = serializer.validated_data.get("workdays", obj.workdays)
            if not isinstance(workdays, list) or not all(isinstance(x, int) and 0 <= x <= 6 for x in workdays):
                return JsonResponse({"detail": "workdays must be a list of integers 0..6"}, status=400)
            serializer.save()
            return JsonResponse(serializer.data)
        return JsonResponse(serializer.errors, status=400)


class HolidayViewSet(viewsets.ModelViewSet):
    queryset = Holiday.objects.all()
    serializer_class = HolidaySerializer
    permission_classes = [IsAdminOrReadOnly]
    pagination_class = DefaultPagination

    def get_queryset(self):
        queryset = Holiday.objects.all()
        
        # Filter by start date
        start_date = self.request.query_params.get('start', None)
        if start_date:
            queryset = queryset.filter(date__gte=start_date)
        
        # Filter by end date
        end_date = self.request.query_params.get('end', None)
        if end_date:
            queryset = queryset.filter(date__lte=end_date)
        
        return queryset.order_by('date')

# -----------------------------
# Role-specific ViewSets (Stage 3)
# -----------------------------

# Admin ViewSets
class AdminDivisionViewSet(viewsets.ModelViewSet):
    queryset = Division.objects.all()
    serializer_class = DivisionSerializer
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination

class AdminPositionViewSet(viewsets.ModelViewSet):
    queryset = Position.objects.all()
    serializer_class = PositionSerializer
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination

@method_decorator(csrf_exempt, name='dispatch')
class AdminEmployeeViewSet(viewsets.ModelViewSet):
    queryset = Employee.objects.select_related('user', 'division', 'position').all()
    serializer_class = EmployeeAdminSerializer
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination

class AdminWorkSettingsViewSet(viewsets.ViewSet):
    permission_classes = [IsAdmin]

    def list(self, request):
        obj, _ = WorkSettings.objects.get_or_create()
        return JsonResponse(WorkSettingsAdminSerializer(obj).data, safe=False)

    def update(self, request, pk=None):
        obj, _ = WorkSettings.objects.get_or_create()
        serializer = WorkSettingsAdminSerializer(obj, data=request.data, partial=False)
        if serializer.is_valid():
            # Basic validation: start < end (no overnight for now)
            start = serializer.validated_data.get("start_time", obj.start_time)
            end = serializer.validated_data.get("end_time", obj.end_time)
            if start >= end:
                return JsonResponse({"detail": "start_time must be earlier than end_time"}, status=400)

            # Friday-specific validation
            friday_start = serializer.validated_data.get("friday_start_time", obj.friday_start_time)
            friday_end = serializer.validated_data.get("friday_end_time", obj.friday_end_time)
            if friday_start >= friday_end:
                return JsonResponse({"detail": "friday_start_time must be earlier than friday_end_time"}, status=400)

            workdays = serializer.validated_data.get("workdays", obj.workdays)
            if not isinstance(workdays, list) or not all(isinstance(x, int) and 0 <= x <= 6 for x in workdays):
                return JsonResponse({"detail": "workdays must be a list of integers 0..6"}, status=400)
            serializer.save()
            return JsonResponse(serializer.data)
        return JsonResponse(serializer.errors, status=400)

class AdminHolidayViewSet(viewsets.ModelViewSet):
    queryset = Holiday.objects.all()
    serializer_class = HolidayAdminSerializer
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination

    def get_queryset(self):
        queryset = Holiday.objects.all()
        
        # Filter by start date
        start_date = self.request.query_params.get('start', None)
        if start_date:
            queryset = queryset.filter(date__gte=start_date)
        
        # Filter by end date
        end_date = self.request.query_params.get('end', None)
        if end_date:
            queryset = queryset.filter(date__lte=end_date)
        
        return queryset.order_by('date')

# Supervisor ViewSets (read-only or scoped)
class SupervisorDivisionViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Division.objects.all()
    serializer_class = DivisionSerializer
    permission_classes = [IsAdminOrSupervisorReadOnly]
    pagination_class = DefaultPagination

class SupervisorPositionViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Position.objects.all()
    serializer_class = PositionSerializer
    permission_classes = [IsAdminOrSupervisorReadOnly]
    pagination_class = DefaultPagination

class SupervisorEmployeeViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = EmployeeSupervisorSerializer
    permission_classes = [IsAdminOrSupervisor]
    pagination_class = DefaultPagination

    def get_queryset(self):
        user = self.request.user
        try:
            division_id = user.employee.division_id  # type: ignore[attr-defined]
        except Exception:
            division_id = None
        if not division_id:
            return Employee.objects.none()
        return Employee.objects.select_related('user', 'division', 'position').filter(division_id=division_id)

class SupervisorWorkSettingsViewSet(viewsets.ViewSet):
    permission_classes = [IsAdminOrSupervisorReadOnly]

    def list(self, request):
        obj, _ = WorkSettings.objects.get_or_create()
        return JsonResponse(WorkSettingsSupervisorSerializer(obj).data, safe=False)

class SupervisorHolidayViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Holiday.objects.all()
    serializer_class = HolidayPublicSerializer
    permission_classes = [IsAdminOrSupervisorReadOnly]
    pagination_class = DefaultPagination

    def get_queryset(self):
        queryset = Holiday.objects.all()
        
        # Filter by start date
        start_date = self.request.query_params.get('start', None)
        if start_date:
            queryset = queryset.filter(date__gte=start_date)
        
        # Filter by end date
        end_date = self.request.query_params.get('end', None)
        if end_date:
            queryset = queryset.filter(date__lte=end_date)
        
        return queryset.order_by('date')

# Employee ViewSets (read-only minimal)
class EmployeeDivisionViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Division.objects.all()
    serializer_class = DivisionSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = DefaultPagination

class EmployeePositionViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Position.objects.all()
    serializer_class = PositionSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = DefaultPagination

class EmployeeEmployeeViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = EmployeeEmployeeSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = DefaultPagination

    def get_queryset(self):
        return Employee.objects.select_related('user', 'division', 'position').filter(user=self.request.user)

class EmployeeHolidayViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Holiday.objects.all()
    serializer_class = HolidayPublicSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = DefaultPagination

    def get_queryset(self):
        queryset = Holiday.objects.all()
        
        # Filter by start date
        start_date = self.request.query_params.get('start', None)
        if start_date:
            queryset = queryset.filter(date__gte=start_date)
        
        # Filter by end date
        end_date = self.request.query_params.get('end', None)
        if end_date:
            queryset = queryset.filter(date__lte=end_date)
        
        return queryset.order_by('date')


class AttendanceViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = AttendanceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = DefaultPagination

    def get_queryset(self):
        user = self.request.user
        qs = Attendance.objects.select_related('employee').filter(user=user)
        start = self.request.query_params.get('start')
        end = self.request.query_params.get('end')
        month = self.request.query_params.get('month')
        
        if start:
            qs = qs.filter(date_local__gte=start)
        if end:
            qs = qs.filter(date_local__lte=end)
        if month:
            # Parse month and filter by year-month
            try:
                year, month_num = month.split('-')
                qs = qs.filter(
                    date_local__year=year,
                    date_local__month=month_num
                )
            except ValueError:
                pass  # Ignore invalid month format
        
        return qs.order_by('-date_local')


class AttendanceCorrectionViewSet(viewsets.ModelViewSet):
    serializer_class = AttendanceCorrectionSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = DefaultPagination
    parser_classes = [JSONParser, MultiPartParser, FormParser]

    def get_queryset(self):
        user = self.request.user
        qs = AttendanceCorrection.objects.select_related('user', 'user__employee__division').all()
        # Scope by role:
        # - admin/superuser: see all
        # - supervisor: only corrections of users within same division
        # - employee: only own corrections
        if user.is_superuser or user.groups.filter(name='admin').exists():
            pass
        elif user.groups.filter(name='supervisor').exists():
            try:
                supervisor_division_id = user.employee.division_id  # type: ignore[attr-defined]
            except Exception:
                supervisor_division_id = None
            if supervisor_division_id is None:
                qs = qs.none()
            else:
                qs = qs.filter(user__employee__division_id=supervisor_division_id)
        else:
            qs = qs.filter(user=user)
        status_ = self.request.query_params.get('status')
        if status_:
            qs = qs.filter(status=status_)
        return qs.order_by('-created_at')

    def perform_create(self, serializer):
        # Force owner to current user
        serializer.save(user=self.request.user)

    @extend_schema(request=None, responses={200: AttendanceCorrectionSerializer})
    def approve(self, request, pk=None):
        """Approve a pending correction (supervisor/admin only). Applies changes to Attendance."""
        if not IsAdminOrSupervisorWithApproval().has_permission(request, self):
            return Response({"detail": "forbidden"}, status=status.HTTP_403_FORBIDDEN)
        try:
            corr = AttendanceCorrection.objects.get(pk=pk)
        except AttendanceCorrection.DoesNotExist:
            return Response({"detail": "not_found"}, status=status.HTTP_404_NOT_FOUND)
        if corr.status != AttendanceCorrection.CorrectionStatus.PENDING:
            return Response({"detail": "invalid_status"}, status=status.HTTP_400_BAD_REQUEST)

        # Division-based authorization for supervisors: must match division of the correction's user
        user = request.user
        if (not (user.is_superuser or user.groups.filter(name='admin').exists())) and user.groups.filter(name='supervisor').exists():
            try:
                supervisor_division_id = user.employee.division_id  # type: ignore[attr-defined]
                employee_division_id = corr.user.employee.division_id  # type: ignore[attr-defined]
            except Exception:
                return Response({"detail": "division_not_configured"}, status=status.HTTP_403_FORBIDDEN)
            if not supervisor_division_id or not employee_division_id or supervisor_division_id != employee_division_id:
                return Response({"detail": "forbidden_division_scope"}, status=status.HTTP_403_FORBIDDEN)

        ws, _ = WorkSettings.objects.get_or_create()
        tzname = ws.timezone or dj_timezone.get_current_timezone_name()
        tz = ZoneInfo(tzname)

        # Prepare target Attendance record
        att, _ = Attendance.objects.get_or_create(
            user=corr.user,
            date_local=corr.date_local,
            defaults={"timezone": tzname},
        )

        # Convert proposed local datetimes to UTC
        def to_utc(local_dt):
            if not local_dt:
                return None
            # Treat stored value as naive local time
            aware_local = local_dt.replace(tzinfo=tz)
            return aware_local.astimezone(ZoneInfo("UTC"))

        cin_utc = to_utc(corr.proposed_check_in_local)
        cout_utc = to_utc(corr.proposed_check_out_local)

        # Apply based on type
        if corr.type == AttendanceCorrection.CorrectionType.MISSING_CHECK_IN:
            if not cin_utc:
                return Response({"detail": "proposed_check_in_local_required"}, status=status.HTTP_400_BAD_REQUEST)
            att.check_in_at_utc = att.check_in_at_utc or cin_utc
        elif corr.type == AttendanceCorrection.CorrectionType.MISSING_CHECK_OUT:
            if not cout_utc:
                return Response({"detail": "proposed_check_out_local_required"}, status=status.HTTP_400_BAD_REQUEST)
            att.check_out_at_utc = att.check_out_at_utc or cout_utc
        else:  # EDIT
            if not cin_utc and not cout_utc:
                return Response({"detail": "one_of_check_in_or_out_required"}, status=status.HTTP_400_BAD_REQUEST)
            if cin_utc:
                att.check_in_at_utc = cin_utc
            if cout_utc:
                att.check_out_at_utc = cout_utc

        # Recompute is_holiday and lateness
        is_holiday = Holiday.objects.filter(date=corr.date_local).exists()
        att.is_holiday = is_holiday
        if att.check_in_at_utc:
            late_res = evaluate_lateness_as_dict(att.check_in_at_utc)
            att.minutes_late = int(late_res.get("minutes_late") or 0)

        # Recompute total_work_minutes when possible
        if att.check_in_at_utc and att.check_out_at_utc and att.check_out_at_utc > att.check_in_at_utc:
            delta = att.check_out_at_utc - att.check_in_at_utc
            att.total_work_minutes = max(0, int(delta.total_seconds() // 60))

        # Link employee if possible (needed for overtime pay calculation)
        try:
            if not att.employee and hasattr(corr.user, "employee"):
                att.employee = corr.user.employee  # type: ignore[attr-defined]
        except Exception:
            pass

        # Recalculate overtime based on total work duration and work settings
        try:
            # Determine required minutes for the day (Friday may differ)
            weekday = int(att.date_local.weekday())
            if weekday == 4:  # Friday
                required_minutes = int(ws.friday_required_minutes or 240)
            else:
                required_minutes = int(ws.required_minutes or 480)

            overtime_threshold = int(ws.overtime_threshold_minutes or 60)  # Default 1 hour buffer
            if att.total_work_minutes and att.total_work_minutes > (required_minutes + overtime_threshold):
                att.overtime_minutes = att.total_work_minutes - required_minutes - overtime_threshold

                # Calculate overtime amount if employee has salary information
                if att.employee and att.employee.gaji_pokok:
                    monthly_hours = 22 * 8  # 22 workdays * 8 hours per day
                    hourly_wage = float(att.employee.gaji_pokok) / monthly_hours

                    # Determine overtime rate based on holiday vs workday
                    if att.is_holiday:
                        rate = float(ws.overtime_rate_holiday or 0.75)
                    else:
                        rate = float(ws.overtime_rate_workday or 0.50)

                    overtime_hours = att.overtime_minutes / 60
                    att.overtime_amount = overtime_hours * hourly_wage * rate
                else:
                    att.overtime_amount = 0

                # Overtime requires manual approval; keep pending
                att.overtime_approved = False
                att.overtime_approved_by = None
                att.overtime_approved_at = None
            else:
                # No overtime for this approval
                att.overtime_minutes = 0
                att.overtime_amount = 0
                att.overtime_approved = False
                att.overtime_approved_by = None
                att.overtime_approved_at = None
        except Exception:
            # Fail-safe: do not block approval due to overtime calc error
            pass

        # Mark manual and attach reason into employee_note (append)
        try:
            reason = (corr.reason or "").strip()
            if reason:
                att.employee_note = ((att.employee_note or "").strip() + ("\n" if att.employee_note else "") + f"[Correction]: {reason}")
        except Exception:
            pass

        att.save()

        # Close correction
        corr.status = AttendanceCorrection.CorrectionStatus.APPROVED
        corr.reviewed_by = request.user
        corr.reviewed_at = dj_timezone.now()
        corr.decision_note = (request.data.get("decision_note") or corr.decision_note)
        corr.save(update_fields=["status", "reviewed_by", "reviewed_at", "decision_note"])

        return Response(AttendanceCorrectionSerializer(corr).data)

    @extend_schema(request=None, responses={200: AttendanceCorrectionSerializer})
    def reject(self, request, pk=None):
        """Reject a pending correction (supervisor/admin only)."""
        if not IsAdminOrSupervisorWithApproval().has_permission(request, self):
            return Response({"detail": "forbidden"}, status=status.HTTP_403_FORBIDDEN)
        try:
            corr = AttendanceCorrection.objects.get(pk=pk)
        except AttendanceCorrection.DoesNotExist:
            return Response({"detail": "not_found"}, status=status.HTTP_404_NOT_FOUND)
        if corr.status != AttendanceCorrection.CorrectionStatus.PENDING:
            return Response({"detail": "invalid_status"}, status=status.HTTP_400_BAD_REQUEST)
        corr.status = AttendanceCorrection.CorrectionStatus.REJECTED
        corr.reviewed_by = request.user
        corr.reviewed_at = dj_timezone.now()
        corr.decision_note = request.data.get("decision_note")
        corr.save(update_fields=["status", "reviewed_by", "reviewed_at", "decision_note"])
        return Response(AttendanceCorrectionSerializer(corr).data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def supervisor_team_attendance(request):
    """Get attendance data for supervisor's team members."""
    user = request.user
    
    # Check if user is admin or has supervisor capabilities (position approval level >= 1)
    if user.is_superuser or user.groups.filter(name='admin').exists():
        pass  # Admin can see all data
    else:
        from .utils import ApprovalChecker
        approval_level = ApprovalChecker.get_user_approval_level(user)
        if approval_level < 1:
            return JsonResponse({"detail": "forbidden"}, status=403)
    
    # Get supervisor's division
    try:
        supervisor_division_id = user.employee.division_id
        if not supervisor_division_id:
            return JsonResponse({"detail": "supervisor_division_not_configured"}, status=400)
    except Exception:
        return JsonResponse({"detail": "supervisor_employee_record_not_found"}, status=400)
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    employee_id = request.GET.get('employee_id')
    
    # Build base queryset for team members in same division
    team_employees = Employee.objects.filter(
        division_id=supervisor_division_id
    ).select_related('user', 'division', 'position')
    
    # Filter by specific employee if requested
    if employee_id:
        team_employees = team_employees.filter(id=employee_id)
    
    # Get attendance data for team members
    attendance_data = []
    
    for employee in team_employees:
        # Get attendance records for this employee
        attendance_qs = Attendance.objects.filter(
            user=employee.user
        ).select_related('user', 'user__employee', 'user__employee__division', 'user__employee__position')
        
        # Apply date filters
        if start_date:
            attendance_qs = attendance_qs.filter(date_local__gte=start_date)
        if end_date:
            attendance_qs = attendance_qs.filter(date_local__lte=end_date)
        
        # Get attendance records
        attendances = attendance_qs.order_by('-date_local')
        
        # Calculate summary statistics
        total_days = attendances.count()
        present_days = attendances.filter(check_in_at_utc__isnull=False).count()
        late_days = attendances.filter(minutes_late__gt=0).count()
        absent_days = total_days - present_days
        
        # Get recent attendance (last 7 days)
        recent_attendance = attendances[:7]
        
        employee_data = {
            "employee": {
                "id": employee.id,
                "nip": employee.nip,
                "fullname": employee.fullname,
                "user": {
                    "id": employee.user.id,
                    "username": employee.user.username,
                    "first_name": employee.user.first_name,
                    "last_name": employee.user.last_name,
                    "email": employee.user.email,
                },
                "division": {
                    "id": employee.division.id,
                    "name": employee.division.name
                } if employee.division else None,
                "position": {
                    "id": employee.position.id,
                    "name": employee.position.name
                } if employee.position else None,
            },
            "summary": {
                "total_days": total_days,
                "present_days": present_days,
                "late_days": late_days,
                "absent_days": absent_days,
                "attendance_rate": round((present_days / total_days * 100) if total_days > 0 else 0, 2)
            },
            "recent_attendance": [
                {
                    "id": att.id,
                    "date_local": str(att.date_local),
                    "check_in_at_utc": att.check_in_at_utc.isoformat() if att.check_in_at_utc else None,
                    "check_out_at_utc": att.check_out_at_utc.isoformat() if att.check_out_at_utc else None,
                    "minutes_late": att.minutes_late,
                    "total_work_minutes": att.total_work_minutes,
                    "is_holiday": att.is_holiday,
                    "within_geofence": att.within_geofence,
                    "note": att.note,
                    "employee_note": att.employee_note,
                }
                for att in recent_attendance
            ]
        }
        
        attendance_data.append(employee_data)
    
    return JsonResponse({
        "team_attendance": attendance_data,
        "filters": {
            "start_date": start_date,
            "end_date": end_date,
            "employee_id": employee_id,
            "division_id": supervisor_division_id
        }
    }, safe=False)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def supervisor_attendance_detail(request, employee_id):
    """Get detailed attendance data for a specific team member."""
    user = request.user
    
    # Check if user is admin or has supervisor capabilities (position approval level >= 1)
    if user.is_superuser or user.groups.filter(name='admin').exists():
        pass  # Admin can see all data
    else:
        from .utils import ApprovalChecker
        approval_level = ApprovalChecker.get_user_approval_level(user)
        if approval_level < 1:
            return JsonResponse({"detail": "forbidden"}, status=403)
    
    # Get supervisor's division
    try:
        supervisor_division_id = user.employee.division_id
        if not supervisor_division_id:
            return JsonResponse({"detail": "supervisor_division_not_configured"}, status=400)
    except Exception:
        return JsonResponse({"detail": "supervisor_employee_record_not_found"}, status=400)
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    month = request.GET.get('month')  # Format: YYYY-MM
    
    # Get the employee and verify they're in supervisor's division
    try:
        employee = Employee.objects.select_related(
            'user', 'division', 'position'
        ).get(
            id=employee_id,
            division_id=supervisor_division_id
        )
    except Employee.DoesNotExist:
        return JsonResponse({"detail": "employee_not_found_or_not_in_team"}, status=404)
    
    # Build attendance queryset
    attendance_qs = Attendance.objects.filter(
        user=employee.user
    ).select_related('user', 'user__employee', 'user__employee__division', 'user__employee__position')
    
    # Apply date filters
    if start_date:
        attendance_qs = attendance_qs.filter(date_local__gte=start_date)
    if end_date:
        attendance_qs = attendance_qs.filter(date_local__lte=end_date)
    if month:
        # Parse month and filter by year-month
        try:
            year, month_num = month.split('-')
            attendance_qs = attendance_qs.filter(
                date_local__year=year,
                date_local__month=month_num
            )
        except ValueError:
            return JsonResponse({"detail": "invalid_month_format_use_yyyy_mm"}, status=400)
    
    # Get attendance records
    attendances = attendance_qs.order_by('-date_local')
    
    # Calculate detailed statistics
    total_days = attendances.count()
    present_days = attendances.filter(check_in_at_utc__isnull=False).count()
    late_days = attendances.filter(minutes_late__gt=0).count()
    absent_days = total_days - present_days
    total_late_minutes = attendances.filter(minutes_late__gt=0).aggregate(
        total=models.Sum('minutes_late')
    )['total'] or 0
    total_work_minutes = attendances.filter(
        total_work_minutes__gt=0
    ).aggregate(
        total=models.Sum('total_work_minutes')
    )['total'] or 0
    
    # Get all attendance records for the period
    attendance_records = [
        {
            "id": att.id,
            "date_local": str(att.date_local),
            "check_in_at_utc": att.check_in_at_utc.isoformat() if att.check_in_at_utc else None,
            "check_out_at_utc": att.check_out_at_utc.isoformat() if att.check_out_at_utc else None,
            "check_in_lat": float(att.check_in_lat) if att.check_in_lat else None,
            "check_in_lng": float(att.check_in_lng) if att.check_in_lng else None,
            "check_out_lat": float(att.check_out_lat) if att.check_out_lat else None,
            "check_out_lng": float(att.check_out_lng) if att.check_out_lng else None,
            "check_in_ip": att.check_in_ip,
            "check_out_ip": att.check_out_ip,
            "minutes_late": att.minutes_late,
            "total_work_minutes": att.total_work_minutes,
            "is_holiday": att.is_holiday,
            "within_geofence": att.within_geofence,
            "note": att.note,
            "employee_note": att.employee_note,
            "created_at": att.created_at.isoformat(),
            "updated_at": att.updated_at.isoformat(),
        }
        for att in attendances
    ]
    
    return JsonResponse({
        "employee": {
            "id": employee.id,
            "nip": employee.nip,
            "fullname": employee.fullname,
            "user": {
                "id": employee.user.id,
                "username": employee.user.username,
                "first_name": employee.user.first_name,
                "last_name": employee.user.last_name,
                "email": employee.user.email,
            },
            "division": {
                "id": employee.division.id,
                "name": employee.division.name
            } if employee.division else None,
            "position": {
                "id": employee.position.id,
                "name": employee.position.name
            } if employee.position else None,
        },
        "summary": {
            "total_days": total_days,
            "present_days": present_days,
            "late_days": late_days,
            "absent_days": absent_days,
            "attendance_rate": round((present_days / total_days * 100) if total_days > 0 else 0, 2),
            "total_late_minutes": total_late_minutes,
            "total_work_minutes": total_work_minutes,
            "average_work_minutes": round(total_work_minutes / present_days, 2) if present_days > 0 else 0
        },
        "attendance_records": attendance_records,
        "filters": {
            "start_date": start_date,
            "end_date": end_date,
            "month": month
        }
    }, safe=False)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def attendance_precheck(request):
    """Return current day attendance status for confirmation UI."""
    user = request.user
    now = dj_timezone.now()
    ws, _ = WorkSettings.objects.get_or_create()
    tz = ZoneInfo(ws.timezone or dj_timezone.get_current_timezone_name())
    local_now = now.astimezone(tz)
    date_local = local_now.date()
    att = Attendance.objects.filter(user=user, date_local=date_local).first()
    has_check_in = bool(att and att.check_in_at_utc)
    has_check_out = bool(att and att.check_out_at_utc)
    return JsonResponse({
        'date_local': str(date_local),
        'has_check_in': has_check_in,
        'has_check_out': has_check_out,
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def attendance_check_in(request):
    user = request.user
    now = dj_timezone.now()
    ws, _ = WorkSettings.objects.get_or_create()
    tz = ZoneInfo(ws.timezone or dj_timezone.get_current_timezone_name())
    local_now = now.astimezone(tz)
    date_local = local_now.date()
    # Prevent double check-in today
    att, _created = Attendance.objects.get_or_create(user=user, date_local=date_local, defaults={
        'timezone': ws.timezone or dj_timezone.get_current_timezone_name(),
    })
    if att.check_in_at_utc:
        return JsonResponse({'detail': 'Already checked in.'}, status=400)

    # Parse location
    data = request.data if hasattr(request, 'data') else {}
    try:
        lat = float(data.get('lat'))
        lng = float(data.get('lng'))
        acc = int(data.get('accuracy_m') or 0)
        client_ip = data.get('client_ip', 'unknown')
    except Exception:
        return JsonResponse({'detail': 'Invalid location payload'}, status=400)

    # Geofence validation
    if ws.office_latitude is None or ws.office_longitude is None:
        return JsonResponse({'detail': 'Office location is not configured'}, status=400)
    dist = haversine_meters(float(ws.office_latitude), float(ws.office_longitude), lat, lng)
    within = dist <= float(ws.office_radius_meters or 100)
    
    # Allow check-in/out even if outside geofence, but mark it
    # if not within:
    #     return JsonResponse({'detail': 'Outside office radius'}, status=403)

    # Holiday/workday determination
    is_holiday = Holiday.objects.filter(date=date_local).exists()
    minutes_late = 0
    if not is_holiday:
        # Compute lateness using local time now
        # evaluate_lateness expects a datetime; we pass now (server) which will be normalized
        res = evaluate_lateness_as_dict(now)
        minutes_late = int(res.get('minutes_late') or 0)
        
        # Double-check calculation for accuracy
        try:
            workdays = ws.workdays or []
            weekday = local_now.weekday()
            is_workday = weekday in workdays
            
            if is_workday and not is_holiday:
                # Manual calculation for verification
                if weekday == 4:  # Friday
                    start_time = ws.friday_start_time
                    grace_minutes = ws.friday_grace_minutes or 0
                else:
                    start_time = ws.start_time
                    grace_minutes = ws.grace_minutes or 0
                
                base_start = datetime.combine(date_local, start_time, tz)
                grace_delta = local_now - base_start
                grace_minutes_from_start = grace_delta.total_seconds() / 60
                
                if grace_minutes_from_start > grace_minutes:
                    manual_minutes_late = int(grace_minutes_from_start - grace_minutes)
                    # Use manual calculation if it differs significantly
                    if abs(manual_minutes_late - minutes_late) > 1:
                        minutes_late = manual_minutes_late
        except Exception:
            pass  # Keep original calculation if manual calculation fails
    
    # Validate earliest check-in time
    if ws.earliest_check_in_enabled:
        earliest_check_in_local = datetime.combine(date_local, ws.earliest_check_in_time, tz)
        if local_now < earliest_check_in_local:
            return JsonResponse({
                'detail': f'Check-in terlalu awal. Jam absen dibuka mulai {ws.earliest_check_in_time.strftime("%H:%M")}'
            }, status=400)

    # Validate latest check-out time
    if ws.latest_check_out_enabled:
        latest_check_out_local = datetime.combine(date_local, ws.latest_check_out_time, tz)
        if local_now > latest_check_out_local:
            return JsonResponse({
                'detail': f'Check-out terlalu terlambat. Jam checkout ditutup pada {ws.latest_check_out_time.strftime("%H:%M")}'
            }, status=400)

    att.check_in_at_utc = now
    att.check_in_lat = lat
    att.check_in_lng = lng
    att.check_in_accuracy_m = acc
    att.check_in_ip = client_ip
    att.within_geofence = within
    att.is_holiday = is_holiday
    att.minutes_late = minutes_late
    
    # Add warning note if outside geofence
    if not within:
        warning_note = f"Check-in di luar area kantor (jarak: {dist:.0f}m dari radius {ws.office_radius_meters}m)"
        att.note = warning_note if not att.note else f"{att.note}; {warning_note}"
    # Mark off-day attendance if today is not a configured workday and not a holiday
    try:
        workdays = ws.workdays or []
        weekday = local_now.weekday()
        is_workday = weekday in workdays
        if (not is_workday) and (not is_holiday):
            att.note = (att.note or 'Off-day attendance')
    except Exception:
        pass
    # Optional employee note for lateness
    try:
        employee_note = (data.get('employee_note') or '').strip()
        if minutes_late > 0 and employee_note:
            att.employee_note = employee_note
    except Exception:
        pass
    # Try link to employee
    try:
        att.employee = user.employee
    except Exception:
        pass
    att.save()
    return JsonResponse(AttendanceSerializer(att).data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def attendance_check_out(request):
    user = request.user
    now = dj_timezone.now()
    ws, _ = WorkSettings.objects.get_or_create()
    tz = ZoneInfo(ws.timezone or dj_timezone.get_current_timezone_name())
    local_now = now.astimezone(tz)
    date_local = local_now.date()
    att = Attendance.objects.filter(user=user, date_local=date_local).first()
    if not att or not att.check_in_at_utc:
        return JsonResponse({'detail': 'No active attendance for today'}, status=400)
    if att.check_out_at_utc:
        return JsonResponse({'detail': 'Already checked out'}, status=400)

    data = request.data if hasattr(request, 'data') else {}
    try:
        lat = float(data.get('lat'))
        lng = float(data.get('lng'))
        acc = int(data.get('accuracy_m') or 0)
        client_ip = data.get('client_ip', 'unknown')
    except Exception:
        return JsonResponse({'detail': 'Invalid location payload'}, status=400)

    # Geofence validation (optional at checkout; still enforce by default)
    if ws.office_latitude is None or ws.office_longitude is None:
        return JsonResponse({'detail': 'Office location is not configured'}, status=400)
    dist = haversine_meters(float(ws.office_latitude), float(ws.office_longitude), lat, lng)
    within = dist <= float(ws.office_radius_meters or 100)
    
    # Allow check-out even if outside geofence, but mark it
    # if not within:
    #     return JsonResponse({'detail': 'Outside office radius'}, status=403)

    att.check_out_at_utc = now
    att.check_out_lat = lat
    att.check_out_lng = lng
    att.check_out_accuracy_m = acc
    att.check_out_ip = client_ip
    
    # Add warning note if outside geofence
    if not within:
        warning_note = f"Check-out di luar area kantor (jarak: {dist:.0f}m dari radius {ws.office_radius_meters}m)"
        if att.note:
            att.note = f"{att.note}; {warning_note}"
        else:
            att.note = warning_note
    # Compute total work minutes
    delta = now - (att.check_in_at_utc or now)
    att.total_work_minutes = max(0, int(delta.total_seconds() // 60))
    
    # Calculate overtime based on total work duration
    weekday = local_now.weekday()  # Monday=0..Sunday=6
    
    # Get required minutes based on weekday
    if weekday == 4:  # Friday
        required_minutes = int(ws.friday_required_minutes or 240)  # Default 4 jam
    else:  # Monday-Thursday
        required_minutes = int(ws.required_minutes or 480)  # Default 8 jam
    
    # Overtime calculation: total work minutes - required minutes - threshold
    overtime_threshold = int(ws.overtime_threshold_minutes or 60)  # Default 1 hour buffer
    if att.total_work_minutes > (required_minutes + overtime_threshold):
        att.overtime_minutes = att.total_work_minutes - required_minutes - overtime_threshold
        
        # Calculate overtime amount if employee has salary
        if att.employee and att.employee.gaji_pokok:
            # Calculate hourly wage (assuming monthly salary)
            monthly_hours = 22 * 8  # 22 workdays * 8 hours per day
            hourly_wage = float(att.employee.gaji_pokok) / monthly_hours
            
            # Determine overtime rate
            if att.is_holiday:
                rate = float(ws.overtime_rate_holiday or 0.75)
            else:
                rate = float(ws.overtime_rate_workday or 0.50)
            
            # Calculate overtime amount
            overtime_hours = att.overtime_minutes / 60
            att.overtime_amount = overtime_hours * hourly_wage * rate
        else:
            att.overtime_amount = 0
            
        # For now, all overtime needs manual approval
        att.overtime_approved = False
        att.overtime_approved_by = None
        att.overtime_approved_at = None
        
    else:
        # No overtime
        att.overtime_minutes = 0
        att.overtime_amount = 0
        att.overtime_approved = False
        att.overtime_approved_by = None
        att.overtime_approved_at = None
    
    # Optional employee note for underwork (less than required minutes)
    try:
        employee_note = (data.get('employee_note') or '').strip()
        if att.total_work_minutes < required_minutes and employee_note:
            att.employee_note = employee_note
    except Exception:
        pass
    
    att.save(update_fields=[
        'check_out_at_utc', 'check_out_lat', 'check_out_lng', 
        'check_out_accuracy_m', 'total_work_minutes', 'employee_note',
        'overtime_minutes', 'overtime_amount', 'overtime_approved',
        'overtime_approved_by', 'overtime_approved_at'
    ])
    return JsonResponse(AttendanceSerializer(att).data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def attendance_report(request):
    """Get detailed attendance report with summary statistics for the current user."""
    user = request.user
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    month = request.GET.get('month')  # Format: YYYY-MM
    
    # Build attendance queryset
    attendance_qs = Attendance.objects.filter(user=user)
    
    # Apply date filters
    # Priority: month filter overrides date range filters
    if month:
        # Parse month and filter by year-month
        try:
            year, month_num = month.split('-')
            attendance_qs = attendance_qs.filter(
                date_local__year=year,
                date_local__month=month_num
            )
            # Log the filter being applied
            print(f"Applied month filter: {year}-{month_num}")
        except ValueError:
            return JsonResponse({"detail": "invalid_month_format_use_yyyy_mm"}, status=400)
    else:
        # Apply date range filters (can be partial - just start or just end)
        if start_date:
            attendance_qs = attendance_qs.filter(date_local__gte=start_date)
            print(f"Applied start_date filter: {start_date}")
        if end_date:
            attendance_qs = attendance_qs.filter(date_local__lte=end_date)
            print(f"Applied end_date filter: {end_date}")
    
    # Get attendance records
    attendances = attendance_qs.order_by('-date_local')
    
    # Calculate detailed statistics
    total_days = attendances.count()
    present_days = attendances.filter(check_in_at_utc__isnull=False).count()
    late_days = attendances.filter(minutes_late__gt=0).count()
    absent_days = total_days - present_days
    total_late_minutes = attendances.filter(minutes_late__gt=0).aggregate(
        total=models.Sum('minutes_late')
    )['total'] or 0
    total_work_minutes = attendances.filter(
        total_work_minutes__gt=0
    ).aggregate(
        total=models.Sum('total_work_minutes')
    )['total'] or 0
    
    # Get all attendance records for the period
    attendance_records = [
        {
            "id": att.id,
            "date_local": str(att.date_local),
            "check_in_at_utc": att.check_in_at_utc.isoformat() if att.check_in_at_utc else None,
            "check_out_at_utc": att.check_out_at_utc.isoformat() if att.check_out_at_utc else None,
            "check_in_lat": float(att.check_in_lat) if att.check_in_lat else None,
            "check_in_lng": float(att.check_in_lng) if att.check_in_lng else None,
            "check_out_lat": float(att.check_out_lat) if att.check_out_lat else None,
            "check_out_lng": float(att.check_out_lng) if att.check_out_lng else None,
            "check_in_ip": att.check_in_ip,
            "check_out_ip": att.check_out_ip,
            "minutes_late": att.minutes_late,
            "total_work_minutes": att.total_work_minutes,
            "is_holiday": att.is_holiday,
            "within_geofence": att.within_geofence,
            "note": att.note,
            "employee_note": att.employee_note,
            "created_at": att.created_at.isoformat(),
            "updated_at": att.updated_at.isoformat(),
        }
        for att in attendances
    ]
    
    return JsonResponse({
        "summary": {
            "total_days": total_days,
            "present_days": present_days,
            "late_days": late_days,
            "absent_days": absent_days,
            "attendance_rate": round((present_days / total_days * 100) if total_days > 0 else 0, 2),
            "total_late_minutes": total_late_minutes,
            "total_work_minutes": total_work_minutes,
            "average_work_minutes": round(total_work_minutes / present_days, 2) if present_days > 0 else 0
        },
        "attendance_records": attendance_records,
        "filters": {
            "start_date": start_date,
            "end_date": end_date,
            "month": month
        }
    }, safe=False)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def attendance_report_pdf(request):
    """Generate PDF report for attendance data."""
    user = request.user
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    month = request.GET.get('month')
    
    # Validate date formats
    if start_date:
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
        except ValueError:
            return JsonResponse({"detail": "Invalid start_date format. Use YYYY-MM-DD"}, status=400)
    
    if end_date:
        try:
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
        except ValueError:
            return JsonResponse({"detail": "Invalid end_date format. Use YYYY-MM-DD"}, status=400)
    
    # Validate that end_date is not before start_date
    if start_date and end_date:
        if end_dt < start_dt:
            return JsonResponse({"detail": "end_date cannot be before start_date"}, status=400)
    
    # Build attendance queryset
    attendance_qs = Attendance.objects.filter(user=user)
    
    # Apply date filters
    # Priority: month filter overrides date range filters
    if month:
        try:
            year, month_num = month.split('-')
            attendance_qs = attendance_qs.filter(
                date_local__year=year,
                date_local__month=month_num
            )
        except ValueError:
            return JsonResponse({"detail": "invalid_month_format_use_yyyy_mm"}, status=400)
    else:
        # Apply date range filters (can be partial - just start or just end)
        if start_date:
            attendance_qs = attendance_qs.filter(date_local__gte=start_date)
        if end_date:
            attendance_qs = attendance_qs.filter(date_local__lte=end_date)
    
    # Get attendance records
    attendances = attendance_qs.order_by('-date_local')
    
    # Calculate statistics
    total_days = attendances.count()
    present_days = attendances.filter(check_in_at_utc__isnull=False).count()
    late_days = attendances.filter(minutes_late__gt=0).count()
    absent_days = total_days - present_days
    total_late_minutes = attendances.filter(minutes_late__gt=0).aggregate(
        total=models.Sum('minutes_late')
    )['total'] or 0
    total_work_minutes = attendances.filter(
        total_work_minutes__gt=0
    ).aggregate(
        total=models.Sum('total_work_minutes')
    )['total'] or 0
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    # Title
    title = Paragraph("Laporan Absensi Pegawai", title_style)
    elements.append(title)
    
    # Employee Info
    try:
        employee = user.employee
        employee_info = f"""
        <b>Nama:</b> {user.first_name} {user.last_name}<br/>
        <b>NIP:</b> {employee.nip if employee else 'N/A'}<br/>
        <b>Divisi:</b> {employee.division.name if employee and employee.division else 'N/A'}<br/>
        <b>Jabatan:</b> {employee.position.name if employee and employee.position else 'N/A'}<br/>
        <b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}
        """
        if month:
            employee_info += f"<br/><b>Bulan:</b> {month}"
    except:
        employee_info = f"<b>Nama:</b> {user.first_name} {user.last_name}<br/><b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}"
    
    employee_para = Paragraph(employee_info, styles['Normal'])
    elements.append(employee_para)
    elements.append(Spacer(1, 20))
    
    # Summary Statistics
    summary_title = Paragraph("Ringkasan Statistik", styles['Heading2'])
    elements.append(summary_title)
    elements.append(Spacer(1, 10))
    
    summary_data = [
        ['Total Hari', 'Hadir', 'Terlambat', 'Tidak Hadir', 'Tingkat Kehadiran'],
        [
            str(total_days),
            str(present_days),
            str(late_days),
            str(absent_days),
            f"{round((present_days / total_days * 100) if total_days > 0 else 0, 2)}%"
        ]
    ]
    
    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))
    
    # Work Hours Summary
    work_hours_title = Paragraph("Ringkasan Jam Kerja", styles['Heading2'])
    elements.append(work_hours_title)
    elements.append(Spacer(1, 10))
    
    work_hours_data = [
        ['Total Jam Kerja', 'Total Keterlambatan', 'Rata-rata Jam Kerja/Hari'],
        [
            format_work_hours(total_work_minutes),
            format_work_hours(total_late_minutes),
            format_work_hours(total_work_minutes / present_days) if present_days > 0 else '0m'
        ]
    ]
    
    work_hours_table = Table(work_hours_data)
    work_hours_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(work_hours_table)
    elements.append(Spacer(1, 20))
    
    # Detailed Records
    if attendances.exists():
        records_title = Paragraph("Detail Absensi", styles['Heading2'])
        elements.append(records_title)
        elements.append(Spacer(1, 10))
        
        # Prepare table data
        table_data = [['Tanggal', 'Status', 'Check-in', 'Check-out', 'Terlambat', 'Jam Kerja']]
        
        for att in attendances[:50]:  # Limit to 50 records to avoid PDF too large
            # Determine status
            if att.is_holiday:
                status = 'Hari Libur'
            elif att.minutes_late > 0:
                status = f'Terlambat {att.minutes_late}m'
            elif att.check_in_at_utc and att.check_out_at_utc:
                status = 'Hadir Lengkap'
            elif att.check_in_at_utc:
                status = 'Hanya Check-in'
            else:
                status = 'Tidak Hadir'
            
            # Format times
            check_in = att.check_in_at_utc.strftime('%H:%M') if att.check_in_at_utc else '-'
            check_out = att.check_out_at_utc.strftime('%H:%M') if att.check_out_at_utc else '-'
            
            # Format work hours
            work_hours = format_work_hours(att.total_work_minutes)
            
            table_data.append([
                att.date_local.strftime('%d/%m/%Y'),
                status,
                check_in,
                check_out,
                f"{att.minutes_late}m" if att.minutes_late > 0 else '-',
                work_hours
            ])
        
        # Create table
        records_table = Table(table_data, colWidths=[1*inch, 1.5*inch, 0.8*inch, 0.8*inch, 0.8*inch, 1*inch])
        records_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(records_table)
        
        if attendances.count() > 50:
            elements.append(Spacer(1, 10))
            note = Paragraph(f"<i>Catatan: Ditampilkan 50 record terbaru dari total {attendances.count()} record</i>", styles['Normal'])
            elements.append(note)
    
    # Footer
    elements.append(Spacer(1, 30))
    footer = Paragraph(f"<i>Dibuat pada: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</i>", styles['Normal'])
    elements.append(footer)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    # Create response
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="attendance-report-{datetime.now().strftime("%Y%m%d-%H%M%S")}.pdf"'
    
    return response


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def supervisor_team_attendance_pdf(request):
    """Generate PDF report for supervisor team attendance."""
    user = request.user
    
    # Check if user is supervisor or admin
    if not (user.is_superuser or user.groups.filter(name__in=['admin', 'supervisor']).exists()):
        return HttpResponse("Forbidden", status=403)
    
    # Get supervisor's division
    try:
        supervisor_division_id = user.employee.division_id
        if not supervisor_division_id:
            return HttpResponse("Supervisor division not configured", status=400)
    except Exception:
        return HttpResponse("Supervisor employee record not found", status=400)
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    employee_id = request.GET.get('employee_id')
    
    # Build base queryset for team members in same division
    team_employees = Employee.objects.filter(
        division_id=supervisor_division_id
    ).select_related('user', 'division', 'position')
    
    # Filter by specific employee if requested
    if employee_id:
        team_employees = team_employees.filter(id=employee_id)
    
    # Get attendance data for team members
    attendance_data = []
    
    for employee in team_employees:
        # Get attendance records for this employee
        attendance_qs = Attendance.objects.filter(
            user=employee.user
        ).select_related('user', 'user__employee', 'user__employee__division', 'user__employee__position')
        
        # Apply date filters
        if start_date:
            attendance_qs = attendance_qs.filter(date_local__gte=start_date)
        if end_date:
            attendance_qs = attendance_qs.filter(date_local__lte=end_date)
        
        # Get attendance records
        attendances = attendance_qs.order_by('-date_local')
        
        # Calculate summary statistics
        total_days = attendances.count()
        present_days = attendances.filter(check_in_at_utc__isnull=False).count()
        late_days = attendances.filter(minutes_late__gt=0).count()
        absent_days = total_days - present_days
        
        employee_data = {
            "employee": employee,
            "summary": {
                "total_days": total_days,
                "present_days": present_days,
                "late_days": late_days,
                "absent_days": absent_days,
                "attendance_rate": round((present_days / total_days * 100) if total_days > 0 else 0, 2)
            },
            "attendances": attendances
        }
        
        attendance_data.append(employee_data)
    
    try:
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Title
        title = Paragraph("Laporan Absensi Tim Supervisor", title_style)
        elements.append(title)
        
        # Supervisor Info
        try:
            supervisor_info = f"""
            <b>Supervisor:</b> {user.first_name} {user.last_name}<br/>
            <b>Divisi:</b> {user.employee.division.name if user.employee and user.employee.division else 'N/A'}<br/>
            <b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}
            """
            if employee_id:
                supervisor_info += f"<br/><b>Employee ID:</b> {employee_id}"
        except:
            supervisor_info = f"<b>Supervisor:</b> {user.first_name} {user.last_name}<br/><b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}"
        
        supervisor_para = Paragraph(supervisor_info, styles['Normal'])
        elements.append(supervisor_para)
        elements.append(Spacer(1, 20))
        
        # Team Summary Statistics
        summary_title = Paragraph("Ringkasan Tim", styles['Heading2'])
        elements.append(summary_title)
        elements.append(Spacer(1, 10))
        
        total_members = len(attendance_data)
        total_present_days = sum(member['summary']['present_days'] for member in attendance_data)
        total_late_days = sum(member['summary']['late_days'] for member in attendance_data)
        avg_attendance_rate = sum(member['summary']['attendance_rate'] for member in attendance_data) / total_members if total_members > 0 else 0
        
        summary_data = [
            ['Total Anggota Tim', 'Total Hari Hadir', 'Total Hari Terlambat', 'Rata-rata Tingkat Kehadiran'],
            [
                str(total_members),
                str(total_present_days),
                str(total_late_days),
                f"{round(avg_attendance_rate, 2)}%"
            ]
        ]
        
        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 20))
        
        # Individual Employee Summary
        if attendance_data:
            employee_title = Paragraph("Ringkasan per Anggota Tim", styles['Heading2'])
            elements.append(employee_title)
            elements.append(Spacer(1, 10))
            
            # Prepare table data
            table_data = [['Nama', 'NIP', 'Divisi', 'Jabatan', 'Total Hari', 'Hadir', 'Terlambat', 'Rate (%)']]
            
            for member in attendance_data:
                employee = member['employee']
                summary = member['summary']
                
                table_data.append([
                    f"{employee.user.first_name} {employee.user.last_name}",
                    employee.nip,
                    employee.division.name if employee.division else '-',
                    employee.position.name if employee.position else '-',
                    str(summary['total_days']),
                    str(summary['present_days']),
                    str(summary['late_days']),
                    f"{summary['attendance_rate']}%"
                ])
            
            # Create table
            employee_table = Table(table_data, colWidths=[1.2*inch, 0.8*inch, 1*inch, 1*inch, 0.7*inch, 0.6*inch, 0.7*inch, 0.6*inch])
            employee_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(employee_table)
            elements.append(Spacer(1, 20))
            
            # Recent Attendance Records (first 3 employees only to avoid PDF too large)
            if len(attendance_data) <= 3:
                records_title = Paragraph("Detail Absensi Terbaru", styles['Heading2'])
                elements.append(records_title)
                elements.append(Spacer(1, 10))
                
                for member in attendance_data:
                    employee = member['employee']
                    attendances = member['attendances'][:7]  # Last 7 days
                    
                    if attendances.exists():
                        employee_name = f"{employee.user.first_name} {employee.user.last_name}"
                        employee_header = Paragraph(f"<b>{employee_name} ({employee.nip})</b>", styles['Heading3'])
                        elements.append(employee_header)
                        elements.append(Spacer(1, 5))
                        
                        # Prepare records table
                        records_data = [['Tanggal', 'Status', 'Check-in', 'Check-out', 'Jam Kerja', 'Keterlambatan']]
                        
                        for att in attendances:
                            # Determine status
                            if att.is_holiday:
                                status = 'Hari Libur'
                            elif att.minutes_late > 0:
                                status = f'Terlambat {att.minutes_late}m'
                            elif att.check_in_at_utc and att.check_out_at_utc:
                                status = 'Hadir Lengkap'
                            elif att.check_in_at_utc:
                                status = 'Hanya Check-in'
                            else:
                                status = 'Tidak Hadir'
                            
                            # Format times
                            check_in = att.check_in_at_utc.strftime('%H:%M') if att.check_in_at_utc else '-'
                            check_out = att.check_out_at_utc.strftime('%H:%M') if att.check_out_at_utc else '-'
                            
                            # Format work hours
                            work_hours = format_work_hours(att.total_work_minutes)
                            
                            records_data.append([
                                att.date_local.strftime('%d/%m/%Y'),
                                status,
                                check_in,
                                check_out,
                                work_hours,
                                f"{att.minutes_late}m" if att.minutes_late > 0 else '-'
                            ])
                        
                        # Create records table
                        records_table = Table(records_data, colWidths=[0.8*inch, 1.2*inch, 0.7*inch, 0.7*inch, 0.8*inch, 0.8*inch])
                        records_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 8),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('FONTSIZE', (0, 1), (-1, -1), 7),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ]))
                        elements.append(records_table)
                        elements.append(Spacer(1, 15))
        
        # Footer
        elements.append(Spacer(1, 30))
        footer = Paragraph(f"<i>Dibuat pada: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</i>", styles['Normal'])
        elements.append(footer)
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Create response with proper headers
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="supervisor-team-attendance-{datetime.now().strftime("%Y%m%d-%H%M%S")}.pdf"'
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        return response
        
    except Exception as e:
        # Log the error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error generating PDF: {str(e)}")
        
        # Return error response
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)


def supervisor_team_attendance_pdf_alt(request):
    """Alternative PDF generation function without DRF decorators to avoid Accept header issues."""
    from django.contrib.auth.decorators import login_required
    from django.views.decorators.http import require_http_methods
    
    # Check if user is authenticated
    if not request.user.is_authenticated:
        return HttpResponse("Unauthorized", status=401)
    
    # Check if user is supervisor or admin
    if not (request.user.is_superuser or request.user.groups.filter(name__in=['admin', 'supervisor']).exists()):
        return HttpResponse("Forbidden", status=403)
    
    user = request.user
    
    # Get supervisor's division
    try:
        supervisor_division_id = user.employee.division_id
        if not supervisor_division_id:
            return HttpResponse("Supervisor division not configured", status=400)
    except Exception:
        return HttpResponse("Supervisor employee record not found", status=400)
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    employee_id = request.GET.get('employee_id')
    
    # Build base queryset for team members in same division
    team_employees = Employee.objects.filter(
        division_id=supervisor_division_id
    ).select_related('user', 'division', 'position')
    
    # Filter by specific employee if requested
    if employee_id:
        team_employees = team_employees.filter(id=employee_id)
    
    # Get attendance data for team members
    attendance_data = []
    
    for employee in team_employees:
        # Get attendance records for this employee
        attendance_qs = Attendance.objects.filter(
            user=employee.user
        ).select_related('user', 'user__employee', 'user__employee__division', 'user__employee__position')
        
        # Apply date filters
        if start_date:
            attendance_qs = attendance_qs.filter(date_local__gte=start_date)
        if end_date:
            attendance_qs = attendance_qs.filter(date_local__lte=end_date)
        
        # Get attendance records
        attendances = attendance_qs.order_by('-date_local')
        
        # Calculate summary statistics
        total_days = attendances.count()
        present_days = attendances.filter(check_in_at_utc__isnull=False).count()
        late_days = attendances.filter(minutes_late__gt=0).count()
        absent_days = total_days - present_days
        
        employee_data = {
            "employee": employee,
            "summary": {
                "total_days": total_days,
                "present_days": present_days,
                "late_days": late_days,
                "absent_days": absent_days,
                "attendance_rate": round((present_days / total_days * 100) if total_days > 0 else 0, 2)
            },
            "attendances": attendances
        }
        
        attendance_data.append(employee_data)
    
    try:
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Title
        title = Paragraph("Laporan Absensi Tim Supervisor", title_style)
        elements.append(title)
        
        # Supervisor Info
        try:
            supervisor_info = f"""
            <b>Supervisor:</b> {user.first_name} {user.last_name}<br/>
            <b>Divisi:</b> {user.employee.division.name if user.employee and user.employee.division else 'N/A'}<br/>
            <b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}
            """
            if employee_id:
                supervisor_info += f"<br/><b>Employee ID:</b> {employee_id}"
        except:
            supervisor_info = f"<b>Supervisor:</b> {user.first_name} {user.last_name}<br/><b>Periode:</b> {start_date or 'Semua'} - {end_date or 'Semua'}"
        
        supervisor_para = Paragraph(supervisor_info, styles['Normal'])
        elements.append(supervisor_para)
        elements.append(Spacer(1, 20))
        
        # Team Summary Statistics
        summary_title = Paragraph("Ringkasan Tim", styles['Heading2'])
        elements.append(summary_title)
        elements.append(Spacer(1, 10))
        
        total_members = len(attendance_data)
        total_present_days = sum(member['summary']['present_days'] for member in attendance_data)
        total_late_days = sum(member['summary']['late_days'] for member in attendance_data)
        avg_attendance_rate = sum(member['summary']['attendance_rate'] for member in attendance_data) / total_members if total_members > 0 else 0
        
        summary_data = [
            ['Total Anggota Tim', 'Total Hari Hadir', 'Total Hari Terlambat', 'Rata-rata Tingkat Kehadiran'],
            [
                str(total_members),
                str(total_present_days),
                str(total_late_days),
                f"{round(avg_attendance_rate, 2)}%"
            ]
        ]
        
        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 20))
        
        # Individual Employee Summary
        if attendance_data:
            employee_title = Paragraph("Ringkasan per Anggota Tim", styles['Heading2'])
            elements.append(employee_title)
            elements.append(Spacer(1, 10))
            
            # Prepare table data
            table_data = [['Nama', 'NIP', 'Divisi', 'Jabatan', 'Total Hari', 'Hadir', 'Terlambat', 'Rate (%)']]
            
            for member in attendance_data:
                employee = member['employee']
                summary = member['summary']
                
                table_data.append([
                    f"{employee.user.first_name} {employee.user.last_name}",
                    employee.nip,
                    employee.division.name if employee.division else '-',
                    employee.position.name if employee.position else '-',
                    str(summary['total_days']),
                    str(summary['present_days']),
                    str(summary['late_days']),
                    f"{summary['attendance_rate']}%"
                ])
            
            # Create table
            employee_table = Table(table_data, colWidths=[1.2*inch, 0.8*inch, 1*inch, 1*inch, 0.7*inch, 0.6*inch, 0.7*inch, 0.6*inch])
            employee_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(employee_table)
            elements.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="attendance_report_{start_date or "all"}_{end_date or "all"}.pdf"'
        
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)


@extend_schema(
    tags=['Settings'],
    summary='Get work settings for employee',
    description='Get work settings that employees need for overtime calculation',
    responses={200: dict},
)
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def employee_work_settings(request):
    """Get work settings that employees need for overtime calculation"""
    try:
        # Get work settings
        work_settings, _ = WorkSettings.objects.get_or_create()
        
        # Return only the fields needed for overtime calculation
        settings_data = {
            'start_time': work_settings.start_time,
            'end_time': work_settings.end_time,
            'friday_start_time': work_settings.friday_start_time,
            'friday_end_time': work_settings.friday_end_time,
            'workdays': work_settings.workdays,
            'overtime_rate_workday': work_settings.overtime_rate_workday,
            'overtime_rate_holiday': work_settings.overtime_rate_holiday,
            'overtime_threshold_minutes': work_settings.overtime_threshold_minutes,
        }
        
        return Response(settings_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


# ===== OVERTIME APPROVAL FUNCTIONS =====

@extend_schema(
    tags=['Overtime'],
    summary='Approve overtime for attendance',
    description='Approve overtime for a specific attendance record',
    responses={200: dict, 400: dict, 403: dict, 404: dict},
)
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def approve_overtime(request, attendance_id):
    """Approve overtime for a specific attendance record"""
    try:
        # Get the attendance record
        attendance = Attendance.objects.get(id=attendance_id)
        
        # Check if user has permission to approve overtime
        user = request.user
        
        # Admin can approve any overtime
        if user.is_superuser or user.groups.filter(name='admin').exists():
            pass
        # Use position-based approval checking
        else:
            from .utils import ApprovalChecker
            
            # Check if user can approve organization-wide
            if ApprovalChecker.can_approve_organization_overtime(user, attendance.employee):
                pass  # Can approve
            # Check if user can approve division-level
            elif ApprovalChecker.can_approve_division_overtime(user, attendance.employee):
                pass  # Can approve
            else:
                return Response({"detail": "You don't have permission to approve overtime for this employee"}, status=403)
        
        # Check if overtime is already approved
        if attendance.overtime_approved:
            return Response({"detail": "Overtime is already approved"}, status=400)
        
        # Approve the overtime
        attendance.overtime_approved = True
        attendance.overtime_approved_by = user
        attendance.overtime_approved_at = dj_timezone.now()
        attendance.save()
        
        return Response({
            "detail": "Overtime approved successfully",
            "attendance_id": attendance.id,
            "approved_by": user.username,
            "approved_at": attendance.overtime_approved_at.isoformat()
        })
        
    except Attendance.DoesNotExist:
        return Response({"detail": "Attendance record not found"}, status=404)
    except Exception as e:
        return Response({"detail": str(e)}, status=500)


@extend_schema(
    tags=['Overtime'],
    summary='Get overtime report',
    description='Get overtime report with filtering options',
    responses={200: dict, 400: dict},
)
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def overtime_report(request):
    """Get overtime report with filtering options"""
    user = request.user
    
    # Get query parameters
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    month = request.GET.get('month')
    employee_id = request.GET.get('employee_id')
    approved_only = request.GET.get('approved_only', '').lower() == 'true'
    pending_only = request.GET.get('pending_only', '').lower() == 'true'
    
    # Build base queryset
    if user.is_superuser or user.groups.filter(name='admin').exists():
        # Admin can see all overtime
        qs = Attendance.objects.filter(overtime_minutes__gt=0)
    else:
        # Use position-based approval checking
        from .utils import ApprovalChecker
        approval_level = ApprovalChecker.get_user_approval_level(user)
        
        if approval_level >= 1:
            # Has supervisor capabilities - can see overtime for their division
            try:
                supervisor_division = user.employee.division
                qs = Attendance.objects.filter(
                    overtime_minutes__gt=0,
                    employee__division=supervisor_division
                )
            except Exception:
                return Response({"detail": "Division not configured"}, status=400)
        else:
            # Employee can only see their own overtime
            qs = Attendance.objects.filter(
                user=user,
                overtime_minutes__gt=0
            )
    
    # Apply filters
    if start_date:
        qs = qs.filter(date_local__gte=start_date)
    if end_date:
        qs = qs.filter(date_local__lte=end_date)
    if month:
        try:
            year, month_num = month.split('-')
            qs = qs.filter(
                date_local__year=year,
                date_local__month=month_num
            )
        except ValueError:
            return Response({"detail": "Invalid month format"}, status=400)
    if employee_id:
        qs = qs.filter(employee_id=employee_id)
    if approved_only:
        qs = qs.filter(overtime_approved=True)
    if pending_only:
        qs = qs.filter(overtime_approved=False)
    
    # Get overtime records
    overtime_records = []
    total_overtime_minutes = 0
    total_overtime_amount = 0
    pending_overtime_count = 0
    
    for att in qs.select_related('user', 'employee', 'employee__division', 'overtime_approved_by'):
        # Determine required minutes for this day
        weekday = att.date_local.weekday()
        if weekday == 4:  # Friday
            required_minutes = 240  # 4 jam
        else:  # Monday-Thursday
            required_minutes = 480  # 8 jam
        
        overtime_records.append({
            "id": att.id,
            "date": str(att.date_local),
            "weekday": ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu'][weekday],
            "employee": {
                "name": att.user.get_full_name() or att.user.username,
                "fullname": (att.employee.fullname if att.employee else None),
                "nip": att.employee.nip if att.employee else None,
                "division": att.employee.division.name if att.employee and att.employee.division else None
            },
            "check_in": att.check_in_at_utc.isoformat() if att.check_in_at_utc else None,
            "check_out": att.check_out_at_utc.isoformat() if att.check_out_at_utc else None,
            "total_work_minutes": att.total_work_minutes,
            "required_minutes": required_minutes,
            "overtime_minutes": att.overtime_minutes,
            "overtime_amount": float(att.overtime_amount),
            "is_holiday": att.is_holiday,
            "overtime_approved": att.overtime_approved,
            "approved_by": att.overtime_approved_by.username if att.overtime_approved_by else None,
            "approved_at": att.overtime_approved_at.isoformat() if att.overtime_approved_at else None
        })
        
        total_overtime_minutes += att.overtime_minutes
        total_overtime_amount += float(att.overtime_amount or 0)
        if not att.overtime_approved:
            pending_overtime_count += 1
    
    return Response({
        "summary": {
            "total_records": len(overtime_records),
            "total_overtime_minutes": total_overtime_minutes,
            "total_overtime_amount": total_overtime_amount,
            "pending_overtime_count": pending_overtime_count,
            "average_overtime_per_day": round(total_overtime_minutes / len(overtime_records), 2) if overtime_records else 0
        },
        "overtime_records": overtime_records,
        "filters": {
            "start_date": start_date,
            "end_date": end_date,
            "month": month,
            "employee_id": employee_id,
            "approved_only": approved_only,
            "pending_only": pending_only
        }
    })


# ===== OVERTIME REQUEST MANAGEMENT =====

class OvertimeRequestViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing overtime requests with role-based access control
    """
    permission_classes = [IsOvertimeRequestOwnerOrSupervisor]
    pagination_class = DefaultPagination
    
    def get_template_path(self):
        """Get the path to the active overtime template with cache busting"""
        import os
        from django.conf import settings
        from django.core.cache import cache
        
        template_dir = os.path.join(settings.BASE_DIR, 'template')
        
        # Check if we have cached template info
        cache_key = 'overtime_template_path'
        cached_result = cache.get(cache_key)
        
        # Get current template directory state
        current_files = []
        if os.path.exists(template_dir):
            for file in os.listdir(template_dir):
                if file.endswith('.docx') and not file.startswith('~$'):  # Exclude temp files
                    file_path = os.path.join(template_dir, file)
                    current_files.append((file, os.path.getmtime(file_path)))
        
        # Create a hash of current directory state
        current_state = hash(tuple(sorted(current_files)))
        
        # If cached result exists and directory hasn't changed, use cache
        if cached_result and cached_result.get('state') == current_state:
            return cached_result['path'], cached_result['name']
        
        # Directory has changed, recalculate template path
        if current_files:
            # Sort by modification time (newest first)
            current_files.sort(key=lambda x: x[1], reverse=True)
            
            # Priority order for template files
            priority_names = [
                'template_SURAT_PERINTAH_KERJA_LEMBUR.docx',  # Primary template
                'template_overtime_clean.docx',                 # Fallback 1
                'template_overtime_working.docx',               # Fallback 2
                'template_overtime_simple.docx',                # Fallback 3
            ]
            
            # First try priority names
            for priority_name in priority_names:
                for file_name, _ in current_files:
                    if file_name == priority_name:
                        template_path = os.path.join(template_dir, priority_name)
                        result = (template_path, priority_name)
                        # Cache the result
                        cache.set(cache_key, {
                            'path': template_path,
                            'name': priority_name,
                            'state': current_state
                        }, timeout=300)  # Cache for 5 minutes
                        return result
            
            # If no priority names found, use the most recently modified .docx file
            most_recent = current_files[0][0]
            template_path = os.path.join(template_dir, most_recent)
            result = (template_path, most_recent)
            # Cache the result
            cache.set(cache_key, {
                'path': template_path,
                'name': most_recent,
                'state': current_state
            }, timeout=300)  # Cache for 5 minutes
            return result
        
        # If no template found, return None
        return None, None

    def get_monthly_export_template_path(self):
        """Get the path to the active monthly export template with cache busting"""
        import os
        from django.conf import settings
        from django.core.cache import cache
        
        template_dir = os.path.join(settings.BASE_DIR, 'template')
        
        # Check if we have cached template info
        cache_key = 'monthly_export_template_path'
        cached_result = cache.get(cache_key)
        
        # Get current template directory state
        current_files = []
        if os.path.exists(template_dir):
            for file in os.listdir(template_dir):
                if file.endswith('.docx') and not file.startswith('~$'):  # Exclude temp files
                    file_path = os.path.join(template_dir, file)
                    current_files.append((file, os.path.getmtime(file_path)))
        
        # Create a hash of current directory state
        current_state = hash(tuple(sorted(current_files)))
        
        # If cached result exists and directory hasn't changed, use cache
        if cached_result and cached_result.get('state') == current_state:
            return cached_result['path'], cached_result['name']
        
        # Directory has changed, recalculate template path
        if current_files:
            # Sort by modification time (newest first)
            current_files.sort(key=lambda x: x[1], reverse=True)
            
            # Priority order for monthly export template files
            priority_names = [
                'template_rekap_lembur.docx',                   # Primary monthly export template (valid)
                'template_SURAT_PERINTAH_KERJA_LEMBUR.docx',    # Fallback 1 (valid)
                'template_monthly_overtime_export.docx',         # Fallback 2 (small file, may be corrupt)
                'template_monthly_overtime.docx',                # Fallback 3
                'template_monthly_export.docx',                  # Fallback 4
                'template_overtime_monthly.docx',                # Fallback 5
            ]
            
            # First try priority names
            for priority_name in priority_names:
                for file_name, _ in current_files:
                    if file_name == priority_name:
                        template_path = os.path.join(template_dir, priority_name)
                        result = (template_path, priority_name)
                        # Cache the result
                        cache.set(cache_key, {
                            'path': template_path,
                            'name': priority_name,
                            'state': current_state
                        }, timeout=300)  # Cache for 5 minutes
                        return result
            
            # If no priority names found, use the most recently modified .docx file
            most_recent = current_files[0][0]
            template_path = os.path.join(template_dir, most_recent)
            result = (template_path, most_recent)
            # Cache the result
            cache.set(cache_key, {
                'path': template_path,
                'name': most_recent,
                'state': current_state
            }, timeout=300)  # Cache for 5 minutes
            return result
        
        # If no template found, return None
        return None, None
    
    def get_queryset(self):
        """Filter queryset based on user role"""
        user = self.request.user
        
        # Admin can see all overtime requests
        if user.is_superuser or user.groups.filter(name='admin').exists():
            return OvertimeRequest.objects.all().select_related(
                'user', 'employee', 'approved_by', 'level1_approved_by', 'final_approved_by'
            )
        
        # Use position-based approval checking for supervisor capabilities
        else:
            from .utils import ApprovalChecker
            approval_level = ApprovalChecker.get_user_approval_level(user)
            
            if approval_level >= 1:
                try:
                    supervisor_employee = user.employee
                    
                    # Check if supervisor has org-wide approval permission
                    if ApprovalChecker.can_approve_overtime_org_wide(user):
                        # Org-wide supervisors can see all requests
                        return OvertimeRequest.objects.all().select_related(
                            'user', 'employee', 'approved_by', 'level1_approved_by', 'final_approved_by'
                        )
                    else:
                        # Division supervisors can only see requests from their division
                        return OvertimeRequest.objects.filter(
                            employee__division=supervisor_employee.division
                        ).select_related('user', 'employee', 'approved_by', 'level1_approved_by', 'final_approved_by')
                except:
                    return OvertimeRequest.objects.none()
            else:
                # No supervisor capabilities - can only see own requests
                return OvertimeRequest.objects.filter(user=user).select_related(
                    'user', 'employee', 'approved_by', 'level1_approved_by', 'final_approved_by'
                )
        
        return OvertimeRequest.objects.none()
    
    def get_serializer_class(self):
        """Return appropriate serializer based on user role and action"""
        user = self.request.user
        
        # For create action, use create serializer
        if self.action == 'create':
            return OvertimeRequestCreateSerializer
        
        # Role-based serializers for other actions
        if user.is_superuser or user.groups.filter(name='admin').exists():
            return OvertimeRequestAdminSerializer
        
        # Use position-based approval checking for supervisor capabilities
        from .utils import ApprovalChecker
        approval_level = ApprovalChecker.get_user_approval_level(user)
        
        if approval_level >= 1:
            return OvertimeRequestSupervisorSerializer
        else:
            return OvertimeRequestEmployeeSerializer
    
    def perform_create(self, serializer):
        """Auto-set user and employee when creating overtime request"""
        user = self.request.user
        try:
            employee = user.employee
            serializer.save(user=user, employee=employee)
        except AttributeError:
            raise serializers.ValidationError("User tidak memiliki profil employee")
        except IntegrityError as e:
            if 'unique constraint' in str(e).lower() or 'duplicate' in str(e).lower():
                raise serializers.ValidationError(
                    "Anda sudah memiliki pengajuan rekap untuk periode dan jenis laporan yang sama. "
                    "Silakan gunakan pengajuan yang sudah ada atau tunggu hingga selesai."
                )
            raise serializers.ValidationError(f"Terjadi kesalahan database: {str(e)}")
        except Exception as e:
            raise serializers.ValidationError(f"Terjadi kesalahan: {str(e)}")
    
    def create(self, request, *args, **kwargs):
        """Override create to ensure only employees can create requests"""
        if not request.user.groups.filter(name='pegawai').exists():
            return Response(
                {"detail": "Hanya pegawai yang dapat mengajukan lembur"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Check if user already has a pending request for the same date
        date_requested = request.data.get('date_requested')
        if date_requested:
            existing_request = OvertimeRequest.objects.filter(
                user=request.user,
                date_requested=date_requested,
                status='pending'
            ).exists()
            
            if existing_request:
                return Response(
                    {"detail": "Anda sudah memiliki pengajuan lembur yang pending untuk tanggal tersebut"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        return super().create(request, *args, **kwargs)
    
    @action(detail=True, methods=['post'], permission_classes=[IsOvertimeRequestApprover])
    def approve(self, request, pk=None):
        """Approve overtime request with 2-level approval system"""
        overtime_request = self.get_object()
        user = request.user
        
        # Check if user is admin (can bypass 2-level approval)
        is_admin = user.is_superuser or user.groups.filter(name='admin').exists()
        
        # Use position-based approval checking
        from .utils import ApprovalChecker
        
        # Check if user is org-wide supervisor
        is_org_wide_supervisor = ApprovalChecker.can_approve_organization_overtime(user, overtime_request.employee)
        # Check if user has no approval permission (level 0)
        has_no_approval = ApprovalChecker.get_user_approval_level(user) == 0
        
        # Reject if user has no approval permission
        if has_no_approval:
            return Response(
                {"detail": "Anda tidak memiliki permission untuk melakukan approval"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Admin can approve any status
        if is_admin:
            if overtime_request.status in ['rejected']:
                return Response(
                    {"detail": "Pengajuan yang sudah ditolak tidak dapat disetujui"},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Admin approval goes straight to final approval
            overtime_request.status = 'approved'
            overtime_request.final_approved_by = user
            overtime_request.final_approved_at = dj_timezone.now()
            overtime_request.approved_by = user  # Legacy field
            overtime_request.approved_at = dj_timezone.now()  # Legacy field
            overtime_request.rejection_reason = None
            overtime_request.save()

            # Generate and save DOCX document
            try:
                self._generate_and_save_docx(overtime_request)
            except Exception as e:
                # Log error but don't fail the approval
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Failed to generate DOCX for overtime request {overtime_request.id}: {str(e)}")

            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil disetujui (Final Approval)",
                "data": serializer.data
            })
        
        # Org-wide supervisor (final approval)
        elif is_org_wide_supervisor:
            if overtime_request.status != 'level1_approved':
                if overtime_request.status == 'pending':
                    return Response(
                        {"detail": "Pengajuan harus disetujui oleh supervisor divisi terlebih dahulu (Level 1 Approval)"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
                else:
                    return Response(
                        {"detail": "Hanya pengajuan dengan status Level 1 Approved yang dapat disetujui final"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
            
            overtime_request.status = 'approved'
            overtime_request.final_approved_by = user
            overtime_request.final_approved_at = dj_timezone.now()
            overtime_request.approved_by = user  # Legacy field
            overtime_request.approved_at = dj_timezone.now()  # Legacy field
            overtime_request.rejection_reason = None
            overtime_request.save()

            # Generate and save DOCX document
            try:
                self._generate_and_save_docx(overtime_request)
            except Exception as e:
                # Log error but don't fail the approval
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Failed to generate DOCX for overtime request {overtime_request.id}: {str(e)}")

            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil disetujui (Final Approval)",
                "data": serializer.data
            })
        
        # Division supervisor (level 1 approval)
        else:
            if overtime_request.status != 'pending':
                return Response(
                    {"detail": "Hanya pengajuan dengan status pending yang dapat disetujui level 1"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            overtime_request.status = 'level1_approved'
            overtime_request.level1_approved_by = user
            overtime_request.level1_approved_at = dj_timezone.now()
            overtime_request.rejection_reason = None
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil disetujui (Level 1 Approval). Menunggu approval final.",
                "data": serializer.data
            })
    
    @action(detail=True, methods=['post'], permission_classes=[IsOvertimeRequestApprover])
    def reject(self, request, pk=None):
        """Reject overtime request with 2-level rejection system"""
        overtime_request = self.get_object()
        user = request.user
        
        # Check if user is admin (can bypass 2-level approval)
        is_admin = user.is_superuser or user.groups.filter(name='admin').exists()
        
        # Use position-based approval checking
        from .utils import ApprovalChecker
        
        # Check if user is org-wide supervisor
        is_org_wide_supervisor = ApprovalChecker.can_approve_organization_overtime(user, overtime_request.employee)
        # Check if user has no approval permission (level 0)
        has_no_approval = ApprovalChecker.get_user_approval_level(user) == 0
        
        # Reject if user has no approval permission
        if has_no_approval:
            return Response(
                {"detail": "Anda tidak memiliki permission untuk melakukan rejection"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Can reject pending or level1_approved requests
        if overtime_request.status not in ['pending', 'level1_approved']:
            return Response(
                {"detail": "Hanya pengajuan dengan status pending atau level1_approved yang dapat ditolak"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        rejection_reason = request.data.get('rejection_reason', '')
        if not rejection_reason.strip():
            return Response(
                {"detail": "Alasan penolakan harus diisi"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        overtime_request.status = 'rejected'
        overtime_request.rejection_reason = rejection_reason
        
        # Admin can reject any status and it goes to final rejection
        if is_admin:
            overtime_request.final_rejected_by = user
            overtime_request.final_rejected_at = dj_timezone.now()
            # Clear all approval fields
            overtime_request.level1_approved_by = None
            overtime_request.level1_approved_at = None
            overtime_request.final_approved_by = None
            overtime_request.final_approved_at = None
            overtime_request.approved_by = None  # Legacy field
            overtime_request.approved_at = None  # Legacy field
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil ditolak (Final Rejection)",
                "data": serializer.data
            })
        
        # Org-wide supervisor (final rejection)
        elif is_org_wide_supervisor:
            overtime_request.final_rejected_by = user
            overtime_request.final_rejected_at = dj_timezone.now()
            # Clear all approval fields
            overtime_request.level1_approved_by = None
            overtime_request.level1_approved_at = None
            overtime_request.final_approved_by = None
            overtime_request.final_approved_at = None
            overtime_request.approved_by = None  # Legacy field
            overtime_request.approved_at = None  # Legacy field
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil ditolak (Final Rejection)",
                "data": serializer.data
            })
        
        # Division supervisor (level 1 rejection)
        else:
            overtime_request.level1_rejected_by = user
            overtime_request.level1_rejected_at = dj_timezone.now()
            # Clear all approval fields
            overtime_request.level1_approved_by = None
            overtime_request.level1_approved_at = None
            overtime_request.final_approved_by = None
            overtime_request.final_approved_at = None
            overtime_request.approved_by = None  # Legacy field
            overtime_request.approved_at = None  # Legacy field
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response({
                "detail": "Pengajuan lembur berhasil ditolak (Level 1 Rejection)",
                "data": serializer.data
            })

    def _generate_and_save_docx(self, overtime_request):
        """Generate and save DOCX document for approved overtime request"""
        from .models import OvertimeDocument
        import os
        import tempfile
        from django.core.files.base import ContentFile

        try:
            # Generate DOCX content using existing logic
            docx_content = self._generate_docx_content(overtime_request)

            # Create filename
            filename = f"Surat_Perintah_Kerja_Lembur_{overtime_request.id}_{overtime_request.employee.user.username}_{overtime_request.final_approved_at.strftime('%Y%m%d_%H%M%S')}.docx"

            # Create or update OvertimeDocument
            doc, created = OvertimeDocument.objects.get_or_create(
                overtime_request=overtime_request,
                document_type='individual',
                defaults={
                    'status': 'generated',
                }
            )

            # Save the DOCX file
            doc.docx_file.save(filename, ContentFile(docx_content), save=True)
            doc.status = 'generated'
            doc.downloaded_at = None  # Reset download timestamp
            doc.save()

            # Update the overtime request reference
            overtime_request.docx_document = doc
            overtime_request.save(update_fields=['docx_document'])

        except Exception as e:
            # Create error record if generation fails
            OvertimeDocument.objects.update_or_create(
                overtime_request=overtime_request,
                document_type='individual',
                defaults={
                    'status': 'error',
                    'error_message': str(e),
                }
            )
            raise

    def _generate_docx_content(self, overtime_request):
        """Generate DOCX content using existing template logic"""
        # Import required libraries
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from io import BytesIO
        from django.utils import timezone
        from django.conf import settings
        import os
        import locale

        # Set locale for Indonesian formatting
        try:
            locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
        except:
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID')
            except:
                pass

        # Get template path
        template_path, template_name = self.get_template_path()

        if not template_path:
            raise Exception("Template overtime tidak ditemukan")

        # Load template document
        doc = Document(template_path)

        # Function to replace text in paragraphs
        def replace_text_in_paragraphs(paragraphs, old_text, new_text):
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

        # Function to replace text in tables
        def replace_text_in_tables(tables, old_text, new_text):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

        # Get current date and time
        current_date = dj_timezone.now()
        current_year = current_date.strftime('%Y')
        current_month = current_date.strftime('%B')
        current_day = current_date.strftime('%d')

        # Format overtime date
        overtime_date = overtime_request.date_requested
        overtime_date_formatted = overtime_date.strftime('%d %B %Y')

        # Get employee information
        employee_name = overtime_request.employee.fullname or overtime_request.employee.user.get_full_name() or overtime_request.employee.user.username
        employee_nip = overtime_request.employee.nip or '-'
        employee_position = overtime_request.employee.position.name if overtime_request.employee.position else '-'
        employee_division = overtime_request.employee.division.name if overtime_request.employee.division else '-'

        # Format NIP variations
        nip_9_digit = employee_nip[:9] if employee_nip and len(employee_nip) >= 9 else employee_nip
        nip_18_digit = employee_nip if employee_nip and len(employee_nip) >= 18 else (employee_nip + '0' * (18 - len(employee_nip))) if employee_nip else '-'

        # Get approval information
        def get_approver_name(user):
            """Get approver name from employee.fullname with fallback to username"""
            if not user:
                return '-'
            if hasattr(user, 'employee') and user.employee and user.employee.fullname:
                return user.employee.fullname
            full_name = user.get_full_name()
            if full_name and full_name.strip():
                return full_name
            return user.username.title()

        def get_approver_nip(user):
            """Get approver NIP with fallback"""
            if not user or not hasattr(user, 'employee') or not user.employee:
                return '-'
            return user.employee.nip or '-'

        level1_approver = get_approver_name(overtime_request.level1_approved_by)
        level1_approver_nip = get_approver_nip(overtime_request.level1_approved_by)
        level1_approval_date = overtime_request.level1_approved_at.strftime('%d %B %Y %H:%M') if overtime_request.level1_approved_at else '-'

        final_approver = get_approver_name(overtime_request.final_approved_by)
        final_approver_nip = get_approver_nip(overtime_request.final_approved_by)
        final_approval_date = overtime_request.final_approved_at.strftime('%d %B %Y %H:%M') if overtime_request.final_approved_at else '-'

        # Format overtime amount
        overtime_amount = f"{overtime_request.overtime_amount}"

        # Define replacement mappings
        replacements = {
            # Document info
            '{{NOMOR_DOKUMEN}}': f"{overtime_request.id}/SPKL/KJRI-DXB/{current_year}",
            '{{TANGGAL_DOKUMEN}}': current_date.strftime('%d %B %Y'),
            '{{TAHUN}}': current_year,
            '{{BULAN}}': current_month,
            '{{HARI}}': current_day,

            # Employee info
            '{{NAMA_PEGAWAI}}': employee_name,
            '{{NIP_PEGAWAI}}': employee_nip,
            '{{NIP}}': employee_nip,
            '{{NIP_LENGKAP}}': nip_18_digit,
            '{{NIP_18_DIGIT}}': nip_18_digit,
            '{{NIP_9_DIGIT}}': nip_9_digit,
            '{{JABATAN_PEGAWAI}}': employee_position,
            '{{DIVISI_PEGAWAI}}': employee_division,

            # Overtime details
            '{{TANGGAL_LEMBUR}}': overtime_date_formatted,
            '{{JAM_LEMBUR}}': f"{overtime_request.overtime_hours} jam",
            '{{DESKRIPSI_PEKERJAAN}}': overtime_request.work_description,
            '{{JUMLAH_GAJI_LEMBUR}}': overtime_amount,

            # Approval info
            '{{LEVEL1_APPROVER}}': level1_approver,
            '{{LEVEL1_APPROVER_NIP}}': level1_approver_nip,
            '{{LEVEL1_APPROVAL_DATE}}': level1_approval_date,
            '{{FINAL_APPROVER}}': final_approver,
            '{{FINAL_APPROVER_NIP}}': final_approver_nip,
            '{{FINAL_APPROVAL_DATE}}': final_approval_date,

            # Company info
            '{{NAMA_PERUSAHAAN}}': 'KJRI DUBAI',
            '{{ALAMAT_PERUSAHAAN}}': 'KONSULAT JENDERAL REPUBLIK INDONESIA DI DUBAI',
            '{{LOKASI}}': 'Dubai',
        }

        # Apply replacements
        for placeholder, value in replacements.items():
            replace_text_in_paragraphs(doc.paragraphs, placeholder, value)
            replace_text_in_tables(doc.tables, placeholder, value)

        # Save to buffer and return content
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get overtime request summary statistics"""
        queryset = self.get_queryset()
        
        # Filter by date range if provided
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        
        if start_date:
            queryset = queryset.filter(date_requested__gte=start_date)
        if end_date:
            queryset = queryset.filter(date_requested__lte=end_date)
        
        # Calculate statistics
        total_requests = queryset.count()
        pending_requests = queryset.filter(status='pending').count()
        level1_approved_requests = queryset.filter(status='level1_approved').count()
        approved_requests = queryset.filter(status='approved').count()
        rejected_requests = queryset.filter(status='rejected').count()
        
        total_hours = queryset.filter(status='approved').aggregate(
            total=models.Sum('overtime_hours')
        )['total'] or 0
        
        total_amount = queryset.filter(status='approved').aggregate(
            total=models.Sum('overtime_amount')
        )['total'] or 0
        
        return Response({
            'total_requests': total_requests,
            'pending_requests': pending_requests,
            'level1_approved_requests': level1_approved_requests,
            'approved_requests': approved_requests,
            'rejected_requests': rejected_requests,
            'total_approved_hours': float(total_hours),
            'total_approved_amount': float(total_amount),
        })
    
    @action(detail=False, methods=['get'])
    def potential_overtime(self, request):
        """Get attendance records that could be submitted as overtime requests"""
        user = self.request.user
        
        # Only employees can view their potential overtime
        if not user.groups.filter(name='pegawai').exists():
            return Response(
                {"detail": "Hanya pegawai yang dapat melihat potensi lembur"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            employee = user.employee
        except:
            return Response(
                {"detail": "User tidak memiliki profil employee"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get date range (default: last 30 days)
        from datetime import date, timedelta
        end_date = date.today()
        start_date = end_date - timedelta(days=30)
        
        # Allow custom date range
        start_param = request.query_params.get('start_date')
        end_param = request.query_params.get('end_date')
        
        if start_param:
            try:
                start_date = date.fromisoformat(start_param)
            except ValueError:
                pass
                
        if end_param:
            try:
                end_date = date.fromisoformat(end_param)
            except ValueError:
                pass
        
        # Get attendance records with potential overtime
        attendance_records = Attendance.objects.filter(
            user=user,
            date_local__gte=start_date,
            date_local__lte=end_date,
            check_out_at_utc__isnull=False,  # Only completed attendance
            total_work_minutes__gt=0
        ).select_related('user').order_by('-date_local')
        
        # Get work settings for overtime threshold calculation
        try:
            ws = WorkSettings.objects.first()
            overtime_threshold = int(ws.overtime_threshold_minutes or 60) if ws else 60
        except:
            overtime_threshold = 60
        
        # Get existing overtime requests to avoid duplicates
        existing_requests = set(
            OvertimeRequest.objects.filter(
                user=user,
                date_requested__gte=start_date,
                date_requested__lte=end_date
            ).values_list('date_requested', flat=True)
        )
        
        potential_records = []
        
        for att in attendance_records:
            # Skip if already has overtime request for this date
            if att.date_local in existing_requests:
                continue
            
            # Determine required minutes for the day
            weekday = att.date_local.weekday()
            if weekday == 4:  # Friday
                required_minutes = int(ws.friday_required_minutes or 240) if ws else 240
            else:
                required_minutes = int(ws.required_minutes or 480) if ws else 480
            
            # Calculate potential overtime (total work - required - threshold)
            potential_overtime_minutes = att.total_work_minutes - required_minutes - overtime_threshold
            
            # Only include if there's potential overtime (worked more than required + threshold)
            if potential_overtime_minutes > 0:
                potential_overtime_hours = potential_overtime_minutes / 60
                
                # Calculate potential overtime amount
                potential_amount = 0
                if employee.gaji_pokok and ws:
                    monthly_hours = 22 * 8
                    hourly_wage = float(employee.gaji_pokok) / monthly_hours
                    
                    # Determine rate
                    if att.is_holiday:
                        rate = float(ws.overtime_rate_holiday or 0.75)
                    else:
                        rate = float(ws.overtime_rate_workday or 0.50)
                    
                    potential_amount = potential_overtime_hours * hourly_wage * rate
                
                # Format times
                check_in_time = None
                check_out_time = None
                
                if att.check_in_at_utc:
                    check_in_time = att.check_in_at_utc.strftime('%H:%M')
                if att.check_out_at_utc:
                    check_out_time = att.check_out_at_utc.strftime('%H:%M')
                
                potential_records.append({
                    'date_local': att.date_local.isoformat(),
                    'weekday': att.date_local.strftime('%A'),
                    'check_in_time': check_in_time,
                    'check_out_time': check_out_time,
                    'total_work_minutes': att.total_work_minutes,
                    'required_minutes': required_minutes,
                    'overtime_threshold_minutes': overtime_threshold,
                    'potential_overtime_minutes': potential_overtime_minutes,
                    'potential_overtime_hours': round(potential_overtime_hours, 2),
                    'potential_overtime_amount': round(potential_amount, 2),
                    'is_holiday': att.is_holiday,
                    'within_geofence': att.within_geofence,
                    'can_submit': True,  # All records here are eligible for submission
                })
        
        return Response({
            'start_date': start_date.isoformat(),
            'end_date': end_date.isoformat(),
            'overtime_threshold_minutes': overtime_threshold,
            'total_potential_records': len(potential_records),
            'total_potential_hours': sum(r['potential_overtime_hours'] for r in potential_records),
            'total_potential_amount': sum(r['potential_overtime_amount'] for r in potential_records),
            'potential_records': potential_records,
        })

    @action(detail=False, methods=['get'])
    def preview_template(self, request):
        """Preview overtime request template with sample data"""
        try:
            # Import required libraries
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            from io import BytesIO
            from django.utils import timezone
            from django.conf import settings
            import os
            import locale
            
            # Set locale for Indonesian formatting
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_ALL, 'id_ID')
                except:
                    pass
            
            # Get template path dynamically
            template_path, template_name = self.get_template_path()
            
            # Check if template exists
            if not template_path:
                return Response(
                    {"detail": "Template overtime tidak ditemukan. Silakan upload template di admin settings."}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Load template document
            doc = Document(template_path)
            
            # Function to replace text in paragraphs
            def replace_text_in_paragraphs(paragraphs, old_text, new_text):
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
            
            # Function to replace text in tables
            def replace_text_in_tables(tables, old_text, new_text):
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
            
            # Get current date and time
            current_date = dj_timezone.now()
            current_year = current_date.strftime('%Y')
            current_month = current_date.strftime('%B')
            current_day = current_date.strftime('%d')
            
            # Sample data for preview
            sample_data = {
                # Document info
                '{{NOMOR_DOKUMEN}}': f"001/SPKL/KJRI-DXB/{current_year}",
                '{{TANGGAL_DOKUMEN}}': current_date.strftime('%d %B %Y'),
                '{{TAHUN}}': current_year,
                '{{BULAN}}': current_month,
                '{{HARI}}': current_day,
                
                # Employee info
                '{{NAMA_PEGAWAI}}': 'CONTOH NAMA PEGAWAI',
                '{{NIP_PEGAWAI}}': '123456789',
                '{{NIP}}': '123456789',
                '{{NIP_LENGKAP}}': '123456789012345678',
                '{{NIP_18_DIGIT}}': '123456789012345678',
                '{{NIP_9_DIGIT}}': '123456789',
                '{{JABATAN_PEGAWAI}}': 'STAFF',
                '{{DIVISI_PEGAWAI}}': 'IT',
                
                # Overtime details
                '{{TANGGAL_LEMBUR}}': current_date.strftime('%d %B %Y'),
                '{{JAM_LEMBUR}}': '2 jam',
                '{{DESKRIPSI_PEKERJAAN}}': 'Contoh deskripsi pekerjaan lembur untuk maintenance server dan backup data',
                '{{JUMLAH_GAJI_LEMBUR}}': '50000',
                
                # Approval info
                '{{LEVEL1_APPROVER}}': 'SUPERVISOR DIVISI',
                '{{LEVEL1_APPROVER_NIP}}': '198501012010012001',
                '{{LEVEL1_APPROVAL_DATE}}': current_date.strftime('%d %B %Y %H:%M'),
                '{{FINAL_APPROVER}}': 'MANAGER',
                '{{FINAL_APPROVER_NIP}}': '198001011990012001',
                '{{FINAL_APPROVAL_DATE}}': current_date.strftime('%d %B %Y %H:%M'),
                
                # Company info
                '{{NAMA_PERUSAHAAN}}': 'KJRI DUBAI',
                '{{ALAMAT_PERUSAHAAN}}': 'KONSULAT JENDERAL REPUBLIK INDONESIA DI DUBAI',
                '{{LOKASI}}': 'Dubai',
            }
            
            # Replace all placeholders with sample data
            for placeholder, value in sample_data.items():
                replace_text_in_paragraphs(doc.paragraphs, placeholder, value)
                replace_text_in_tables(doc.tables, placeholder, value)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Return document response
            from django.http import HttpResponse
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="Preview_Template_Overtime_{current_date.strftime("%Y%m%d")}.docx"'
            return response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal preview template: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['post'])
    def upload_template(self, request):
        """Upload new overtime template"""
        try:
            # Check if user is admin
            if not request.user.is_staff:
                return Response(
                    {"detail": "Hanya admin yang dapat upload template"}, 
                    status=status.HTTP_403_FORBIDDEN
                )
            
            # Get template file from request
            template_file = request.FILES.get('template')
            if not template_file:
                return Response(
                    {"detail": "Template file tidak ditemukan"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate file type
            if not template_file.name.endswith('.docx'):
                return Response(
                    {"detail": "File harus berformat .docx"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate file size (max 10MB)
            if template_file.size > 10 * 1024 * 1024:
                return Response(
                    {"detail": "Ukuran file maksimal 10MB"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Import required libraries
            import os
            from django.conf import settings
            
            # Path to template directory
            template_dir = os.path.join(settings.BASE_DIR, 'template')
            if not os.path.exists(template_dir):
                os.makedirs(template_dir)
            
            # Save new template with original filename (but ensure .docx extension)
            original_filename = template_file.name
            if not original_filename.endswith('.docx'):
                original_filename += '.docx'
            
            # Save as the new primary template
            template_path = os.path.join(template_dir, original_filename)
            
            with open(template_path, 'wb+') as destination:
                for chunk in template_file.chunks():
                    destination.write(chunk)
            
            # Clear template cache to force reload
            from django.core.cache import cache
            cache.delete('overtime_template_path')
            
            return Response({
                "detail": "Template berhasil diupload",
                "filename": original_filename,
                "size": template_file.size,
                "path": template_path,
                "message": f"Template '{original_filename}' berhasil disimpan dan akan digunakan sebagai template utama. Cache template telah di-clear."
            })
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal upload template: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['post'], permission_classes=[IsAdminUser])
    def reload_template(self, request):
        """Force reload template cache (admin only)"""
        try:
            from django.core.cache import cache
            cache.delete('overtime_template_path')
            
            # Get current template info
            template_path, template_name = self.get_template_path()
            
            return Response({
                "detail": "Template cache berhasil di-clear",
                "current_template": template_name,
                "message": f"Template '{template_name}' akan digunakan untuk request selanjutnya"
            })
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal reload template: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['post'], permission_classes=[IsAdminUser])
    def upload_monthly_export_template(self, request):
        """Upload new monthly export template for overtime"""
        try:
            # Check if user is admin
            if not request.user.is_staff:
                return Response(
                    {"detail": "Hanya admin yang dapat upload template"}, 
                    status=status.HTTP_403_FORBIDDEN
                )
            
            # Get template file from request
            template_file = request.FILES.get('template')
            if not template_file:
                return Response(
                    {"detail": "Template file tidak ditemukan"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate file type
            if not template_file.name.endswith('.docx'):
                return Response(
                    {"detail": "File harus berformat .docx"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate file size (max 10MB)
            if template_file.size > 10 * 1024 * 1024:
                return Response(
                    {"detail": "Ukuran file maksimal 10MB"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Import required libraries
            import os
            from django.conf import settings
            
            # Path to template directory
            template_dir = os.path.join(settings.BASE_DIR, 'template')
            if not os.path.exists(template_dir):
                os.makedirs(template_dir)
            
            # Save new template with original filename (but ensure .docx extension)
            original_filename = template_file.name
            if not original_filename.endswith('.docx'):
                original_filename += '.docx'
            
            # Save as the new monthly export template
            template_path = os.path.join(template_dir, original_filename)
            
            with open(template_path, 'wb+') as destination:
                for chunk in template_file.chunks():
                    destination.write(chunk)
            
            # Clear template cache to force reload
            from django.core.cache import cache
            cache.delete('monthly_export_template_path')
            
            return Response({
                "detail": "Template monthly export berhasil diupload",
                "filename": original_filename,
                "size": template_file.size,
                "path": template_path,
                "message": f"Template '{original_filename}' berhasil disimpan dan akan digunakan sebagai template monthly export. Cache template telah di-clear."
            })
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal upload template monthly export: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['post'], permission_classes=[IsAdminUser])
    def reload_monthly_export_template(self, request):
        """Force reload monthly export template cache (admin only)"""
        try:
            from django.core.cache import cache
            cache.delete('monthly_export_template_path')
            
            # Get current template info
            template_path, template_name = self.get_monthly_export_template_path()
            
            return Response({
                "detail": "Template monthly export cache berhasil di-clear",
                "current_template": template_name,
                "message": f"Template '{template_name}' akan digunakan untuk export monthly overtime selanjutnya"
            })
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal reload template monthly export: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download overtime request document - uses stored file if available, otherwise generates new"""
        overtime_request = self.get_object()

        # Only allow download for approved overtime requests
        if overtime_request.status != 'approved':
            return Response(
                {"detail": "Hanya overtime yang sudah disetujui yang dapat didownload"},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # First, try to use stored DOCX file if available
            if overtime_request.docx_document and overtime_request.docx_document.docx_file:
                try:
                    # Update download timestamp
                    overtime_request.docx_document.downloaded_at = dj_timezone.now()
                    overtime_request.docx_document.save(update_fields=['downloaded_at'])

                    # Return stored file
                    with open(overtime_request.docx_document.docx_file.path, 'rb') as f:
                        file_content = f.read()

                    from django.http import HttpResponse
                    response = HttpResponse(file_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename="{overtime_request.docx_document.docx_file.name.split("/")[-1]}"'
                    return response

                except (FileNotFoundError, OSError) as e:
                    # File not found or corrupted, fall back to generating new one
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Stored DOCX file not found for overtime request {overtime_request.id}: {str(e)}")

            # If no stored file or file is missing, generate new DOCX
            # Generate and save the document first
            self._generate_and_save_docx(overtime_request)

            # Now return the newly generated file
            if overtime_request.docx_document and overtime_request.docx_document.docx_file:
                # Update download timestamp
                overtime_request.docx_document.downloaded_at = dj_timezone.now()
                overtime_request.docx_document.save(update_fields=['downloaded_at'])

                with open(overtime_request.docx_document.docx_file.path, 'rb') as f:
                    file_content = f.read()

                from django.http import HttpResponse
                response = HttpResponse(file_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{overtime_request.docx_document.docx_file.name.split("/")[-1]}"'
                return response

            # Fallback: generate on-the-fly if saving failed
            return self._generate_docx_on_fly(overtime_request)

        except Exception as e:
            return Response(
                {"detail": f"Gagal generate dokumen: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _generate_docx_on_fly(self, overtime_request):
        """Generate DOCX document on-the-fly (fallback method)"""
        # Import required libraries
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from io import BytesIO
        # timezone already imported as dj_timezone at the top
        from django.conf import settings
        import os
        import locale

        # Set locale for Indonesian formatting
        try:
            locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
        except:
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID')
            except:
                pass

        # Get template path dynamically
        template_path, template_name = self.get_template_path()

        # Check if template exists
        if not template_path:
            from django.http import HttpResponse
            response = HttpResponse("Template overtime tidak ditemukan", status=404)
            return response

        # Load template document
        doc = Document(template_path)

        # Function to replace text in paragraphs
        def replace_text_in_paragraphs(paragraphs, old_text, new_text):
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

        # Function to replace text in tables
        def replace_text_in_tables(tables, old_text, new_text):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

        # Get current date and time
        current_date = dj_timezone.now()
        current_year = current_date.strftime('%Y')
        current_month = current_date.strftime('%B')
        current_day = current_date.strftime('%d')

        # Format overtime date
        overtime_date = overtime_request.date_requested
        overtime_date_formatted = overtime_date.strftime('%d %B %Y')

        # Get employee information
        employee_name = overtime_request.employee.fullname or overtime_request.employee.user.get_full_name() or overtime_request.employee.user.username
        employee_nip = overtime_request.employee.nip or '-'
        employee_position = overtime_request.employee.position.name if overtime_request.employee.position else '-'
        employee_division = overtime_request.employee.division.name if overtime_request.employee.division else '-'

        # Format NIP variations
        nip_9_digit = employee_nip[:9] if employee_nip and len(employee_nip) >= 9 else employee_nip
        nip_18_digit = employee_nip if employee_nip and len(employee_nip) >= 18 else (employee_nip + '0' * (18 - len(employee_nip))) if employee_nip else '-'

        # Get approval information
        def get_approver_name(user):
            """Get approver name from employee.fullname with fallback to username"""
            if not user:
                return '-'
            if hasattr(user, 'employee') and user.employee and user.employee.fullname:
                return user.employee.fullname
            full_name = user.get_full_name()
            if full_name and full_name.strip():
                return full_name
            return user.username.title()

        def get_approver_nip(user):
            """Get approver NIP with fallback"""
            if not user or not hasattr(user, 'employee') or not user.employee:
                return '-'
            return user.employee.nip or '-'

        level1_approver = get_approver_name(overtime_request.level1_approved_by)
        level1_approver_nip = get_approver_nip(overtime_request.level1_approved_by)
        level1_approval_date = overtime_request.level1_approved_at.strftime('%d %B %Y %H:%M') if overtime_request.level1_approved_at else '-'

        final_approver = get_approver_name(overtime_request.final_approved_by)
        final_approver_nip = get_approver_nip(overtime_request.final_approved_by)
        final_approval_date = overtime_request.final_approved_at.strftime('%d %B %Y %H:%M') if overtime_request.final_approved_at else '-'

        # Format overtime amount
        overtime_amount = f"{overtime_request.overtime_amount}"

        # Define replacement mappings
        replacements = {
            # Document info
            '{{NOMOR_DOKUMEN}}': f"{overtime_request.id}/SPKL/KJRI-DXB/{current_year}",
            '{{TANGGAL_DOKUMEN}}': current_date.strftime('%d %B %Y'),
            '{{TAHUN}}': current_year,
            '{{BULAN}}': current_month,
            '{{HARI}}': current_day,

            # Employee info
            '{{NAMA_PEGAWAI}}': employee_name,
            '{{NIP_PEGAWAI}}': employee_nip,
            '{{NIP}}': employee_nip,
            '{{NIP_LENGKAP}}': nip_18_digit,
            '{{NIP_18_DIGIT}}': nip_18_digit,
            '{{NIP_9_DIGIT}}': nip_9_digit,
            '{{JABATAN_PEGAWAI}}': employee_position,
            '{{DIVISI_PEGAWAI}}': employee_division,

            # Overtime details
            '{{TANGGAL_LEMBUR}}': overtime_date_formatted,
            '{{JAM_LEMBUR}}': f"{overtime_request.overtime_hours} jam",
            '{{DESKRIPSI_PEKERJAAN}}': overtime_request.work_description,
            '{{JUMLAH_GAJI_LEMBUR}}': overtime_amount,

            # Approval info
            '{{LEVEL1_APPROVER}}': level1_approver,
            '{{LEVEL1_APPROVER_NIP}}': level1_approver_nip,
            '{{LEVEL1_APPROVAL_DATE}}': level1_approval_date,
            '{{FINAL_APPROVER}}': final_approver,
            '{{FINAL_APPROVER_NIP}}': final_approver_nip,
            '{{FINAL_APPROVAL_DATE}}': final_approval_date,

            # Company info
            '{{NAMA_PERUSAHAAN}}': 'KJRI DUBAI',
            '{{ALAMAT_PERUSAHAAN}}': 'KONSULAT JENDERAL REPUBLIK INDONESIA DI DUBAI',
            '{{LOKASI}}': 'Dubai',
        }

        # Apply replacements
        for placeholder, value in replacements.items():
            replace_text_in_paragraphs(doc.paragraphs, placeholder, value)
            replace_text_in_tables(doc.tables, placeholder, value)

        # Save to buffer and return response
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Surat_Perintah_Kerja_Lembur_{overtime_request.id}_{overtime_request.employee.user.username}.docx"'
        return response


    @action(detail=False, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def export_monthly_docx(self, request):
        """Export monthly overtime data to DOCX using template with dynamic table"""
        try:
            # Get export parameters
            month = request.query_params.get('month')  # Format: YYYY-MM
            employee_id = request.query_params.get('employee_id')
            
            if not month:
                return Response(
                    {"detail": "Parameter month (YYYY-MM) diperlukan"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate month format
            try:
                year, month_num = month.split('-')
                year = int(year)
                month_num = int(month_num)
                if month_num < 1 or month_num > 12:
                    raise ValueError
            except ValueError:
                return Response(
                    {"detail": "Format month harus YYYY-MM (contoh: 2025-01)"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Get employee (current user or specified employee)
            if employee_id and request.user.is_superuser:
                try:
                    employee = Employee.objects.get(id=employee_id)
                except Employee.DoesNotExist:
                    return Response(
                        {"detail": "Employee tidak ditemukan"}, 
                        status=status.HTTP_404_NOT_FOUND
                    )
            else:
                try:
                    employee = request.user.employee
                except:
                    return Response(
                        {"detail": "User tidak memiliki profil employee"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
            
            # Get overtime data for the month
            start_date = datetime.date(year, month_num, 1)
            if month_num == 12:
                end_date = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
            else:
                end_date = datetime.date(year, month_num + 1, 1) - datetime.timedelta(days=1)
            
            overtime_records = OvertimeRequest.objects.filter(
                employee=employee,
                date_requested__gte=start_date,
                date_requested__lte=end_date
            ).order_by('date_requested')
            
            if not overtime_records.exists():
                return Response(
                    {"detail": f"Tidak ada data overtime untuk periode {month}"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Generate DOCX document
            docx_file = self.generate_monthly_export_docx(
                overtime_records, 
                employee, 
                month, 
                start_date, 
                end_date
            )
            
            # Return file response
            from django.http import FileResponse
            import os
            
            filename = f"Laporan_Overtime_{employee.nip}_{month}.docx"
            
            response = FileResponse(
                open(docx_file, 'rb'),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            # Clean up temporary file
            os.unlink(docx_file)
            
            return response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate export: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def export_list_pdf(self, request):
        """Export list of overtime requests to PDF"""
        try:
            print(f"Debug - export_list_pdf called")
            print(f"Debug - User: {request.user.username}")
            print(f"Debug - User groups: {[g.name for g in request.user.groups.all()]}")

            # Get filter parameters
            month = request.query_params.get('month')
            status_filter = request.query_params.get('status')
            employee_id = request.query_params.get('employee_id')

            print(f"Debug - Filters: month={month}, status_filter={status_filter}, employee_id={employee_id}")

            # Get queryset based on user role and filters
            queryset = self.get_queryset()
            print(f"Debug - Initial queryset count: {queryset.count()}")

            # Apply filters
            if month:
                try:
                    print(f"Debug - Applying month filter: {month}")
                    year, month_num = month.split('-')
                    year = int(year)
                    month_num = int(month_num)
                    start_date = datetime(year, month_num, 1).date()
                    if month_num == 12:
                        end_date = datetime(year + 1, 1, 1).date() - timedelta(days=1)
                    else:
                        end_date = datetime(year, month_num + 1, 1).date() - timedelta(days=1)
                    print(f"Debug - Date range: {start_date} to {end_date}")
                    queryset = queryset.filter(date_requested__range=[start_date, end_date])
                    print(f"Debug - After month filter count: {queryset.count()}")
                except ValueError:
                    return Response(
                        {"detail": "Format bulan tidak valid. Gunakan format YYYY-MM"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
            
            if status_filter:
                queryset = queryset.filter(status=status_filter)
            
            if employee_id:
                queryset = queryset.filter(employee_id=employee_id)
            
            # Get overtime requests
            overtime_requests = queryset.order_by('-date_requested', '-created_at')
            
            if not overtime_requests.exists():
                return Response(
                    {"detail": "Tidak ada data pengajuan lembur yang ditemukan"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Debug: Check the data
            print(f"Debug - Found {overtime_requests.count()} overtime requests")
            for req in overtime_requests:
                print(f"Debug - ID: {req.id}, Date: {req.date_requested}, Employee: {req.employee.fullname if req.employee else 'No Employee'}")
                break
            
            # Debug: Check overtime_requests
            print(f"Debug - Before PDF generation:")
            print(f"Debug - overtime_requests count: {overtime_requests.count()}")
            print(f"Debug - overtime_requests type: {type(overtime_requests)}")

            # Generate PDF directly
            pdf_response = self.create_overtime_list_pdf(overtime_requests, month)

            return pdf_response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate export: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
            
            if status:
                queryset = queryset.filter(status=status)
            
            if employee_id and request.user.is_superuser:
                try:
                    employee = Employee.objects.get(id=employee_id)
                    queryset = queryset.filter(employee=employee)
                except Employee.DoesNotExist:
                    return Response(
                        {"detail": "Employee tidak ditemukan"}, 
                        status=status.HTTP_404_NOT_FOUND
                    )
            
            # Order by date
            queryset = queryset.order_by('-date_requested')
            
            if not queryset.exists():
                return Response(
                    {"detail": "Tidak ada data pengajuan lembur yang sesuai dengan filter"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Generate PDF for list
            pdf_response = self.generate_overtime_list_pdf(queryset, request.user)
            
            return pdf_response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate PDF: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=False, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def export_monthly_pdf(self, request):
        """Export monthly overtime data directly to PDF (no DOCX conversion)"""
        try:
            # Get export parameters
            month = request.query_params.get('month')  # Format: YYYY-MM
            employee_id = request.query_params.get('employee_id')
            
            if not month:
                return Response(
                    {"detail": "Parameter month (YYYY-MM) diperlukan"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Validate month format
            try:
                year, month_num = month.split('-')
                year = int(year)
                month_num = int(month_num)
                if month_num < 1 or month_num > 12:
                    raise ValueError
            except ValueError:
                return Response(
                    {"detail": "Format month harus YYYY-MM (contoh: 2025-01)"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Get employee (current user or specified employee)
            if employee_id and request.user.is_superuser:
                try:
                    employee = Employee.objects.get(id=employee_id)
                except Employee.DoesNotExist:
                    return Response(
                        {"detail": "Employee tidak ditemukan"}, 
                        status=status.HTTP_404_NOT_FOUND
                    )
            else:
                try:
                    employee = request.user.employee
                except:
                    return Response(
                        {"detail": "User tidak memiliki profil employee"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
            
            # Get overtime data for the month
            start_date = datetime.date(year, month_num, 1)
            if month_num == 12:
                end_date = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
            else:
                end_date = datetime.date(year, month_num + 1, 1) - datetime.timedelta(days=1)
            
            overtime_records = OvertimeRequest.objects.filter(
                employee=employee,
                date_requested__gte=start_date,
                date_requested__lte=end_date
            ).order_by('date_requested')
            
            if not overtime_records.exists():
                return Response(
                    {"detail": f"Tidak ada data overtime untuk periode {month}"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Debug: Check the data types
            print(f"Debug - overtime_records count: {overtime_records.count()}")
            for record in overtime_records:
                print(f"Debug - record.date_requested type: {type(record.date_requested)}, value: {record.date_requested}")
                print(f"Debug - record.overtime_hours type: {type(record.overtime_hours)}, value: {record.overtime_hours}")
                print(f"Debug - record.overtime_amount type: {type(record.overtime_amount)}, value: {record.overtime_amount}")
                break  # Just check first record
            
            # Generate PDF directly
            pdf_response = self.generate_monthly_overtime_pdf(
                overtime_records, 
                employee, 
                month, 
                start_date, 
                end_date
            )
            
            return pdf_response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate PDF: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'], permission_classes=[permissions.AllowAny])
    def export_pdf(self, request, pk=None):
        """
        Export individual overtime request to PDF with improved workflow:
        1. Generate DOCX if not exists
        2. Save DOCX to media/overtime_docx and database
        3. Convert DOCX to PDF using DOCX converter service
        4. Save PDF to media/overtime_pdf and database
        5. Return PDF for download
        """
        # Handle AnonymousUser for AllowAny permission
        if request.user.is_anonymous:
            try:
                overtime_request = OvertimeRequest.objects.get(pk=pk)
            except OvertimeRequest.DoesNotExist:
                return Response(
                    {"detail": "Overtime request tidak ditemukan."},
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            overtime_request = self.get_object()

        # Only allow export for approved overtime requests
        if overtime_request.status != 'approved':
            return Response(
                {"detail": "Overtime request must be approved to export PDF."},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # Step 1 & 2: Generate and save DOCX if not exists
            if not overtime_request.docx_document or not overtime_request.docx_document.docx_file:
                # Generate and save DOCX
                self._generate_and_save_docx(overtime_request)

            # Verify DOCX exists after generation
            if not overtime_request.docx_document or not overtime_request.docx_document.docx_file:
                return Response(
                    {"detail": "Gagal generate DOCX file."},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            # Get the DOCX file path
            docx_file_path = overtime_request.docx_document.docx_file.path

            # Check if file exists on disk
            import os
            if not os.path.exists(docx_file_path):
                # File doesn't exist, regenerate DOCX
                print(f"DOCX file not found at {docx_file_path}, regenerating...")
                self._generate_and_save_docx(overtime_request)

                # Re-check after regeneration
                if not overtime_request.docx_document or not overtime_request.docx_document.docx_file:
                    return Response(
                        {"detail": "Gagal regenerate DOCX file."},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )

                # Update file path after regeneration
                docx_file_path = overtime_request.docx_document.docx_file.path

                # Final check
                if not os.path.exists(docx_file_path):
                    return Response(
                        {"detail": "File DOCX masih tidak ditemukan setelah regenerasi."},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )

            # Step 3 & 4: Convert DOCX to PDF using DOCX converter service
            pdf_content = self._convert_docx_to_pdf_via_service(docx_file_path)

            if not pdf_content:
                return Response(
                    {"detail": "Gagal convert DOCX ke PDF."},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            # Step 5: Save PDF to database
            from .models import OvertimeDocument
            from django.core.files.base import ContentFile

            # Update existing OvertimeDocument or create new one
            doc = overtime_request.docx_document
            if not doc:
                doc = OvertimeDocument.objects.create(
                    overtime_request=overtime_request,
                    document_type='individual',
                    status='converted'
                )

            # Create filename for PDF
            pdf_filename = f"Surat_Perintah_Lembur_{overtime_request.employee.fullname}_{overtime_request.date_requested}.pdf"

            # Save PDF file
            doc.pdf_file.save(pdf_filename, ContentFile(pdf_content), save=True)
            doc.status = 'converted'
            doc.converted_at = dj_timezone.now()
            doc.save()

            # Update overtime request reference if needed
            if not overtime_request.docx_document:
                overtime_request.docx_document = doc
                overtime_request.save(update_fields=['docx_document'])

            # Step 6: Return PDF for download
            from django.http import HttpResponse
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'

            return response

        except Exception as e:
            return Response(
                {"detail": f"Gagal export PDF: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _convert_docx_to_pdf_via_service(self, docx_file_path):
        """
        Convert DOCX to PDF using the DOCX converter service
        Returns PDF content as bytes
        """
        import requests
        import os
        from django.conf import settings

        try:
            # Prepare the file for upload
            with open(docx_file_path, 'rb') as f:
                files = {'file': ('document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}

                # Send POST request to DOCX converter service
                # Use internal network URL for better performance
                converter_url = "http://docx_converter:5000/convert"

                data = {'method': 'file'}
                response = requests.post(converter_url, files=files, data=data, timeout=60)

                if response.status_code == 200:
                    return response.content
                else:
                    # Log error for debugging
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.error(f"DOCX converter service error: {response.status_code} - {response.text}")
                    return None

        except requests.exceptions.RequestException as e:
            # Log network error
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"DOCX converter service network error: {str(e)}")
            return None
        except Exception as e:
            # Log general error
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"DOCX to PDF conversion error: {str(e)}")
            return None

    def generate_monthly_export_docx(self, overtime_records, employee, month, start_date, end_date):
        """Generate monthly overtime export DOCX with dynamic table"""
        try:
            # Import required libraries
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            from io import BytesIO
            from django.utils import timezone
            import os
            import locale
            import tempfile
            
            # Set locale for Indonesian formatting
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_ALL, 'id_ID')
                except:
                    pass
            
            # Get monthly export template path (priority) or fallback to regular template
            template_path, template_name = self.get_monthly_export_template_path()
            
            if not template_path:
                # Try to get regular template as fallback
                template_path, template_name = self.get_template_path()
            
            if not template_path:
                # Create basic document if no template found
                doc = Document()
            else:
                # Load existing template
                doc = Document(template_path)
            
            # Calculate summary data
            total_days = overtime_records.count()
            total_hours = sum(float(record.overtime_hours) for record in overtime_records)
            total_amount = sum(float(record.overtime_amount) for record in overtime_records)
            avg_per_day = total_hours / total_days if total_days > 0 else 0
            
            # Get approval information
            level1_approver = None
            final_approver = None
            
            # Find the most recent approved record for approval info
            approved_records = overtime_records.filter(status='approved')
            if approved_records.exists():
                latest_approved = approved_records.latest('final_approved_at')
                level1_approver = latest_approved.level1_approved_by
                final_approver = latest_approved.final_approved_by
            
            # Format dates
            current_date = dj_timezone.now()
            month_name = start_date.strftime('%B %Y')
            export_date = current_date.strftime('%d %B %Y')
            
            # Define replacement mappings
            replacements = {
                # Header info
                '{{PERIODE_EXPORT}}': month_name,
                '{{TANGGAL_EXPORT}}': export_date,
                '{{NOMOR_EXPORT}}': f"EXP-{month.replace('-', '')}/2025",
                
                # Summary data
                '{{TOTAL_HARI_LEMBUR}}': f"{total_days} hari",
                '{{TOTAL_JAM_LEMBUR}}': f"{total_hours:.1f} jam",
                '{{TOTAL_GAJI_LEMBUR}}': f"{total_amount:.2f} AED",
                '{{RATA_RATA_PER_HARI}}': f"{avg_per_day:.1f} jam",
                
                # Employee info
                '{{NAMA_PEGAWAI}}': employee.fullname or employee.user.get_full_name() or employee.user.username,
                '{{NIP_PEGAWAI}}': employee.nip or '-',
                '{{DIVISI_PEGAWAI}}': employee.division.name if employee.division else '-',
                '{{JABATAN_PEGAWAI}}': employee.position.name if employee.position else '-',
                
                # Approval info
                '{{LEVEL1_APPROVER}}': self.get_approver_name(level1_approver),
                '{{FINAL_APPROVER}}': self.get_approver_name(final_approver),
                '{{TANGGAL_APPROVAL}}': export_date,
                
                # Company info
                '{{NAMA_PERUSAHAAN}}': 'KJRI DUBAI',
                '{{ALAMAT_PERUSAHAAN}}': 'KONSULAT JENDERAL REPUBLIK INDONESIA DI DUBAI',
                '{{LOKASI}}': 'Dubai',
            }
            
            # Replace placeholders in document
            self.replace_placeholders_in_document(doc, replacements)
            
            # If no template exists, create basic structure
            if not template_path:
                self.create_basic_document_structure(doc, replacements)
            
            # Add overtime table to document
            self.add_overtime_table_to_document(doc, overtime_records, total_hours, total_amount)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Gagal generate DOCX: {str(e)}")

    def add_overtime_table_to_document(self, doc, overtime_records, total_hours, total_amount):
        """Add overtime table directly to document with professional styling"""
        from docx.shared import Inches
        
        # Add heading
        doc.add_heading('Detail Pengajuan Overtime', level=2)
        
        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(0.5)   # No
        table.columns[1].width = Inches(1.0)   # Tanggal
        table.columns[2].width = Inches(1.0)   # Jam Lembur
        table.columns[3].width = Inches(2.5)   # Deskripsi
        table.columns[4].width = Inches(1.2)   # Status
        table.columns[5].width = Inches(1.2)   # Gaji Lembur
        
        # Add header row
        header_row = table.rows[0]
        headers = ['No', 'Tanggal', 'Jam Lembur', 'Deskripsi Pekerjaan', 'Status', 'Gaji Lembur']
        
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Style header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for i, record in enumerate(overtime_records, 1):
            row = table.add_row()
            
            # Format data
            date_str = record.date_requested.strftime('%d/%m/%Y')
            hours_str = f"{float(record.overtime_hours):.1f} jam"
            
            # Truncate description
            description = record.work_description
            if len(description) > 50:
                description = description[:50] + '...'
            
            # Format status
            status_map = {
                'pending': 'Menunggu',
                'level1_approved': 'Level 1 Approved',
                'approved': 'Final Approved',
                'rejected': 'Ditolak'
            }
            status_display = status_map.get(record.status, record.status)
            
            # Format amount
            amount_str = f"{float(record.overtime_amount):.2f} AED"
            
            # Populate row
            row_data = [str(i), date_str, hours_str, description, status_display, amount_str]
            
            for j, data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = data
                
                # Style cells
                for paragraph in cell.paragraphs:
                    if j == 0:  # No column
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 2:  # Jam Lembur
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 4:  # Status
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 5:  # Gaji Lembur
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add summary row
        summary_row = table.add_row()
        summary_row.cells[0].text = 'TOTAL'
        summary_row.cells[1].text = ''
        summary_row.cells[2].text = f"{total_hours:.1f} jam"
        summary_row.cells[3].text = ''
        summary_row.cells[4].text = ''
        summary_row.cells[5].text = f"{total_amount:.2f} AED"
        
        # Style summary row
        for cell in summary_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add spacing
        doc.add_paragraph()

    def create_basic_document_structure(self, doc, replacements):
        """Create basic document structure if no template exists"""
        from docx.shared import Inches
        
        # Title
        title = doc.add_heading('LAPORAN OVERTIME BULANAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Header information
        doc.add_heading('INFORMASI DOKUMEN', level=1)
        
        # Create header table
        header_table = doc.add_table(rows=4, cols=2)
        header_table.style = 'Table Grid'
        
        # Header data
        header_data = [
            ['Periode', replacements['{{PERIODE_EXPORT}}']],
            ['Tanggal Export', replacements['{{TANGGAL_EXPORT}}']],
            ['Nomor Export', replacements['{{NOMOR_EXPORT}}']],
            ['Status', 'Draft']
        ]
        
        for i, (label, value) in enumerate(header_data):
            header_table.cell(i, 0).text = label
            header_table.cell(i, 1).text = value
        
        # Employee information
        doc.add_heading('INFORMASI PEGAWAI', level=1)
        
        emp_table = doc.add_table(rows=4, cols=2)
        emp_table.style = 'Table Grid'
        
        emp_data = [
            ['Nama', replacements['{{NAMA_PEGAWAI}}']],
            ['NIP', replacements['{{NIP_PEGAWAI}}']],
            ['Divisi', replacements['{{DIVISI_PEGAWAI}}']],
            ['Jabatan', replacements['{{JABATAN_PEGAWAI}}']]
        ]
        
        for i, (label, value) in enumerate(emp_data):
            emp_table.cell(i, 0).text = label
            emp_table.cell(i, 1).text = value
        
        # Summary
        doc.add_heading('RINGKASAN BULANAN', level=1)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ['Total Hari Lembur', replacements['{{TOTAL_HARI_LEMBUR}}']],
            ['Total Jam Lembur', replacements['{{TOTAL_JAM_LEMBUR}}']],
            ['Total Gaji Lembur', replacements['{{TOTAL_GAJI_LEMBUR}}']],
            ['Rata-rata per Hari', replacements['{{RATA_RATA_PER_HARI}}']]
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        # Approval section
        doc.add_heading('APPROVAL', level=1)
        
        approval_table = doc.add_table(rows=3, cols=2)
        approval_table.style = 'Table Grid'
        
        approval_data = [
            ['Level 1', replacements['{{LEVEL1_APPROVER}}']],
            ['Final', replacements['{{FINAL_APPROVER}}']],
            ['Tanggal Approval', replacements['{{TANGGAL_APPROVAL}}']]
        ]
        
        for i, (label, value) in enumerate(approval_data):
            approval_table.cell(i, 0).text = label
            approval_table.cell(i, 1).text = value
        
        # Company info
        doc.add_heading('PERUSAHAAN', level=1)
        doc.add_paragraph(replacements['{{NAMA_PERUSAHAAN}}'])
        doc.add_paragraph(replacements['{{ALAMAT_PERUSAHAAN}}'])
        doc.add_paragraph(replacements['{{LOKASI}}'])
        
        # Add spacing before overtime table
        doc.add_paragraph()

    def replace_placeholders_in_document(self, doc, replacements):
        """Replace placeholders in document paragraphs and tables"""
        # Function to replace text in paragraphs
        def replace_text_in_paragraphs(paragraphs, old_text, new_text):
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
        
        # Function to replace text in tables
        def replace_text_in_tables(tables, old_text, new_text):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
        
        # Replace all placeholders
        for placeholder, value in replacements.items():
            replace_text_in_paragraphs(doc.paragraphs, placeholder, value)
            replace_text_in_tables(doc.tables, placeholder, value)

    def get_approver_name(self, user):
        """Get approver name from employee.fullname with fallback to username"""
        if not user:
            return '-'
        # Try to get name from employee.fullname first
        if hasattr(user, 'employee') and user.employee and user.employee.fullname:
            return user.employee.fullname
        # Fallback to user.get_full_name()
        full_name = user.get_full_name()
        if full_name and full_name.strip():
            return full_name
        # Final fallback to username
        return user.username.title()  # Capitalize username as fallback

    def create_overtime_list_pdf(self, overtime_requests, month=None):
        """Generate simple PDF for list of overtime requests"""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from io import BytesIO
        from django.http import HttpResponse
        import datetime

        # Create buffer for PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

        # Get styles
        styles = getSampleStyleSheet()

        # Build content
        story = []

        # Title
        title = Paragraph("DAFTAR PENGAJUAN LEMBUR", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 20))

         # Build data table from overtime_requests
        table_data = [['No', 'Tanggal', 'Nama Pegawai', 'Jam Lembur', 'Status']]
        for idx, request in enumerate(overtime_requests, 1):
            employee_name = request.employee.fullname if request.employee and request.employee.fullname else "N/A"
            date_str = request.date_requested.strftime('%d/%m/%Y') if request.date_requested else "N/A"
            overtime_hours = f"{request.overtime_hours:.2f}" if request.overtime_hours else "0.00"
            status_display = request.get_status_display() if hasattr(request, 'get_status_display') else request.status

            table_data.append([
                str(idx),
                date_str,
                employee_name,
                overtime_hours,
                status_display
            ])

        # Create table
        main_table = Table(table_data, colWidths=[0.5*inch, 1*inch, 2*inch, 1*inch, 1*inch])
        main_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
            ('TOPPADDING', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(main_table)
        story.append(Spacer(1, 20))

        # Footer
        footer_text = f"Dicetak pada: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"
        footer = Paragraph(footer_text, styles['Normal'])
        story.append(footer)

        # Build PDF
        doc.build(story)

        # Get PDF content
        buffer.seek(0)
        pdf_content = buffer.getvalue()
        buffer.close()

        # Create response
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="daftar-pengajuan-lembur-{month or "all"}.pdf"'

        return response

    def convert_docx_to_pdf(self, docx_file_path):
        """Convert DOCX file to PDF using reportlab (primary) or pypandoc (fallback)"""
        import tempfile
        import os

        # Create temporary PDF file
        pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        pdf_temp.close()

        try:
            # Primary method: Use reportlab (fast and reliable)
            self._convert_docx_to_pdf_reportlab(docx_file_path, pdf_temp.name)

            if os.path.exists(pdf_temp.name) and os.path.getsize(pdf_temp.name) > 0:
                return pdf_temp.name
            else:
                raise Exception("PDF file tidak berhasil dibuat dengan reportlab")

        except Exception as reportlab_error:
            # Fallback: Use pypandoc if reportlab fails
            try:
                import pypandoc
                pypandoc.convert_file(
                    docx_file_path,
                    'pdf',
                    outputfile=pdf_temp.name,
                    extra_args=[
                        '--pdf-engine=pdflatex',
                        '--variable', 'geometry:margin=2cm',
                        '--variable', 'fontsize=10pt'
                    ]
                )
                if os.path.exists(pdf_temp.name) and os.path.getsize(pdf_temp.name) > 0:
                    return pdf_temp.name
                else:
                    raise Exception("PDF file tidak berhasil dibuat dengan pypandoc")
            except Exception as pypandoc_error:
                # Clean up temp file on error
                if os.path.exists(pdf_temp.name):
                    os.unlink(pdf_temp.name)
                raise Exception(f"Error saat konversi DOCX ke PDF - Reportlab: {str(reportlab_error)}, Pypandoc: {str(pypandoc_error)}")


    def _convert_docx_to_pdf_reportlab(self, docx_file_path, output_pdf_path):
        """Convert DOCX to PDF using reportlab with advanced template matching"""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

            # Load the DOCX document
            doc = Document(docx_file_path)

            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_pdf_path,
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )

            # Get styles
            styles = getSampleStyleSheet()

            # Create custom styles
            custom_styles = {
                'Normal': styles['Normal'],
                'Title': ParagraphStyle(
                    'Title',
                    parent=styles['Normal'],
                    fontSize=16,
                    spaceAfter=30,
                    alignment=TA_CENTER,
                    fontName='Helvetica-Bold'
                ),
                'Heading1': ParagraphStyle(
                    'Heading1',
                    parent=styles['Normal'],
                    fontSize=14,
                    spaceAfter=20,
                    fontName='Helvetica-Bold'
                ),
                'Heading2': ParagraphStyle(
                    'Heading2',
                    parent=styles['Normal'],
                    fontSize=12,
                    spaceAfter=15,
                    fontName='Helvetica-Bold'
                ),
                'Signature': ParagraphStyle(
                    'Signature',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER,
                    spaceAfter=50
                )
            }

            # Create story (content)
            story = []

            # Process each paragraph in the DOCX
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    # Create paragraph style based on DOCX formatting
                    style = self._determine_paragraph_style(paragraph, custom_styles)

                    # Format text with runs
                    formatted_text = self._format_paragraph_text(paragraph)

                    if formatted_text:
                        p = Paragraph(formatted_text, style)
                        story.append(p)
                        story.append(Spacer(1, 0.2*cm))

            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = self._clean_cell_text(cell)
                        row_data.append(cell_text)
                    table_data.append(row_data)

                if table_data:
                    # Create table
                    pdf_table = Table(table_data)
                    table_style = self._create_table_style(table_data, len(table_data), len(table_data[0]) if table_data else 0)
                    pdf_table.setStyle(table_style)

                    story.append(pdf_table)
                    story.append(Spacer(1, 0.5*cm))

            # Build PDF
            pdf_doc.build(story)

        except Exception as e:
            raise Exception(f"Error converting DOCX to PDF with reportlab: {str(e)}")

    def _determine_paragraph_style(self, paragraph, custom_styles):
        """Determine the appropriate style for a paragraph based on DOCX properties"""
        try:
            # Check for title-like paragraphs
            if paragraph.style.name.startswith('Title') or 'title' in paragraph.text.lower():
                return custom_styles['Title']

            # Check for heading-like paragraphs
            if paragraph.style.name.startswith('Heading'):
                if '1' in paragraph.style.name:
                    return custom_styles['Heading1']
                else:
                    return custom_styles['Heading2']

            # Check for signature-like paragraphs
            if any(keyword in paragraph.text.lower() for keyword in ['ttd', 'signature', 'dibuat', 'disetujui']):
                return custom_styles['Signature']

            # Default to normal style
            return custom_styles['Normal']

        except:
            return custom_styles['Normal']

    def _format_paragraph_text(self, paragraph):
        """Format paragraph text with special characters and formatting"""
        try:
            text = paragraph.text

            # Handle common Indonesian document formatting
            text = text.replace(':', ' : ')
            text = text.replace(';', ' ; ')

            # Handle bullet points and numbering
            if text.strip().startswith(('•', '-', '*', '1.', '2.', '3.')):
                text = '&nbsp;&nbsp;&nbsp;&nbsp;' + text

            return text

        except:
            return paragraph.text

    def _clean_cell_text(self, cell_text):
        """Clean and format cell text for table display"""
        try:
            # Remove extra whitespace
            cleaned = ' '.join(cell_text.split())

            # Handle empty cells
            if not cleaned:
                return '-'

            # Truncate very long text
            if len(cleaned) > 50:
                cleaned = cleaned[:47] + '...'

            return cleaned

        except:
            return cell_text if cell_text else '-'

    def _create_table_style(self, table_data, rows, cols):
        """Create enhanced table styling that matches DOCX template"""
        try:
            from reportlab.lib import colors

            # Base table style
            style = [
                # Header row styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),

                # Data rows styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 6),

                # Grid lines
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.black),
            ]

            return TableStyle(style)

        except:
            # Return basic table style if custom styling fails
            from reportlab.platypus import TableStyle
            return TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ])


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def supervisor_approvals_summary(request):
    """Get summary of pending approvals for supervisor dashboard card"""
    user = request.user
    
    # Check if user has supervisor capabilities (position approval level >= 1)
    from .utils import ApprovalChecker
    approval_level = ApprovalChecker.get_user_approval_level(user)
    
    if approval_level < 1:
        return Response(
            {"detail": "Hanya supervisor yang dapat mengakses endpoint ini"}, 
            status=status.HTTP_403_FORBIDDEN
        )
    
    try:
        supervisor_employee = user.employee
    except:
        return Response(
            {"detail": "User tidak memiliki profil employee"}, 
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Check if supervisor has org-wide approval permission
    is_org_wide_supervisor = (supervisor_employee.position and 
                            supervisor_employee.position.can_approve_overtime_org_wide)
    
    # Build overtime requests queryset based on supervisor scope
    if is_org_wide_supervisor:
        # Org-wide supervisors can see all requests
        overtime_qs = OvertimeRequest.objects.all()
        # For org-wide supervisors, show requests that need final approval
        pending_overtime_count = overtime_qs.filter(status__in=['pending', 'level1_approved']).count()
    else:
        # Division supervisors can only see requests from their division
        overtime_qs = OvertimeRequest.objects.filter(
            employee__division=supervisor_employee.division
        )
        # For division supervisors, show only pending requests (need level 1 approval)
        pending_overtime_count = overtime_qs.filter(status='pending').count()
    
    # Build attendance corrections queryset based on supervisor scope
    if is_org_wide_supervisor:
        # Org-wide supervisors can see all corrections
        corrections_qs = AttendanceCorrection.objects.all()
    else:
        # Division supervisors can only see corrections from their division
        corrections_qs = AttendanceCorrection.objects.filter(
            user__employee__division=supervisor_employee.division
        )
    
    pending_corrections_count = corrections_qs.filter(
        status=AttendanceCorrection.CorrectionStatus.PENDING
    ).count()
    
    return Response({
        'pending_overtime_count': pending_overtime_count,
        'pending_corrections_count': pending_corrections_count,
        'is_org_wide_supervisor': is_org_wide_supervisor,
        'supervisor_division': supervisor_employee.division.name if supervisor_employee.division else None,
        'can_approve_overtime_org_wide': is_org_wide_supervisor,
        'is_admin': user.is_superuser or user.groups.filter(name='admin').exists(),
    })

# ============================================================================
# MONTHLY SUMMARY REQUEST VIEWSET
# ============================================================================

class OvertimeSummaryRequestViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing monthly summary requests with role-based access control
    """
    permission_classes = [IsOvertimeRequestOwnerOrSupervisor]
    pagination_class = DefaultPagination
    
    def get_queryset(self):
        """Filter queryset based on user role"""
        user = self.request.user
        
        # Admin can see all requests
        if user.is_superuser or user.groups.filter(name='admin').exists():
            return OvertimeSummaryRequest.objects.all().select_related(
                'user', 'employee', 'level1_approved_by', 'final_approved_by'
            )
        
        # Use position-based approval checking for supervisor capabilities
        else:
            from .utils import ApprovalChecker
            approval_level = ApprovalChecker.get_user_approval_level(user)
            
            if approval_level >= 1:
                try:
                    supervisor_employee = user.employee
                    
                    # Check if supervisor has org-wide approval permission
                    if ApprovalChecker.can_approve_overtime_org_wide(user):
                        # Org-wide supervisors can see all requests
                        return OvertimeSummaryRequest.objects.all().select_related(
                            'user', 'employee', 'level1_approved_by', 'final_approved_by'
                        )
                    else:
                        # Division supervisors can only see requests from their division
                        return OvertimeSummaryRequest.objects.filter(
                            employee__division=supervisor_employee.division
                        ).select_related('user', 'employee', 'level1_approved_by', 'final_approved_by')
                except:
                    return OvertimeSummaryRequest.objects.none()
            else:
                # No supervisor capabilities - can only see own requests
                return OvertimeSummaryRequest.objects.filter(user=user).select_related(
                    'user', 'employee', 'level1_approved_by', 'final_approved_by'
                )
        
        return OvertimeSummaryRequest.objects.none()
    
    def get_serializer_class(self):
        """Return appropriate serializer based on user role and action"""
        user = self.request.user
        
        # For create action, use create serializer
        if self.action == 'create':
            return OvertimeSummaryRequestCreateSerializer
        
        # Role-based serializers for other actions
        if user.is_superuser or user.groups.filter(name='admin').exists():
            return OvertimeSummaryRequestAdminSerializer
        
        # Use position-based approval checking for supervisor capabilities
        from .utils import ApprovalChecker
        approval_level = ApprovalChecker.get_user_approval_level(user)
        
        if approval_level >= 1:
            return OvertimeSummaryRequestSupervisorSerializer
        else:
            return OvertimeSummaryRequestEmployeeSerializer
    
    def perform_create(self, serializer):
        """Auto-set user and employee when creating request"""
        user = self.request.user
        try:
            employee = user.employee
            serializer.save(user=user, employee=employee)
        except AttributeError:
            raise serializers.ValidationError("User tidak memiliki profil employee")
        except IntegrityError as e:
            if 'unique constraint' in str(e).lower() or 'duplicate' in str(e).lower():
                raise serializers.ValidationError(
                    "Anda sudah memiliki pengajuan rekap untuk periode dan jenis laporan yang sama. "
                    "Silakan gunakan pengajuan yang sudah ada atau tunggu hingga selesai."
                )
            raise serializers.ValidationError(f"Terjadi kesalahan database: {str(e)}")
        except Exception as e:
            raise serializers.ValidationError(f"Terjadi kesalahan: {str(e)}")
    
    def create(self, request, *args, **kwargs):
        """Override create to ensure only employees can create requests"""
        if not request.user.groups.filter(name='pegawai').exists():
            return Response(
                {"detail": "Hanya pegawai yang dapat mengajukan rekap bulanan"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        return super().create(request, *args, **kwargs)
    
    @action(detail=True, methods=['post'], permission_classes=[IsOvertimeRequestApprover])
    def approve(self, request, pk=None):
        """Approve monthly summary request with 2-level approval system"""
        summary_request = self.get_object()
        user = request.user
        
        # Check if user is admin (can bypass 2-level approval)
        is_admin = user.is_superuser or user.groups.filter(name='admin').exists()
        
        # Use position-based approval checking
        from .utils import ApprovalChecker
        
        # Check if user is org-wide supervisor
        is_org_wide_supervisor = ApprovalChecker.can_approve_organization_overtime(user, summary_request.employee)
        # Check if user has no approval permission (level 0)
        has_no_approval = ApprovalChecker.get_user_approval_level(user) == 0
        
        # Reject if user has no approval permission
        if has_no_approval:
            return Response(
                {"detail": "Anda tidak memiliki permission untuk melakukan approval"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Admin can approve any status
        if is_admin:
            if summary_request.status in ['rejected', 'cancelled']:
                return Response(
                    {"detail": "Pengajuan yang sudah ditolak/dibatalkan tidak dapat disetujui"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Admin approval goes straight to final approval
            summary_request.status = 'approved'
            summary_request.final_approved_by = user
            summary_request.final_approved_at = dj_timezone.now()
            summary_request.rejection_reason = None
            summary_request.save()
            
            serializer = self.get_serializer(summary_request)
            return Response({
                "detail": "Pengajuan rekap bulanan berhasil disetujui (Final Approval)",
                "data": serializer.data
            })
        
        # Org-wide supervisor (final approval)
        elif is_org_wide_supervisor:
            if summary_request.status != 'level1_approved':
                if summary_request.status == 'pending':
                    return Response(
                        {"detail": "Pengajuan harus disetujui oleh supervisor divisi terlebih dahulu (Level 1 Approval)"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
                else:
                    return Response(
                        {"detail": "Hanya pengajuan dengan status Level 1 Approved yang dapat disetujui final"}, 
                        status=status.HTTP_400_BAD_REQUEST
                    )
            
            summary_request.status = 'approved'
            summary_request.final_approved_by = user
            summary_request.final_approved_at = dj_timezone.now()
            summary_request.rejection_reason = None
            summary_request.save()
            
            serializer = self.get_serializer(summary_request)
            return Response({
                "detail": "Pengajuan rekap bulanan berhasil disetujui (Final Approval)",
                "data": serializer.data
            })
        
        # Division supervisor (level 1 approval)
        else:
            if summary_request.status != 'pending':
                return Response(
                    {"detail": "Hanya pengajuan dengan status pending yang dapat disetujui level 1"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            summary_request.status = 'level1_approved'
            summary_request.level1_approved_by = user
            summary_request.level1_approved_at = dj_timezone.now()
            summary_request.rejection_reason = None
            summary_request.save()
            
            serializer = self.get_serializer(summary_request)
            return Response({
                "detail": "Pengajuan rekap bulanan berhasil disetujui (Level 1 Approval). Menunggu approval final.",
                "data": serializer.data
            })
    
    @action(detail=True, methods=['post'], permission_classes=[IsOvertimeRequestApprover])
    def reject(self, request, pk=None):
        """Reject monthly summary request"""
        summary_request = self.get_object()
        
        # Can reject pending or level1_approved requests
        if summary_request.status not in ['pending', 'level1_approved']:
            return Response(
                {"detail": "Hanya pengajuan dengan status pending atau level1_approved yang dapat ditolak"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        rejection_reason = request.data.get('rejection_reason', '')
        if not rejection_reason.strip():
            return Response(
                {"detail": "Alasan penolakan harus diisi"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        summary_request.status = 'rejected'
        summary_request.rejection_reason = rejection_reason
        # Clear all approval fields
        summary_request.level1_approved_by = None
        summary_request.level1_approved_at = None
        summary_request.final_approved_by = None
        summary_request.final_approved_at = None
        summary_request.save()
        
        serializer = self.get_serializer(summary_request)
        return Response({
            "detail": "Pengajuan rekap bulanan berhasil ditolak",
            "data": serializer.data
        })
    
    @action(detail=True, methods=['post'], permission_classes=[IsOvertimeRequestApprover])
    def complete(self, request, pk=None):
        """Mark request as completed"""
        summary_request = self.get_object()
        
        # Only approved requests can be completed
        if summary_request.status != 'approved':
            return Response(
                {"detail": "Hanya pengajuan yang sudah disetujui yang dapat ditandai selesai"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        completion_notes = request.data.get('completion_notes', '')
        summary_request.status = 'completed'
        summary_request.completed_at = dj_timezone.now()
        summary_request.completion_notes = completion_notes
        summary_request.save()
        
        serializer = self.get_serializer(summary_request)
        return Response({
            "detail": "Pengajuan rekap bulanan berhasil ditandai selesai",
            "data": serializer.data
        })
    
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get monthly summary request summary statistics"""
        queryset = self.get_queryset()
        
        # Filter by date range if provided
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        
        if start_date:
            queryset = queryset.filter(request_period__gte=start_date)
        if end_date:
            queryset = queryset.filter(request_period__lte=end_date)
        
        # Calculate statistics
        total_requests = queryset.count()
        pending_requests = queryset.filter(status='pending').count()
        level1_approved_requests = queryset.filter(status='level1_approved').count()
        approved_requests = queryset.filter(status='approved').count()
        completed_requests = queryset.filter(status='completed').count()
        rejected_requests = queryset.filter(status='rejected').count()
        
        # Priority breakdown
        low_priority = queryset.filter(priority='low').count()
        medium_priority = queryset.filter(priority='medium').count()
        high_priority = queryset.filter(priority='high').count()
        urgent_priority = queryset.filter(priority='urgent').count()
        
        return Response({
            'total_requests': total_requests,
            'pending_requests': pending_requests,
            'level1_approved_requests': level1_approved_requests,
            'approved_requests': approved_requests,
            'completed_requests': completed_requests,
            'rejected_requests': rejected_requests,
            'priority_breakdown': {
                'low': low_priority,
                'medium': medium_priority,
                'high': high_priority,
                'urgent': urgent_priority,
            }
        })

    @action(detail=True, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def generate_report(self, request, pk=None):
        """Generate overtime report data for approved monthly summary request"""
        summary_request = self.get_object()
        user = request.user
        
        # Only generate report for approved requests
        if summary_request.status != 'approved':
            return Response(
                {"detail": "Hanya pengajuan yang sudah disetujui yang dapat di-generate laporannya"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Parse period
            year, month = summary_request.request_period.split('-')
            year = int(year)
            month = int(month)
            
            # Get overtime data for the period
            overtime_data = self.get_overtime_data_for_period(
                summary_request.employee, 
                summary_request.request_period,
                include_approved_only=True
            )
            
            # Generate comprehensive summary based on request scope
            summary_stats = self.generate_comprehensive_summary(
                summary_request.employee, 
                summary_request.request_period,
                summary_request
            )
            
            # Prepare response data
            report_data = {
                'request_info': {
                    'id': summary_request.id,
                    'period': summary_request.request_period,
                    'report_type': summary_request.report_type,
                    'request_title': summary_request.request_title,
                    'request_description': summary_request.request_description,
                    'data_scope': {
                        'include_attendance': summary_request.include_attendance,
                        'include_overtime': summary_request.include_overtime,
                        'include_corrections': summary_request.include_corrections,
                        'include_summary_stats': summary_request.include_summary_stats,
                    }
                },
                'employee_info': {
                    'nip': summary_request.employee.nip,
                    'fullname': summary_request.employee.fullname or f"{summary_request.employee.user.first_name} {summary_request.employee.user.last_name}".strip(),
                    'division': summary_request.employee.division.name if summary_request.employee.division else None,
                    'position': summary_request.employee.position.name if summary_request.employee.position else None,
                },
                'approval_info': {
                    'level1_approved_by': summary_request.level1_approved_by.username if summary_request.level1_approved_by else None,
                    'level1_approved_at': summary_request.level1_approved_at.isoformat() if summary_request.level1_approved_at else None,
                    'final_approved_by': summary_request.final_approved_by.username if summary_request.final_approved_by else None,
                    'final_approved_at': summary_request.final_approved_at.isoformat() if summary_request.final_approved_at else None,
                },
                'overtime_summary': summary_stats,
                'generated_at': dj_timezone.now().isoformat(),
            }
            
            return Response(report_data)
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate laporan: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def export_docx(self, request, pk=None):
        """Export monthly summary to DOCX using existing template system"""
        summary_request = self.get_object()
        user = request.user
        
        # Only export for approved requests
        if summary_request.status != 'approved':
            return Response(
                {"detail": "Hanya pengajuan yang sudah disetujui yang dapat di-export"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Parse period
            year, month = summary_request.request_period.split('-')
            year = int(year)
            month = int(month)
            
            # Get overtime data for the period
            overtime_data = self.get_overtime_data_for_period(
                summary_request.employee, 
                summary_request.request_period,
                include_approved_only=True
            )
            
            if not overtime_data.exists():
                return Response(
                    {"detail": f"Tidak ada data overtime untuk periode {summary_request.request_period}"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Generate DOCX document using existing template system
            docx_file = self.generate_monthly_summary_docx(
                summary_request,
                overtime_data,
                summary_request.request_period
            )
            
            # Return file response
            from django.http import FileResponse
            import os
            
            filename = f"Rekap_Lembur_Bulanan_{summary_request.employee.nip}_{summary_request.request_period}.docx"
            
            response = FileResponse(
                open(docx_file, 'rb'),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            # Clean up temporary file
            os.unlink(docx_file)
            
            return response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal export DOCX: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'], permission_classes=[IsOvertimeRequestOwnerOrSupervisor])
    def export_pdf(self, request, pk=None):
        """Export monthly summary to PDF by converting DOCX"""
        summary_request = self.get_object()
        user = request.user
        
        # Only export for approved requests
        if summary_request.status != 'approved':
            return Response(
                {"detail": "Hanya pengajuan yang sudah disetujui yang dapat di-export"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Parse period
            year, month = summary_request.request_period.split('-')
            year = int(year)
            month = int(month)
            
            # Get overtime data for the period
            overtime_data = self.get_overtime_data_for_period(
                summary_request.employee, 
                summary_request.request_period,
                include_approved_only=True
            )
            
            if not overtime_data.exists():
                return Response(
                    {"detail": f"Tidak ada data overtime untuk periode {summary_request.request_period}"}, 
                    status=status.HTTP_404_NOT_FOUND
                )
            
            # Generate DOCX document first using existing template system
            docx_file = self.generate_monthly_summary_docx(
                summary_request,
                overtime_data,
                summary_request.request_period
            )
            
            # Convert DOCX to PDF
            pdf_file = self.convert_docx_to_pdf(docx_file)
            
            # Return PDF file response
            from django.http import FileResponse
            import os
            
            filename = f"Rekap_Lembur_Bulanan_{summary_request.employee.nip}_{summary_request.request_period}.pdf"
            
            response = FileResponse(
                open(pdf_file, 'rb'),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            # Clean up temporary files
            os.unlink(docx_file)
            os.unlink(pdf_file)
            
            return response
            
        except Exception as e:
            return Response(
                {"detail": f"Gagal export PDF: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def get_overtime_data_for_period(self, employee, period, include_approved_only=True):
        """
        Get overtime data for specific period and employee
        """
        try:
            year, month = period.split('-')
            year = int(year)
            month = int(month)
            
            # Base query for overtime requests
            queryset = OvertimeRequest.objects.filter(
                employee=employee,
                date_requested__year=year,
                date_requested__month=month
            ).select_related(
                'level1_approved_by', 
                'final_approved_by'
            ).order_by('date_requested')
            
            # Filter by approval status if needed
            if include_approved_only:
                queryset = queryset.filter(status='approved')
            
            return queryset
            
        except (ValueError, AttributeError) as e:
            # Return empty queryset if period format is invalid
            return OvertimeRequest.objects.none()

    def generate_overtime_summary(self, employee, period, overtime_data):
        """
        Generate overtime summary report with detailed statistics
        """
        try:
            # Basic statistics
            total_overtime_hours = overtime_data.aggregate(
                total=models.Sum('overtime_hours')
            )['total'] or 0
            
            total_overtime_amount = overtime_data.aggregate(
                total=models.Sum('overtime_amount')
            )['total'] or 0
            
            total_requests = overtime_data.count()
            
            # Get work settings for calculation
            work_settings = WorkSettings.objects.first()
            
            # Calculate average hourly rate
            avg_hourly_rate = 0
            if total_overtime_hours > 0:
                avg_hourly_rate = float(total_overtime_amount) / float(total_overtime_hours)
            
            # Detailed breakdown by date
            overtime_by_date = {}
            for overtime in overtime_data:
                date_str = overtime.date_requested.strftime('%Y-%m-%d')
                if date_str not in overtime_by_date:
                    overtime_by_date[date_str] = {
                        'date': overtime.date_requested.strftime('%d %B %Y'),
                        'total_hours': 0,
                        'total_amount': 0,
                        'requests': []
                    }
                
                overtime_by_date[date_str]['total_hours'] += float(overtime.overtime_hours)
                overtime_by_date[date_str]['total_amount'] += float(overtime.overtime_amount)
                overtime_by_date[date_str]['requests'].append({
                    'id': overtime.id,
                    'hours': float(overtime.overtime_hours),
                    'amount': float(overtime.overtime_amount),
                    'work_description': overtime.work_description,
                    'level1_approver': overtime.level1_approved_by.username if overtime.level1_approved_by else None,
                    'final_approver': overtime.final_approved_by.username if overtime.final_approved_by else None,
                    'final_approved_at': overtime.final_approved_at.isoformat() if overtime.final_approved_at else None,
                })
            
            # Convert to list and sort by date
            overtime_by_date_list = list(overtime_by_date.values())
            overtime_by_date_list.sort(key=lambda x: x['date'])
            
            # Summary statistics
            summary_stats = {
                'period': period,
                'total_requests': total_requests,
                'total_overtime_hours': float(total_overtime_hours),
                'total_overtime_amount': float(total_overtime_amount),
                'average_hourly_rate': round(avg_hourly_rate, 2),
                'average_daily_hours': round(float(total_overtime_hours) / len(overtime_by_date) if overtime_by_date else 0, 2),
                'average_daily_amount': round(float(total_overtime_amount) / len(overtime_by_date) if overtime_by_date else 0, 2),
                'workdays_with_overtime': len(overtime_by_date),
                'overtime_by_date': overtime_by_date_list,
            }
            
            return summary_stats
            
        except Exception as e:
            # Return basic stats if detailed calculation fails
            return {
                'period': period,
                'total_requests': overtime_data.count(),
                'total_overtime_hours': 0,
                'total_overtime_amount': 0,
                'error': f"Gagal generate detail statistik: {str(e)}"
            }

    def get_attendance_data_for_period(self, employee, period):
        """
        Get attendance data for specific period and employee
        """
        try:
            year, month = period.split('-')
            year = int(year)
            month = int(month)
            
            # Get attendance records for the period
            from datetime import date
            start_date = date(year, month, 1)
            if month == 12:
                end_date = date(year + 1, 1, 1) - timedelta(days=1)
            else:
                end_date = date(year, month + 1, 1) - timedelta(days=1)
            
            queryset = Attendance.objects.filter(
                user=employee.user,
                date_local__gte=start_date,
                date_local__lte=end_date
            ).select_related('user').order_by('date_local')
            
            return queryset
            
        except (ValueError, AttributeError) as e:
            return Attendance.objects.none()

    def get_corrections_data_for_period(self, employee, period):
        """
        Get attendance corrections data for specific period and employee
        """
        try:
            year, month = period.split('-')
            year = int(year)
            month = int(month)
            
            # Get corrections for the period
            from datetime import date
            start_date = date(year, month, 1)
            if month == 12:
                end_date = date(year + 1, 1, 1) - timedelta(days=1)
            else:
                end_date = date(year, month + 1, 1) - timedelta(days=1)
            
            queryset = AttendanceCorrection.objects.filter(
                employee=employee,
                correction_date__gte=start_date,
                correction_date__lte=end_date,
                status='approved'  # Only approved corrections
            ).select_related('employee', 'approved_by').order_by('correction_date')
            
            return queryset
            
        except (ValueError, AttributeError) as e:
            return AttendanceCorrection.objects.none()

    def generate_comprehensive_summary(self, employee, period, summary_request):
        """
        Generate comprehensive monthly summary including overtime, attendance, and corrections
        """
        try:
            # Get all data types based on request scope
            overtime_data = None
            attendance_data = None
            corrections_data = None
            
            if summary_request.include_overtime:
                overtime_data = self.get_overtime_data_for_period(employee, period, include_approved_only=True)
            
            if summary_request.include_attendance:
                attendance_data = self.get_attendance_data_for_period(employee, period)
            
            if summary_request.include_corrections:
                corrections_data = self.get_corrections_data_for_period(employee, period)
            
            # Generate overtime summary if requested
            overtime_summary = None
            if overtime_data is not None:
                overtime_summary = self.generate_overtime_summary(employee, period, overtime_data)
            
            # Generate attendance summary if requested
            attendance_summary = None
            if attendance_data is not None:
                attendance_summary = self.generate_attendance_summary(employee, period, attendance_data)
            
            # Generate corrections summary if requested
            corrections_summary = None
            if corrections_data is not None:
                corrections_summary = self.generate_corrections_summary(employee, period, corrections_data)
            
            # Combine all summaries
            comprehensive_summary = {
                'period': period,
                'data_scope': {
                    'overtime_included': summary_request.include_overtime,
                    'attendance_included': summary_request.include_attendance,
                    'corrections_included': summary_request.include_corrections,
                    'summary_stats_included': summary_request.include_summary_stats,
                },
                'overtime_summary': overtime_summary,
                'attendance_summary': attendance_summary,
                'corrections_summary': corrections_summary,
                'generated_at': dj_timezone.now().isoformat(),
            }
            
            return comprehensive_summary
            
        except Exception as e:
            return {
                'period': period,
                'error': f"Gagal generate comprehensive summary: {str(e)}"
            }

    def generate_monthly_summary_docx(self, summary_request, overtime_data, period):
        """Generate monthly summary DOCX using existing template system"""
        try:
            # Import required libraries
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            from io import BytesIO
            from django.utils import timezone
            import os
            import locale
            import tempfile
            from datetime import datetime
            
            # Set locale for Indonesian formatting
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_ALL, 'id_ID')
                except:
                    pass
            
            # Get monthly export template path (priority) or fallback to regular template
            template_path, template_name = self.get_monthly_export_template_path()
            
            if not template_path:
                # Try to get regular template as fallback
                template_path, template_name = self.get_monthly_export_template_path()
            
            if not template_path:
                # Create basic document if no template found
                doc = Document()
            else:
                # Load existing template
                doc = Document(template_path)
            
            # Calculate summary data
            total_days = overtime_data.count()
            total_hours = sum(float(record.overtime_hours) for record in overtime_data)
            total_amount = sum(float(record.overtime_amount) for record in overtime_data)
            avg_per_day = total_hours / total_days if total_days > 0 else 0
            
            # Get approval information
            level1_approver = summary_request.level1_approved_by
            final_approver = summary_request.final_approved_by
            
            # Format dates
            current_date = dj_timezone.now()
            month_name = datetime.strptime(period, '%Y-%m').strftime('%B %Y')
            export_date = current_date.strftime('%d %B %Y')
            
            # Format approval dates
            level1_approval_date = summary_request.level1_approved_at.strftime('%d %B %Y') if summary_request.level1_approved_at else '-'
            final_approval_date = summary_request.final_approved_at.strftime('%d %B %Y') if summary_request.final_approved_at else '-'
            
            # Define replacement mappings
            replacements = {
                # Header info
                '{{PERIODE_EXPORT}}': month_name,
                '{{TANGGAL_EXPORT}}': export_date,
                '{{NOMOR_EXPORT}}': f"REKAP-{period.replace('-', '')}/2025",
                
                # Summary data
                '{{TOTAL_HARI_LEMBUR}}': f"{total_days} hari",
                '{{TOTAL_JAM_LEMBUR}}': f"{total_hours:.1f} jam",
                '{{TOTAL_GAJI_LEMBUR}}': f"{total_amount:.2f} AED",
                '{{RATA_RATA_PER_HARI}}': f"{avg_per_day:.1f} jam",
                
                # Employee info
                '{{NAMA_PEGAWAI}}': summary_request.employee.fullname or summary_request.employee.user.get_full_name() or summary_request.employee.user.username,
                '{{NIP_PEGAWAI}}': summary_request.employee.nip or '-',
                '{{DIVISI_PEGAWAI}}': summary_request.employee.division.name if summary_request.employee.division else '-',
                '{{JABATAN_PEGAWAI}}': summary_request.employee.position.name if summary_request.employee.position else '-',
                
                # Approval info
                '{{LEVEL1_APPROVER}}': level1_approver.username if level1_approver else '-',
                '{{LEVEL1_APPROVER_NIP}}': level1_approver.employee.nip if level1_approver and hasattr(level1_approver, 'employee') and level1_approver.employee else '-',
                '{{LEVEL1_APPROVAL_DATE}}': level1_approval_date,
                '{{FINAL_APPROVER}}': final_approver.username if final_approver else '-',
                '{{FINAL_APPROVER_NIP}}': final_approver.employee.nip if final_approver and hasattr(final_approver, 'employee') and final_approver.employee else '-',
                '{{FINAL_APPROVAL_DATE}}': final_approval_date,
                '{{TANGGAL_APPROVAL}}': export_date,
                
                # Company info
                '{{NAMA_PERUSAHAAN}}': 'KJRI DUBAI',
                '{{ALAMAT_PERUSAHAAN}}': 'KONSULAT JENDERAL REPUBLIK INDONESIA DI DUBAI',
                '{{LOKASI}}': 'Dubai',
                
                # Special placeholders (will be replaced with actual content)
                '{{TABEL_OVERTIME}}': '[TABEL_OVERTIME_WILL_BE_REPLACED]',  # Placeholder for overtime table
            }
            
            # Replace placeholders in document
            self.replace_placeholders_in_document(doc, replacements)
            
            # If no template exists, create basic structure
            if not template_path:
                self.create_basic_monthly_summary_structure(doc, replacements)
            
            # Add overtime table to document
            self.add_overtime_table_to_document(doc, overtime_data, total_hours, total_amount)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Gagal generate DOCX monthly summary: {str(e)}")

    def get_monthly_export_template_path(self):
        """Get the path to the active monthly export template"""
        import os
        from django.conf import settings
        
        template_dir = os.path.join(settings.BASE_DIR, 'template')
        
        # Priority order for monthly export template files
        priority_names = [
            'template_rekap_lembur.docx',                   # Primary monthly export template (valid)
            'template_SURAT_PERINTAH_KERJA_LEMBUR.docx',    # Fallback 1 (valid)
            'template_monthly_overtime_export.docx',         # Fallback 2 (small file, may be corrupt)
        ]
        
        # Check each priority template
        for template_name in priority_names:
            template_path = os.path.join(template_dir, template_name)
            if os.path.exists(template_path):
                return template_path, template_name
        
        # If no priority template found, return None
        return None, None

    def create_basic_monthly_summary_structure(self, doc, replacements):
        """Create basic document structure for monthly summary if no template exists"""
        from docx.shared import Inches
        
        # Title
        title = doc.add_heading('REKAP LEMBUR BULANAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Header information
        doc.add_heading('INFORMASI REKAP', level=1)
        
        # Create header table
        header_table = doc.add_table(rows=4, cols=2)
        header_table.style = 'Table Grid'
        
        # Header data
        header_data = [
            ['Periode', replacements['{{PERIODE_EXPORT}}']],
            ['Tanggal Export', replacements['{{TANGGAL_EXPORT}}']],
            ['Nomor Export', replacements['{{NOMOR_EXPORT}}']],
            ['Status', 'Approved']
        ]
        
        for i, (label, value) in enumerate(header_data):
            header_table.cell(i, 0).text = label
            header_table.cell(i, 1).text = value

    def generate_attendance_summary(self, employee, period, attendance_data):
        """
        Generate attendance summary for the period
        """
        try:
            total_days = attendance_data.count()
            present_days = attendance_data.filter(check_in_at_utc__isnull=False, check_out_at_utc__isnull=False).count()
            late_days = attendance_data.filter(lateness_minutes__gt=0).count()
            absent_days = attendance_data.filter(check_in_at_utc__isnull=True).count()
            
            # Calculate total work hours
            total_work_minutes = attendance_data.aggregate(
                total=models.Sum('total_work_minutes')
            )['total'] or 0
            
            total_work_hours = total_work_minutes / 60 if total_work_minutes > 0 else 0
            
            # Calculate average daily work hours
            avg_daily_hours = total_work_hours / total_days if total_days > 0 else 0
            
            attendance_summary = {
                'total_days': total_days,
                'present_days': present_days,
                'late_days': late_days,
                'absent_days': absent_days,
                'total_work_hours': round(total_work_hours, 2),
                'average_daily_hours': round(avg_daily_hours, 2),
                'attendance_rate': round((present_days / total_days * 100) if total_days > 0 else 0, 2),
            }
            
            return attendance_summary
            
        except Exception as e:
            return {
                'error': f"Gagal generate attendance summary: {str(e)}"
            }

    def generate_corrections_summary(self, employee, period, corrections_data):
        """
        Generate attendance corrections summary for the period
        """
        try:
            total_corrections = corrections_data.count()
            approved_corrections = corrections_data.filter(status='approved').count()
            pending_corrections = corrections_data.filter(status='pending').count()
            rejected_corrections = corrections_data.filter(status='rejected').count()
            
            corrections_summary = {
                'total_corrections': total_corrections,
                'approved_corrections': approved_corrections,
                'pending_corrections': pending_corrections,
                'rejected_corrections': rejected_corrections,
                'approval_rate': round((approved_corrections / total_corrections * 100) if total_corrections > 0 else 0, 2),
            }
            
            return corrections_summary
            
        except Exception as e:
            return {
                'error': f"Gagal generate corrections summary: {str(e)}"
            }

    def replace_placeholders_in_document(self, doc, replacements):
        """Replace placeholders in document with actual values"""
        try:
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Special handling for {{TABEL_OVERTIME}} placeholder
            if '{{TABEL_OVERTIME}}' in replacements:
                self.replace_tabel_overtime_placeholder(doc, overtime_data, total_hours, total_amount)
                
        except Exception as e:
            print(f"Warning: Could not replace all placeholders: {str(e)}")

    def replace_tabel_overtime_placeholder(self, doc, overtime_data, total_hours, total_amount):
        """Replace {{TABEL_OVERTIME}} placeholder with actual overtime table"""
        try:
            # Find paragraphs containing the placeholder and replace with table
            for paragraph in doc.paragraphs:
                if '{{TABEL_OVERTIME}}' in paragraph.text:
                    # Clear the paragraph
                    paragraph.clear()
                    # Add heading for overtime table
                    heading = paragraph.add_run('Detail Pengajuan Overtime')
                    heading.bold = True
                    heading.font.size = Pt(14)
                    break
            
            # Find tables containing the placeholder and replace with actual data
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{TABEL_OVERTIME}}' in paragraph.text:
                                # Clear the cell content
                                paragraph.clear()
                                # Add note that table will be inserted
                                paragraph.add_run('Tabel overtime akan di-insert di bawah ini')
                                break
            
            # Insert the actual overtime table after the placeholder
            self.insert_overtime_table_after_placeholder(doc, overtime_data, total_hours, total_amount)
                                
        except Exception as e:
            print(f"Warning: Could not replace {{TABEL_OVERTIME}} placeholder: {str(e)}")

    def insert_overtime_table_after_placeholder(self, doc, overtime_data, total_hours, total_amount):
        """Insert overtime table after finding {{TABEL_OVERTIME}} placeholder"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            
            # Find the position where to insert the table
            insert_position = None
            
            # Look for paragraphs containing the placeholder
            for i, paragraph in enumerate(doc.paragraphs):
                if '{{TABEL_OVERTIME}}' in paragraph.text or 'Tabel overtime akan di-insert di bawah ini' in paragraph.text:
                    insert_position = i
                    break
            
            # Look for tables containing the placeholder
            if insert_position is None:
                for i, table in enumerate(doc.tables):
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '{{TABEL_OVERTIME}}' in paragraph.text or 'Tabel overtime akan di-insert di bawah ini' in paragraph.text:
                                    # Find the paragraph after this table
                                    table_index = doc.element.index(table.element)
                                    insert_position = table_index + 1
                                    break
                            if insert_position is not None:
                                break
                        if insert_position is not None:
                            break
                    if insert_position is not None:
                        break
            
            # If we found a position, insert the table
            if insert_position is not None:
                # Create the overtime table
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                
                # Set column widths
                table.columns[0].width = Inches(0.5)   # No
                table.columns[1].width = Inches(1.0)   # Tanggal
                table.columns[2].width = Inches(1.0)   # Jam Lembur
                table.columns[3].width = Inches(2.5)   # Deskripsi
                table.columns[4].width = Inches(1.2)   # Status
                table.columns[5].width = Inches(1.2)   # Gaji Lembur
                
                # Add header row
                header_row = table.rows[0]
                headers = ['No', 'Tanggal', 'Jam Lembur', 'Deskripsi Pekerjaan', 'Status', 'Gaji Lembur']
                
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    # Style header
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for i, record in enumerate(overtime_data, 1):
                    row = table.add_row()
                    
                    # Format data
                    date_str = record.date_requested.strftime('%d/%m/%Y')
                    hours_str = f"{float(record.overtime_hours):.1f} jam"
                    
                    # Truncate description
                    description = record.work_description
                    if len(description) > 50:
                        description = description[:50] + '...'
                    
                    # Format status
                    status_map = {
                        'pending': 'Menunggu',
                        'level1_approved': 'Level 1 Approved',
                        'approved': 'Final Approved',
                        'rejected': 'Ditolak'
                    }
                    status_display = status_map.get(record.status, record.status)
                    
                    # Format amount
                    amount_str = f"{float(record.overtime_amount):.2f} AED"
                    
                    # Populate row
                    row_data = [str(i), date_str, hours_str, description, status_display, amount_str]
                    
                    for j, data in enumerate(row_data):
                        cell = row.cells[j]
                        cell.text = data
                        
                        # Style cells
                        for paragraph in cell.paragraphs:
                            if j == 0:  # No column
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif j == 2:  # Jam Lembur
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif j == 4:  # Status
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif j == 5:  # Gaji Lembur
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add summary row
                summary_row = table.add_row()
                summary_row.cells[0].text = 'TOTAL'
                summary_row.cells[1].text = ''
                summary_row.cells[2].text = f"{total_hours:.1f} jam"
                summary_row.cells[3].text = ''
                summary_row.cells[4].text = ''
                summary_row.cells[5].text = f"{total_amount:.2f} AED"
                
                # Style summary row
                for cell in summary_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add spacing after table
                doc.add_paragraph()
                
        except Exception as e:
            print(f"Warning: Could not insert overtime table: {str(e)}")

    def add_overtime_table_to_document(self, doc, overtime_data, total_hours, total_amount):
        """Add overtime table directly to document with professional styling"""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Add heading
        doc.add_heading('Detail Pengajuan Overtime', level=2)
        
        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(0.5)   # No
        table.columns[1].width = Inches(1.0)   # Tanggal
        table.columns[2].width = Inches(1.0)   # Jam Lembur
        table.columns[3].width = Inches(2.5)   # Deskripsi
        table.columns[4].width = Inches(1.2)   # Status
        table.columns[5].width = Inches(1.2)   # Gaji Lembur
        
        # Add header row
        header_row = table.rows[0]
        headers = ['No', 'Tanggal', 'Jam Lembur', 'Deskripsi Pekerjaan', 'Status', 'Gaji Lembur']
        
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Style header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for i, record in enumerate(overtime_data, 1):
            row = table.add_row()
            
            # Format data
            date_str = record.date_requested.strftime('%d/%m/%Y')
            hours_str = f"{float(record.overtime_hours):.1f} jam"
            
            # Truncate description
            description = record.work_description
            if len(description) > 50:
                description = description[:50] + '...'
            
            # Format status
            status_map = {
                'pending': 'Menunggu',
                'level1_approved': 'Level 1 Approved',
                'approved': 'Final Approved',
                'rejected': 'Ditolak'
            }
            status_display = status_map.get(record.status, record.status)
            
            # Format amount
            amount_str = f"{float(record.overtime_amount):.2f} AED"
            
            # Populate row
            row_data = [str(i), date_str, hours_str, description, status_display, amount_str]
            
            for j, data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = data
                
                # Style cells
                for paragraph in cell.paragraphs:
                    if j == 0:  # No column
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 2:  # Jam Lembur
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 4:  # Status
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 5:  # Gaji Lembur
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add summary row
        summary_row = table.add_row()
        summary_row.cells[0].text = 'TOTAL'
        summary_row.cells[1].text = ''
        summary_row.cells[2].text = f"{total_hours:.1f} jam"
        summary_row.cells[3].text = ''
        summary_row.cells[4].text = ''
        summary_row.cells[5].text = f"{total_amount:.2f} AED"
        
        # Style summary row
        for cell in summary_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add spacing
        doc.add_paragraph()

    def convert_docx_to_pdf(self, docx_file_path):
        """Convert DOCX file to PDF using docx2pdf library"""
        try:
            from docx2pdf import convert
            import tempfile
            import os
            
            # Create temporary PDF file
            pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            pdf_temp.close()
            
            # Convert DOCX to PDF
            convert(docx_file_path, pdf_temp.name)
            
            return pdf_temp.name
            
        except ImportError:
            # Fallback: if docx2pdf is not available, try alternative method
            try:
                from docx2pdf import convert
                import tempfile
                import os
                
                # Create temporary PDF file
                pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                pdf_temp.close()
                
                # Convert DOCX to PDF
                convert(docx_file_path, pdf_temp.name)
                
                return pdf_temp.name
                
            except Exception as e:
                # If all methods fail, raise error
                raise Exception(f"Tidak dapat mengkonversi DOCX ke PDF: {str(e)}")
                
        except Exception as e:
            raise Exception(f"Error saat konversi DOCX ke PDF: {str(e)}")

    def _convert_docx_to_pdf_reportlab(self, docx_file_path):
        """Convert DOCX to PDF using reportlab with advanced template matching"""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            import tempfile
            import os
            
            # Read DOCX content
            doc = Document(docx_file_path)
            
            # Create temporary PDF file
            pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            pdf_temp.close()
            
            # Create PDF using reportlab
            pdf_doc = SimpleDocTemplate(
                pdf_temp.name, 
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            # Create custom styles that match DOCX template
            styles = getSampleStyleSheet()
            
            # Custom styles for better template matching
            custom_styles = {
                'Title': ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontSize=18,
                    spaceAfter=20,
                    alignment=TA_CENTER,
                    fontName='Helvetica-Bold',
                    textColor=colors.black
                ),
                'Heading1': ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontSize=16,
                    spaceAfter=15,
                    spaceBefore=20,
                    alignment=TA_LEFT,
                    fontName='Helvetica-Bold',
                    textColor=colors.darkblue
                ),
                'Heading2': ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontSize=14,
                    spaceAfter=12,
                    spaceBefore=15,
                    alignment=TA_LEFT,
                    fontName='Helvetica-Bold',
                    textColor=colors.darkblue
                ),
                'Normal': ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    spaceAfter=8,
                    alignment=TA_JUSTIFY,
                    fontName='Helvetica',
                    textColor=colors.black,
                    leftIndent=0
                ),
                'Signature': ParagraphStyle(
                    'CustomSignature',
                    parent=styles['Normal'],
                    fontSize=11,
                    spaceAfter=20,
                    spaceBefore=30,
                    alignment=TA_RIGHT,
                    fontName='Helvetica',
                    textColor=colors.black
                )
            }
            
            story = []
            
            # Process paragraphs with advanced formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Determine style based on paragraph properties
                    style = self._determine_paragraph_style(paragraph, custom_styles)
                    
                    # Handle special formatting
                    formatted_text = self._format_paragraph_text(paragraph)
                    
                    p = Paragraph(formatted_text, style)
                    story.append(p)
                    
                    # Add spacing based on paragraph type
                    if paragraph.style.name.startswith('Heading'):
                        story.append(Spacer(1, 15))
                    else:
                        story.append(Spacer(1, 8))
            
            # Process tables with enhanced styling
            for table in doc.tables:
                table_data = []
                table_style = []
                
                for i, row in enumerate(table.rows):
                    row_data = []
                    for j, cell in enumerate(row.cells):
                        # Clean cell text and handle formatting
                        cell_text = self._clean_cell_text(cell.text)
                        row_data.append(cell_text)
                    
                    table_data.append(row_data)
                
                if table_data:
                    # Create enhanced table styling
                    pdf_table = Table(table_data, repeatRows=1)
                    table_style = self._create_table_style(table_data, len(table_data), len(table_data[0]))
                    pdf_table.setStyle(table_style)
                    
                    story.append(Spacer(1, 15))
                    story.append(pdf_table)
                    story.append(Spacer(1, 20))
            
            # Build PDF
            pdf_doc.build(story)
            
            return pdf_temp.name
            
        except Exception as e:
            raise Exception(f"Gagal konversi menggunakan reportlab: {str(e)}")
    
    def _determine_paragraph_style(self, paragraph, custom_styles):
        """Determine the appropriate style for a paragraph based on DOCX properties"""
        try:
            # Check for title-like paragraphs
            if paragraph.style.name.startswith('Title') or 'title' in paragraph.text.lower():
                return custom_styles['Title']
            
            # Check for heading-like paragraphs
            if paragraph.style.name.startswith('Heading'):
                if '1' in paragraph.style.name:
                    return custom_styles['Heading1']
                else:
                    return custom_styles['Heading2']
            
            # Check for signature-like paragraphs
            if any(keyword in paragraph.text.lower() for keyword in ['ttd', 'signature', 'dibuat', 'disetujui']):
                return custom_styles['Signature']
            
            # Default to normal style
            return custom_styles['Normal']
            
        except:
            return custom_styles['Normal']
    
    def _format_paragraph_text(self, paragraph):
        """Format paragraph text with special characters and formatting"""
        try:
            text = paragraph.text
            
            # Handle common Indonesian document formatting
            text = text.replace(':', ' : ')
            text = text.replace(';', ' ; ')
            
            # Handle bullet points and numbering
            if text.strip().startswith(('•', '-', '*', '1.', '2.', '3.')):
                text = '&nbsp;&nbsp;&nbsp;&nbsp;' + text
            
            return text
            
        except:
            return paragraph.text
    
    def _clean_cell_text(self, cell_text):
        """Clean and format cell text for table display"""
        try:
            # Remove extra whitespace
            cleaned = ' '.join(cell_text.split())
            
            # Handle empty cells
            if not cleaned:
                return '-'
            
            # Truncate very long text
            if len(cleaned) > 50:
                cleaned = cleaned[:47] + '...'
            
            return cleaned
            
        except:
            return cell_text if cell_text else '-'
    
    def _create_table_style(self, table_data, rows, cols):
        """Create enhanced table styling that matches DOCX template"""
        try:
            from reportlab.lib import colors
            
            # Base table style
            style = [
                # Header row styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),
                
                # Data rows styling
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 8),
                
                # Grid styling
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                
                # Column width adjustments
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]
            
            # Add specific column alignments for common overtime request fields
            if cols >= 4:  # Assuming standard overtime request table
                # Date column - center aligned
                style.append(('ALIGN', (0, 1), (0, -1), 'CENTER'))
                # Hours column - center aligned
                style.append(('ALIGN', (1, 1), (1, -1), 'CENTER'))
                # Amount column - right aligned
                style.append(('ALIGN', (3, 1), (3, -1), 'RIGHT'))
            
            return TableStyle(style)
            
        except Exception as e:
            # Fallback to basic styling
            return TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ])

    def generate_overtime_list_pdf_simple(self, overtime_requests, month=None):
        """Generate PDF for list of overtime requests"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from io import BytesIO
            from django.http import HttpResponse
            import datetime
            
            # Create buffer for PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            header_style = ParagraphStyle(
                'CustomHeader',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=12,
                alignment=TA_LEFT,
                textColor=colors.darkblue
            )
            
            # Build content
            story = []
            
            # Title
            title = Paragraph("DAFTAR PENGAJUAN LEMBUR", title_style)
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Summary information
            total_requests = overtime_requests.count()
            total_hours = sum(float(req.overtime_hours) for req in overtime_requests)
            total_amount = sum(float(req.overtime_amount) for req in overtime_requests)
            
            # Status summary
            status_counts = {}
            for req in overtime_requests:
                status = req.get_status_display()
                status_counts[status] = status_counts.get(status, 0) + 1
            
            # Summary table
            summary_data = [
                ['Total Pengajuan', str(total_requests)],
                ['Total Jam Lembur', f"{total_hours:.2f} jam"],
                ['Total Biaya Lembur', f"Rp {total_amount:,.2f}"],
                ['Rata-rata per Pengajuan', f"{total_hours/total_requests:.2f} jam" if total_requests > 0 else "0 jam"]
            ]
            
            # Add status breakdown
            for status, count in status_counts.items():
                summary_data.append([f"Status: {status}", str(count)])
            
            summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 20))
            
            # Main data table
            table_data = [['No', 'Tanggal', 'Nama Pegawai', 'NIP', 'Jam Lembur', 'Status', 'Biaya (Rp)']]
            
            for idx, request in enumerate(overtime_requests, 1):
                # Format date - handle both date and datetime objects
                if hasattr(request.date_requested, 'strftime'):
                    date_str = request.date_requested.strftime('%d/%m/%Y')
                else:
                    # If it's already a string or other format, convert to date first
                    if isinstance(request.date_requested, str):
                        date_obj = datetime.datetime.strptime(request.date_requested, '%Y-%m-%d').date()
                    else:
                        date_obj = request.date_requested
                    date_str = date_obj.strftime('%d/%m/%Y')
                
                # Get employee info
                employee_name = request.employee.fullname if request.employee else "N/A"
                employee_nip = request.employee.nip if request.employee else "N/A"
                
                # Format overtime hours
                hours_str = f"{request.overtime_hours:.2f}"
                
                # Get status display
                status_display = request.get_status_display()
                
                # Format amount
                amount_str = f"{request.overtime_amount:,.2f}"
                
                table_data.append([
                    str(idx),
                    date_str,
                    employee_name,
                    employee_nip,
                    hours_str,
                    status_display,
                    amount_str
                ])
            
            # Create table
            table = Table(table_data, colWidths=[0.5*inch, 1*inch, 2*inch, 1.2*inch, 0.8*inch, 1.2*inch, 1.2*inch])
            table.setStyle(self._create_table_style(table_data, len(table_data), len(table_data[0])))
            
            story.append(table)
            
            # Add footer with generation info
            story.append(Spacer(1, 20))
            footer_text = f"Dicetak pada: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
            if month:
                footer_text += f" | Periode: {month}"
            
            footer = Paragraph(footer_text, styles['Normal'])
            story.append(footer)
            
            # Build PDF
            doc.build(story)
            
            # Get PDF content
            pdf_content = buffer.getvalue()
            buffer.close()
            
            # Create response
            response = HttpResponse(content_type='application/pdf')
            filename = f"Daftar_Pengajuan_Lembur_{month or 'All'}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response.write(pdf_content)
            
            return response
            
        except Exception as e:
            raise Exception(f"Error generating PDF: {str(e)}")

    def generate_monthly_overtime_pdf(self, overtime_records, employee, month, start_date, end_date):
        """Generate monthly overtime PDF directly using ReportLab (no DOCX conversion)"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from io import BytesIO
            from django.utils import timezone
            import locale
            
            # Set locale for Indonesian formatting
            try:
                locale.setlocale(locale.LC_ALL, 'id_ID.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_ALL, 'id_ID')
                except:
                    pass
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            # Create custom styles
            styles = getSampleStyleSheet()
            
            # Title style (matching DOCX template)
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=16,
                spaceAfter=20,
                spaceBefore=0,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                textColor=colors.black
            )
            
            # Header style
            header_style = ParagraphStyle(
                'CustomHeader',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=15,
                spaceBefore=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                textColor=colors.darkblue
            )
            
            # Subheader style
            subheader_style = ParagraphStyle(
                'CustomSubheader',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=10,
                spaceBefore=15,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                textColor=colors.darkblue
            )
            
            # Normal style
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=8,
                alignment=TA_JUSTIFY,
                fontName='Helvetica',
                textColor=colors.black
            )
            
            # Signature style
            signature_style = ParagraphStyle(
                'CustomSignature',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=20,
                spaceBefore=30,
                alignment=TA_RIGHT,
                fontName='Helvetica',
                textColor=colors.black
            )
            
            # Build content
            story = []
            
            # Title
            title = Paragraph("DAFTAR PENGAJUAN LEMBUR", title_style)
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Employee Information
            employee_name = employee.fullname or employee.user.get_full_name() or employee.user.username
            employee_nip = employee.nip or '-'
            employee_position = employee.position.name if employee.position else '-'
            employee_division = employee.division.name if employee.division else '-'
            
            # Format month name in Indonesian
            month_names = [
                'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
                'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
            ]
            year, month_num = month.split('-')
            month_name = month_names[int(month_num) - 1]
            
            # Employee info table
            employee_info_data = [
                ['Nama', ':', employee_name],
                ['NIP', ':', employee_nip],
                ['Jabatan', ':', employee_position],
                ['Divisi', ':', employee_division],
                ['Periode', ':', f"{month_name} {year}"],
            ]
            
            employee_table = Table(employee_info_data, colWidths=[2*cm, 0.5*cm, 8*cm])
            employee_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTNAME', (2, 0), (2, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('ALIGN', (2, 0), (2, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            story.append(employee_table)
            story.append(Spacer(1, 20))
            
            # Summary Statistics
            total_requests = overtime_records.count()
            total_hours = sum(float(record.overtime_hours) for record in overtime_records)
            total_amount = sum(float(record.overtime_amount) for record in overtime_records)
            avg_hours = total_hours / total_requests if total_requests > 0 else 0
            avg_amount = total_amount / total_requests if total_requests > 0 else 0
            
            summary_data = [
                ['Total Pengajuan', ':', f"{total_requests} kali"],
                ['Total Jam Lembur', ':', f"{total_hours:.2f} jam"],
                ['Total Biaya Lembur', ':', f"Rp {total_amount:,.2f}"],
                ['Rata-rata per Pengajuan', ':', f"{avg_hours:.2f} jam / Rp {avg_amount:,.2f}"],
            ]
            
            summary_table = Table(summary_data, colWidths=[4*cm, 0.5*cm, 6*cm])
            summary_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTNAME', (2, 0), (2, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('ALIGN', (2, 0), (2, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 20))
            
            # Detailed Overtime Records
            if overtime_records.exists():
                # Table header
                table_data = [['No', 'Tanggal', 'Jam Lembur', 'Deskripsi Pekerjaan', 'Status', 'Jumlah (Rp)']]
                
                for idx, record in enumerate(overtime_records, 1):
                    # Format date - handle both date and datetime objects
                    if hasattr(record.date_requested, 'strftime'):
                        date_str = record.date_requested.strftime('%d/%m/%Y')
                    else:
                        # If it's already a string or other format, convert to date first
                        from datetime import datetime
                        if isinstance(record.date_requested, str):
                            date_obj = datetime.strptime(record.date_requested, '%Y-%m-%d').date()
                        else:
                            date_obj = record.date_requested
                        date_str = date_obj.strftime('%d/%m/%Y')
                    
                    # Format hours
                    hours_str = f"{record.overtime_hours:.2f} jam"
                    
                    # Format status
                    status_map = {
                        'pending': 'Menunggu',
                        'level1_approved': 'Disetujui Level 1',
                        'approved': 'Disetujui',
                        'rejected': 'Ditolak'
                    }
                    status_str = status_map.get(record.status, record.status.title())
                    
                    # Format amount
                    amount_str = f"Rp {record.overtime_amount:,.2f}"
                    
                    # Truncate description if too long
                    description = record.work_description[:50] + '...' if len(record.work_description) > 50 else record.work_description
                    
                    table_data.append([
                        str(idx),
                        date_str,
                        hours_str,
                        description,
                        status_str,
                        amount_str
                    ])
                
                # Create table
                records_table = Table(table_data, colWidths=[1*cm, 2*cm, 2*cm, 6*cm, 2.5*cm, 2.5*cm])
                records_table.setStyle(TableStyle([
                    # Header styling
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('TOPPADDING', (0, 0), (-1, 0), 8),
                    
                    # Data rows styling
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                    ('TOPPADDING', (0, 1), (-1, -1), 6),
                    
                    # Alignment
                    ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # No
                    ('ALIGN', (1, 1), (1, -1), 'CENTER'),  # Tanggal
                    ('ALIGN', (2, 1), (2, -1), 'CENTER'),  # Jam Lembur
                    ('ALIGN', (3, 1), (3, -1), 'LEFT'),    # Deskripsi
                    ('ALIGN', (4, 1), (4, -1), 'CENTER'),  # Status
                    ('ALIGN', (5, 1), (5, -1), 'RIGHT'),   # Jumlah
                    
                    # Grid
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    
                    # Alternating row colors
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ]))
                
                story.append(records_table)
                story.append(Spacer(1, 20))
            
            # Footer with date and signature
            current_date = dj_timezone.now()
            footer_text = f"Dibuat pada: {current_date.strftime('%d %B %Y')}"
            footer = Paragraph(footer_text, normal_style)
            story.append(footer)
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            # Create response
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            filename = f"Laporan_Overtime_{employee.nip}_{month}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response['Pragma'] = 'no-cache'
            response['Expires'] = '0'
            
            return response
            
        except Exception as e:
            raise Exception(f"Gagal generate PDF: {str(e)}")


# Group ViewSets
class GroupViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk manajemen Group dengan role-based access control
    """
    permission_classes = [IsAdminOrReadOnly]
    pagination_class = DefaultPagination
    
    def get_queryset(self):
        return Group.objects.all().order_by('name')
    
    def get_serializer_class(self):
        if self.action == 'create':
            return GroupCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return GroupUpdateSerializer
        elif self.action == 'retrieve':
            return GroupDetailSerializer
        return GroupSerializer
    
    def perform_create(self, serializer):
        serializer.save()
    
    def perform_update(self, serializer):
        serializer.save()
    
    def perform_destroy(self, instance):
        # Check if group has users before deletion
        if instance.user_set.exists():
            from rest_framework.exceptions import ValidationError
            raise ValidationError("Cannot delete group that has users. Please remove all users from this group first.")
        instance.delete()


class AdminGroupViewSet(viewsets.ModelViewSet):
    """
    Admin ViewSet untuk manajemen Group - Full access
    """
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination
    
    def get_queryset(self):
        return Group.objects.all().order_by('name')
    
    def get_serializer_class(self):
        if self.action == 'create':
            return GroupCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return GroupUpdateSerializer
        elif self.action == 'retrieve':
            return GroupDetailSerializer
        return GroupAdminSerializer
    
    def perform_create(self, serializer):
        serializer.save()
    
    def perform_update(self, serializer):
        serializer.save()
    
    def perform_destroy(self, instance):
        # Check if group has users before deletion
        if instance.user_set.exists():
            from rest_framework.exceptions import ValidationError
            raise ValidationError("Cannot delete group that has users. Please remove all users from this group first.")
        instance.delete()


class SupervisorGroupViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Supervisor ViewSet untuk Group - Read-only access
    """
    permission_classes = [IsAdminOrSupervisorReadOnly]
    pagination_class = DefaultPagination
    serializer_class = GroupSerializer
    
    def get_queryset(self):
        return Group.objects.all().order_by('name')


class EmployeeGroupViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Employee ViewSet untuk Group - Read-only access
    """
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = DefaultPagination
    serializer_class = GroupSerializer
    
    def get_queryset(self):
        return Group.objects.all().order_by('name')


class GroupPermissionViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk manajemen GroupPermission
    """
    permission_classes = [IsAdmin]
    
    def list(self, request, *args, **kwargs):
        """
        Override list method untuk allow anonymous access di development
        """
        # Allow anonymous access for development
        if settings.DEBUG:
            self.permission_classes = []
        return super().list(request, *args, **kwargs)
    pagination_class = DefaultPagination
    
    def get_queryset(self):
        return GroupPermission.objects.select_related('group').all().order_by('group__name', 'permission_type', 'permission_action')
    
    def get_serializer_class(self):
        if self.action == 'create':
            return GroupPermissionCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return GroupPermissionUpdateSerializer
        elif self.action == 'retrieve':
            return GroupPermissionDetailSerializer
        return GroupPermissionSerializer


class GroupPermissionTemplateViewSet(viewsets.ModelViewSet):
    """
    ViewSet untuk manajemen GroupPermissionTemplate
    """
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination
    serializer_class = GroupPermissionTemplateSerializer
    
    def get_queryset(self):
        return GroupPermissionTemplate.objects.all().order_by('name')
    
    def get_serializer_class(self):
        if self.action == 'create':
            return GroupPermissionTemplateCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return GroupPermissionTemplateUpdateSerializer
        return GroupPermissionTemplateSerializer
    
    @action(detail=True, methods=['post'])
    def apply_to_group(self, request, pk=None):
        """Apply template permissions to a specific group"""
        template = self.get_object()
        group_id = request.data.get('group_id')
        
        if not group_id:
            return JsonResponse({"detail": "group_id is required"}, status=400)
        
        try:
            group = Group.objects.get(id=group_id)
            template.apply_to_group(group)
            return JsonResponse({"detail": f"Template '{template.name}' applied to group '{group.name}' successfully"})
        except Group.DoesNotExist:
            return JsonResponse({"detail": "Group not found"}, status=404)
        except Exception as e:
            return JsonResponse({"detail": str(e)}, status=500)


class AdminGroupWithPermissionsViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Admin ViewSet untuk Group dengan permissions detail
    """
    permission_classes = [IsAdmin]
    pagination_class = DefaultPagination
    serializer_class = GroupWithPermissionsSerializer
    
    def get_queryset(self):
        return Group.objects.prefetch_related('custom_permissions').all().order_by('name')


class PermissionManagementViewSet(viewsets.ViewSet):
    """
    ViewSet untuk manajemen permission secara bulk
    """
    permission_classes = [IsAdmin]
    
    @action(detail=False, methods=['post'])
    def bulk_update(self, request):
        """Bulk update permissions for multiple groups"""
        print(f"DEBUG: Received data: {request.data}")
        print(f"DEBUG: Data type: {type(request.data)}")
        serializer = BulkPermissionUpdateSerializer(data=request.data, many=True)
        if serializer.is_valid():
            try:
                with transaction.atomic():
                    for item in serializer.validated_data:
                        group_id = item['group_id']
                        permissions = item['permissions']
                        
                        # Clear existing permissions for this group
                        GroupPermission.objects.filter(group_id=group_id).delete()
                        
                        # Create new permissions
                        for perm in permissions:
                            GroupPermission.objects.create(
                                group_id=group_id,
                                permission_type=perm['permission_type'],
                                permission_action=perm['permission_action'],
                                is_active=perm.get('is_active', True)
                            )
                
                return JsonResponse({"detail": "Permissions updated successfully"})
            except Exception as e:
                return JsonResponse({"detail": str(e)}, status=500)
        return JsonResponse({"errors": serializer.errors}, status=400)
    
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get permission summary for all groups"""
        groups = Group.objects.prefetch_related('custom_permissions').all()
        summary_data = []
        
        for group in groups:
            permissions = group.custom_permissions.filter(is_active=True)
            permission_types = list(set(perm.permission_type for perm in permissions))
            permission_actions = list(set(perm.permission_action for perm in permissions))
            
            summary_data.append({
                'group_id': group.id,
                'group_name': group.name,
                'total_permissions': group.custom_permissions.count(),
                'active_permissions': permissions.count(),
                'permission_types': permission_types,
                'permission_actions': permission_actions
            })
        
        return JsonResponse(summary_data, safe=False)
    
    @action(detail=False, methods=['post'])
    def copy_permissions(self, request):
        """Copy permissions from one group to another"""
        source_group_id = request.data.get('source_group_id')
        target_group_id = request.data.get('target_group_id')
        
        if not source_group_id or not target_group_id:
            return JsonResponse({"detail": "Both source_group_id and target_group_id are required"}, status=400)
        
        try:
            source_group = Group.objects.get(id=source_group_id)
            target_group = Group.objects.get(id=target_group_id)
            
            # Get source permissions
            source_permissions = GroupPermission.objects.filter(group=source_group, is_active=True)
            
            # Clear target permissions
            GroupPermission.objects.filter(group=target_group).delete()
            
            # Copy permissions
            for perm in source_permissions:
                GroupPermission.objects.create(
                    group=target_group,
                    permission_type=perm.permission_type,
                    permission_action=perm.permission_action,
                    is_active=perm.is_active
                )
            
            return JsonResponse({
                "detail": f"Permissions copied from '{source_group.name}' to '{target_group.name}' successfully",
                "copied_permissions": source_permissions.count()
            })
        except Group.DoesNotExist:
            return JsonResponse({"detail": "One or both groups not found"}, status=404)
        except Exception as e:
            return JsonResponse({"detail": str(e)}, status=500)


class AttendanceViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = AttendanceSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = DefaultPagination

    def get_queryset(self):
        user = self.request.user
        qs = Attendance.objects.select_related('employee').filter(user=user)
        start = self.request.query_params.get('start')
        end = self.request.query_params.get('end')
        month = self.request.query_params.get('month')
        
        if start:
            qs = qs.filter(date_local__gte=start)
        if end:
            qs = qs.filter(date_local__lte=end)
        if month:
            # Parse month and filter by year-month
            try:
                year, month_num = month.split('-')
                qs = qs.filter(
                    date_local__year=year,
                    date_local__month=month_num
                )
            except ValueError:
                pass  # Ignore invalid month format
        
        return qs.order_by('-date_local')


@api_view(['GET'])
@permission_classes([AllowAny])
@ensure_csrf_cookie
def csrf_token_view(request):
    """
    Get CSRF token for frontend
    """
    return JsonResponse({
        'csrfToken': get_token(request)
    })


@api_view(['GET'])
@permission_classes([AllowAny])
def public_permissions_view(request):
    """
    Get all available permissions for frontend (development only)
    """
    from django.contrib.auth.models import Permission
    from django.contrib.contenttypes.models import ContentType
    
    # Get all Django permissions
    permissions = Permission.objects.select_related('content_type').all()
    
    # Transform to match frontend interface
    results = []
    for perm in permissions:
        # Extract action from codename (e.g., "add_attendance" -> "add")
        action = perm.codename.split('_')[0] if '_' in perm.codename else perm.codename
        
        # Create human-readable name
        model_name = perm.content_type.model.replace('_', ' ').title()
        action_name = action.replace('_', ' ').title()
        
        # Fix codename to be just the action
        codename = action
        
        results.append({
            'id': perm.id,
            'name': f"{model_name} - {action_name}",
            'codename': codename,
            'content_type': perm.content_type.model,
            'permission_type': perm.content_type.model,
            'permission_action': action,
            'is_active': True
        })
    
    return JsonResponse({
        'count': len(results),
        'results': results
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def attendance_corrections(request):
    """
    Get attendance records that need corrections (missing attendance, WFA, etc.)
    """
    try:
        # Get user and employee
        user = request.user
        try:
            employee = Employee.objects.get(user=user)
        except Employee.DoesNotExist:
            return Response(
                {'detail': 'Employee not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

        # Get query parameters
        start_date = request.GET.get('start_date')
        end_date = request.GET.get('end_date')
        month = request.GET.get('month')
        status_filter = request.GET.get('status', 'all')

        # Build date filter
        date_filter = Q()
        if month:
            # Parse month (YYYY-MM format)
            try:
                year, month_num = month.split('-')
                start_of_month = datetime.date(int(year), int(month_num), 1)
                if int(month_num) == 12:
                    end_of_month = datetime.date(int(year) + 1, 1, 1) - timedelta(days=1)
                else:
                    end_of_month = datetime.date(int(year), int(month_num) + 1, 1) - timedelta(days=1)
                
                date_filter = Q(date_local__gte=start_of_month) & Q(date_local__lte=end_of_month)
            except (ValueError, TypeError):
                return Response(
                    {'detail': 'Invalid month format. Use YYYY-MM'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        elif start_date and end_date:
            try:
                start = datetime.strptime(start_date, '%Y-%m-%d').date()
                end = datetime.strptime(end_date, '%Y-%m-%d').date()
                date_filter = Q(date_local__gte=start) & Q(date_local__lte=end)
            except ValueError:
                return Response(
                    {'detail': 'Invalid date format. Use YYYY-MM-DD'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            # Default to current month if no date filter
            today = dj_timezone.now().date()
            start_of_month = today.replace(day=1)
            if today.month == 12:
                end_of_month = today.replace(year=today.year + 1, month=1, day=1) - timedelta(days=1)
            else:
                end_of_month = today.replace(month=today.month + 1, day=1) - timedelta(days=1)
            date_filter = Q(date_local__gte=start_of_month) & Q(date_local__lte=end_of_month)

        # Get attendance records for the user (not employee, since employee field might be null)
        attendance_records = Attendance.objects.filter(
            user=user
        ).filter(date_filter).order_by('-date_local')

        # Get work settings for geofence validation
        try:
            work_settings = WorkSettings.objects.first()
        except WorkSettings.DoesNotExist:
            work_settings = None

        # Process records to identify correction needs
        correction_records = []
        total_records = 0
        wfa_records = 0
        missing_attendance = 0
        pending_corrections = 0

        for record in attendance_records:
            total_records += 1
            
            # Check if record needs correction
            needs_correction = False
            correction_reasons = []
            
            # Check for missing attendance
            if not record.check_in_at_utc and not record.check_out_at_utc:
                needs_correction = True
                correction_reasons.append('missing_both')
                missing_attendance += 1
            elif not record.check_in_at_utc:
                needs_correction = True
                correction_reasons.append('missing_check_in')
                missing_attendance += 1
            elif not record.check_out_at_utc:
                needs_correction = True
                correction_reasons.append('missing_check_out')
                missing_attendance += 1
            
            # Check for WFA
            if record.check_in_at_utc and not record.within_geofence:
                needs_correction = True
                correction_reasons.append('wfa')
                wfa_records += 1
            
            # Check for system notes indicating issues
            if record.note and ('luar area kantor' in record.note.lower() or 'outside office' in record.note.lower()):
                needs_correction = True
                correction_reasons.append('system_note')
            
            # Check if correction request exists
            try:
                correction = AttendanceCorrection.objects.filter(
                    user=user,
                    date_local=record.date_local
                ).first()
                if correction:
                    correction_status = correction.status
                    if correction.status == 'pending':
                        pending_corrections += 1
                else:
                    correction_status = None
            except Exception:
                correction_status = None
            
            # Add to correction records if needs correction or has correction request
            if needs_correction or correction_status:
                correction_records.append({
                    'id': record.id,
                    'date_local': record.date_local,
                    'check_in_at_utc': record.check_in_at_utc.isoformat() if record.check_in_at_utc else None,
                    'check_out_at_utc': record.check_out_at_utc.isoformat() if record.check_out_at_utc else None,
                    'check_in_lat': record.check_in_lat,
                    'check_in_lng': record.check_in_lng,
                    'check_out_lat': record.check_out_lat,
                    'check_out_lng': record.check_out_lng,
                    'check_in_ip': record.check_in_ip,
                    'check_out_ip': record.check_out_ip,
                    'minutes_late': record.minutes_late,
                    'total_work_minutes': record.total_work_minutes,
                    'is_holiday': record.is_holiday,
                    'within_geofence': record.within_geofence,
                    'note': record.note,
                    'employee_note': record.employee_note,
                    'created_at': record.created_at.isoformat(),
                    'updated_at': record.updated_at.isoformat(),
                    'correction_status': correction_status,
                    'correction_reasons': correction_reasons,
                    # Add proposed correction times if correction request exists
                    'proposed_check_in_local': correction.proposed_check_in_local.isoformat() if correction and correction.proposed_check_in_local else None,
                    'proposed_check_out_local': correction.proposed_check_out_local.isoformat() if correction and correction.proposed_check_out_local else None,
                    'correction_type': correction.type if correction else None,
                    'correction_reason': correction.reason if correction else None
                })

        # Get manual correction requests that don't have corresponding attendance records
        manual_corrections = AttendanceCorrection.objects.filter(
            user=user
        ).filter(date_filter).exclude(
            date_local__in=[record.date_local for record in attendance_records]
        )

        for manual_correction in manual_corrections:
            # Add manual correction records with proposed times
            correction_records.append({
                'id': f"manual_{manual_correction.id}",
                'date_local': manual_correction.date_local,
                'check_in_at_utc': None,
                'check_out_at_utc': None,
                'check_in_lat': None,
                'check_in_lng': None,
                'check_out_lat': None,
                'check_out_lng': None,
                'check_in_ip': None,
                'check_out_ip': None,
                'minutes_late': 0,
                'total_work_minutes': 0,
                'is_holiday': False,
                'within_geofence': False,
                'note': None,
                'employee_note': None,
                'created_at': manual_correction.created_at.isoformat(),
                'updated_at': manual_correction.updated_at.isoformat(),
                'correction_status': manual_correction.status,
                'correction_reasons': ['manual_request'],
                'is_manual': True,
                # Add proposed correction times
                'proposed_check_in_local': manual_correction.proposed_check_in_local.isoformat() if manual_correction.proposed_check_in_local else None,
                'proposed_check_out_local': manual_correction.proposed_check_out_local.isoformat() if manual_correction.proposed_check_out_local else None,
                'correction_type': manual_correction.type,
                'correction_reason': manual_correction.reason
            })
            
            # Update summary counts
            if manual_correction.status == 'pending':
                pending_corrections += 1
            total_records += 1

        # Apply status filter if specified
        if status_filter != 'all':
            if status_filter == 'not_submitted':
                correction_records = [r for r in correction_records if not r['correction_status']]
            elif status_filter == 'pending':
                correction_records = [r for r in correction_records if r['correction_status'] == 'pending']
            elif status_filter == 'approved':
                correction_records = [r for r in correction_records if r['correction_status'] == 'approved']
            elif status_filter == 'rejected':
                correction_records = [r for r in correction_records if r['correction_status'] == 'rejected']

        # Prepare response
        response_data = {
            'correction_records': correction_records,
            'summary': {
                'total_records': total_records,
                'wfa_records': wfa_records,
                'missing_attendance': missing_attendance,
                'pending_corrections': pending_corrections
            }
        }

        return Response(response_data)

    except Exception as e:
        print(f"Error in attendance_corrections: {str(e)}")
        return Response(
            {'detail': 'Internal server error'}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([JSONParser, MultiPartParser, FormParser])
def attendance_correction_request(request):
    """
    Submit a new attendance correction request
    """
    try:
        # Get employee for the authenticated user
        try:
            employee = Employee.objects.get(user=request.user)
        except Employee.DoesNotExist:
            return Response(
                {'detail': 'Employee not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

        # Validate required fields
        date_local = request.data.get('date_local')
        correction_type = request.data.get('type')
        reason = request.data.get('reason')

        if not all([date_local, correction_type, reason]):
            return Response(
                {'detail': 'Missing required fields: date_local, type, reason'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validate date format
        try:
            date_obj = datetime.strptime(date_local, '%Y-%m-%d').date()
        except ValueError:
            return Response(
                {'detail': 'Invalid date format. Use YYYY-MM-DD'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check if date is not in the future
        if date_obj > dj_timezone.now().date():
            return Response(
                {'detail': 'Cannot submit correction for future dates'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check if correction request already exists for this date
        existing_correction = AttendanceCorrection.objects.filter(
            user=request.user,
            date_local=date_obj
        ).first()

        if existing_correction:
            if existing_correction.status == 'pending':
                return Response(
                    {'detail': 'Correction request already pending for this date'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            elif existing_correction.status == 'approved':
                return Response(
                    {'detail': 'Correction already approved for this date'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )

        # Map frontend type to model type (support both old and new formats)
        type_mapping = {
            'check_in': AttendanceCorrection.CorrectionType.MISSING_CHECK_IN,
            'check_out': AttendanceCorrection.CorrectionType.MISSING_CHECK_OUT,
            'both': AttendanceCorrection.CorrectionType.EDIT,
            'missing_check_in': AttendanceCorrection.CorrectionType.MISSING_CHECK_IN,
            'missing_check_out': AttendanceCorrection.CorrectionType.MISSING_CHECK_OUT,
            'edit': AttendanceCorrection.CorrectionType.EDIT
        }

        if correction_type not in type_mapping:
            return Response(
                {'detail': 'Invalid correction type'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # Create correction request
        correction_data = {
            'user': request.user,
            'date_local': date_obj,
            'type': type_mapping[correction_type],
            'reason': reason,
            'status': AttendanceCorrection.CorrectionStatus.PENDING
        }

        # Handle proposed times if provided (support both old and new field names)
        proposed_check_in = request.data.get('proposed_check_in_local') or request.data.get('proposed_check_in_time')
        proposed_check_out = request.data.get('proposed_check_out_local') or request.data.get('proposed_check_out_time')
        
        if proposed_check_in:
            try:
                # Parse time string to datetime object
                if isinstance(proposed_check_in, str):
                    if 'T' in proposed_check_in:
                        # ISO format: "2025-08-25T08:00:00"
                        dt = datetime.fromisoformat(proposed_check_in.replace('Z', '+00:00'))
                    elif ':' in proposed_check_in and len(proposed_check_in.split(':')) == 2:
                        # Time format: "08:00" - combine with date_local
                        time_obj = datetime.strptime(proposed_check_in, '%H:%M').time()
                        dt = datetime.combine(date_obj, time_obj)
                    else:
                        # Simple format: "2025-08-25 08:00:00"
                        dt = datetime.strptime(proposed_check_in, '%Y-%m-%d %H:%M:%S')
                    correction_data['proposed_check_in_local'] = dt
            except ValueError:
                return Response(
                    {'detail': 'Invalid proposed_check_in_time format'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )

        if proposed_check_out:
            try:
                # Parse time string to datetime object
                if isinstance(proposed_check_out, str):
                    if 'T' in proposed_check_out:
                        # ISO format: "2025-08-25T08:00:00"
                        dt = datetime.fromisoformat(proposed_check_out.replace('Z', '+00:00'))
                    elif ':' in proposed_check_out and len(proposed_check_out.split(':')) == 2:
                        # Time format: "08:00" - combine with date_local
                        time_obj = datetime.strptime(proposed_check_out, '%H:%M').time()
                        dt = datetime.combine(date_obj, time_obj)
                    else:
                        # Simple format: "2025-08-25 08:00:00"
                        dt = datetime.strptime(proposed_check_out, '%H:%M:%S')
                    correction_data['proposed_check_out_local'] = dt
            except ValueError:
                return Response(
                    {'detail': 'Invalid proposed_check_out_time format'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )

        # Handle attachment if provided
        attachment = request.FILES.get('attachment')
        if attachment:
            # Validate file type and size
            allowed_types = ['application/pdf', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'image/jpeg', 'image/png']
            if attachment.content_type not in allowed_types:
                return Response(
                    {'detail': 'Invalid file type. Only PDF, DOC, DOCX, JPG, and PNG are allowed'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            if attachment.size > 10 * 1024 * 1024:  # 10MB limit
                return Response(
                    {'detail': 'File size too large. Maximum size is 10MB'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            correction_data['attachment'] = attachment

        # Create the correction request
        correction = AttendanceCorrection.objects.create(**correction_data)

        return Response({
            'detail': 'Correction request submitted successfully',
            'correction_id': correction.id,
            'status': correction.status
        }, status=status.HTTP_201_CREATED)
        
    except Exception as e:
        print(f"Error in attendance_correction_request: {str(e)}")
        return Response(
            {'detail': 'Internal server error'}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

