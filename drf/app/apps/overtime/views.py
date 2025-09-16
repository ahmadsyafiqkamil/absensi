from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.utils import timezone
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework import filters
from datetime import date
from .models import OvertimeRequest, MonthlySummaryRequest
from .serializers import (
    OvertimeRequestSerializer, OvertimeRequestAdminSerializer,
    OvertimeRequestSupervisorSerializer, OvertimeRequestEmployeeSerializer,
    OvertimeRequestCreateUpdateSerializer, OvertimeRequestApprovalSerializer,
    OvertimeRequestListSerializer,
    MonthlySummaryRequestSerializer, MonthlySummaryRequestAdminSerializer,
    MonthlySummaryRequestSupervisorSerializer, MonthlySummaryRequestEmployeeSerializer,
    MonthlySummaryRequestCreateUpdateSerializer, MonthlySummaryRequestApprovalSerializer,
    MonthlySummaryRequestListSerializer
)
from apps.core.permissions import IsAdmin, IsSupervisor, IsEmployee
from .services import OvertimeService
from django.http import HttpResponse
from django.conf import settings
import os
import tempfile
import requests


# Overtime Request Views
class OvertimeRequestViewSet(viewsets.ModelViewSet):
    """Overtime request management ViewSet with role-based access"""
    serializer_class = OvertimeRequestSerializer
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['status', 'request_type', 'user', 'employee']
    search_fields = ['user__username', 'user__first_name', 'user__last_name', 'employee__name']
    ordering_fields = ['created_at', 'date', 'status']
    ordering = ['-created_at']
    
    def get_serializer_class(self):
        """Return appropriate serializer based on user role"""
        if self.action in ['create', 'update', 'partial_update']:
            return OvertimeRequestCreateUpdateSerializer
        elif self.request.user.is_superuser or self.request.user.groups.filter(name='admin').exists():
            return OvertimeRequestAdminSerializer
        elif self.request.user.groups.filter(name='supervisor').exists():
            return OvertimeRequestSupervisorSerializer
        else:
            return OvertimeRequestEmployeeSerializer
    
    def get_serializer_context(self):
        """Add request context to serializer"""
        context = super().get_serializer_context()
        context['request'] = self.request
        return context
    
    def get_queryset(self):
        """Filter overtime requests based on user role and date range"""
        queryset = OvertimeRequest.objects.all()
        
        # Apply role-based filtering
        if self.request.user.is_superuser or self.request.user.groups.filter(name='admin').exists():
            queryset = queryset
        elif self.request.user.groups.filter(name='supervisor').exists():
            # Supervisors can see overtime requests of employees in their division
            # If they have org-wide approval, they can see all overtime requests
            if hasattr(self.request.user, 'employee_profile') and self.request.user.employee_profile.position:
                if self.request.user.employee_profile.position.can_approve_overtime_org_wide:
                    # Org-wide approval: can see all overtime requests
                    queryset = queryset
                elif self.request.user.employee_profile.division:
                    # Division-only approval: can see only their division's requests
                    queryset = queryset.filter(
                        employee__division=self.request.user.employee_profile.division
                    )
                else:
                    queryset = queryset.none()
            else:
                queryset = queryset.none()
        else:
            # Regular employees can only see their own overtime requests
            queryset = queryset.filter(user=self.request.user)
        
        # Apply date range filtering if start_date and end_date parameters are provided
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        
        if start_date:
            try:
                start_date_obj = date.fromisoformat(start_date)
                queryset = queryset.filter(date__gte=start_date_obj)
            except ValueError:
                pass  # Invalid date format, ignore
        
        if end_date:
            try:
                end_date_obj = date.fromisoformat(end_date)
                queryset = queryset.filter(date__lte=end_date_obj)
            except ValueError:
                pass  # Invalid date format, ignore
        
        return queryset
    
    def perform_create(self, serializer):
        """Auto-set user when creating"""
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download approved overtime request as DOCX (v2)"""
        overtime_request = self.get_object()

        if overtime_request.status != 'approved':
            return Response(
                {"detail": "Hanya overtime yang sudah disetujui yang dapat didownload"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Generate DOCX on-the-fly using template when available
        try:
            docx_bytes, filename = self._generate_docx(overtime_request)
            response = HttpResponse(
                docx_bytes,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            return Response(
                {"detail": f"Gagal generate dokumen: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Export approved overtime request to PDF using converter service (v2)"""
        overtime_request = self.get_object()

        if overtime_request.status != 'approved':
            return Response(
                {"detail": "Overtime request must be approved to export PDF."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            # 1) Generate a temporary DOCX
            docx_bytes, base_filename = self._generate_docx(overtime_request)

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(docx_bytes)
                tmp_path = tmp.name

            try:
                # 2) Convert via converter service
                with open(tmp_path, 'rb') as f:
                    files = {
                        'file': (
                            'document.docx',
                            f,
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        )
                    }
                    data = {'method': 'file'}
                    converter_url = 'http://docx_converter:5000/convert'
                    r = requests.post(converter_url, files=files, data=data, timeout=60)
                    if r.status_code != 200:
                        return Response(
                            {"detail": f"DOCX converter error: {r.status_code}"},
                            status=status.HTTP_502_BAD_GATEWAY,
                        )

                    pdf_content = r.content
            finally:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

            # 3) Return PDF
            pdf_filename = base_filename.replace('.docx', '.pdf')
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
        except Exception as e:
            return Response(
                {"detail": f"Gagal export PDF: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    def _get_template_path(self):
        """Return best template path for overtime document or None."""
        template_dir = os.path.join(settings.BASE_DIR, 'template')
        if not os.path.isdir(template_dir):
            return None

        priority_names = [
            'template_SURAT_PERINTAH_KERJA_LEMBUR.docx',
            'template_overtime_clean.docx',
            'template_overtime_working.docx',
            'template_overtime_simple.docx',
        ]

        # Priority match
        for name in priority_names:
            path = os.path.join(template_dir, name)
            if os.path.exists(path):
                return path

        # Fallback to any latest .docx
        docx_files = [
            os.path.join(template_dir, f)
            for f in os.listdir(template_dir)
            if f.endswith('.docx') and not f.startswith('~$')
        ]
        if not docx_files:
            return None
        docx_files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
        return docx_files[0]

    def _generate_docx(self, overtime_request):
        """Generate DOCX bytes for the overtime_request. Returns (bytes, filename)."""
        try:
            from docx import Document
        except Exception:
            raise RuntimeError('python-docx is not installed')

        template_path = self._get_template_path()
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            # Minimal fallback document
            doc = Document()
            doc.add_heading('Surat Perintah Kerja Lembur', level=1)

        # Build replacements
        employee = overtime_request.employee
        employee_name = None
        if employee:
            # Prefer explicit display name
            try:
                employee_name = employee.display_name
            except Exception:
                employee_name = getattr(employee, 'fullname', None)
        if not employee_name:
            if employee and getattr(employee, 'user', None):
                employee_name = employee.user.get_full_name() or employee.user.username
            else:
                employee_name = '-'

        employee_nip = getattr(employee, 'nip', '-') if employee else '-'
        nip_18 = employee_nip if employee_nip and len(employee_nip) >= 18 else employee_nip
        if nip_18 and len(nip_18) < 18:
            nip_18 = nip_18 + ('0' * (18 - len(nip_18)))
        nip_9 = employee_nip[:9] if employee_nip else '-'

        division_name = employee.get_division_name() if employee else '-'
        position_name = employee.get_position_name() if employee else '-'

        current_dt = timezone.now()
        tahun = current_dt.strftime('%Y')
        bulan = current_dt.strftime('%B')
        hari = current_dt.strftime('%d')
        tanggal_doc = current_dt.strftime('%d %B %Y')

        tanggal_lembur = overtime_request.date.strftime('%d %B %Y') if overtime_request.date else '-'
        jam_lembur = f"{overtime_request.total_hours} jam"
        deskripsi = overtime_request.work_description
        jumlah = f"{overtime_request.total_amount or 0}"

        # Approver info
        def _approver_name(user):
            if not user:
                return '-'
            try:
                emp = getattr(user, 'employee_profile', None)
                if emp and getattr(emp, 'fullname', None):
                    return emp.fullname
            except Exception:
                pass
            return user.get_full_name() or user.username

        def _approver_nip(user):
            try:
                emp = getattr(user, 'employee_profile', None)
                return getattr(emp, 'nip', '-') if emp else '-'
            except Exception:
                return '-'

        lvl1_name = _approver_name(getattr(overtime_request, 'level1_approved_by', None))
        lvl1_nip = _approver_nip(getattr(overtime_request, 'level1_approved_by', None))
        lvl1_at = getattr(overtime_request, 'level1_approved_at', None)
        lvl1_date = lvl1_at.strftime('%d %B %Y') if lvl1_at else '-'

        final_name = _approver_name(getattr(overtime_request, 'final_approved_by', None))
        final_nip = _approver_nip(getattr(overtime_request, 'final_approved_by', None))
        final_at = getattr(overtime_request, 'final_approved_at', None)
        final_date = final_at.strftime('%d %B %Y') if final_at else '-'

        nomor_dok = f"{overtime_request.id}/SPKL/KJRI-DXB/{tahun}"

        replacements = {
            # Document info
            '{{NOMOR_DOKUMEN}}': nomor_dok,
            '{{TANGGAL_DOKUMEN}}': tanggal_doc,
            '{{TAHUN}}': tahun,
            '{{BULAN}}': bulan,
            '{{HARI}}': hari,

            # Employee info
            '{{NAMA_PEGAWAI}}': employee_name,
            '{{NIP_PEGAWAI}}': employee_nip or '-',
            '{{NIP}}': employee_nip or '-',
            '{{NIP_LENGKAP}}': nip_18 or '-',
            '{{NIP_18_DIGIT}}': nip_18 or '-',
            '{{NIP_9_DIGIT}}': nip_9 or '-',
            '{{JABATAN_PEGAWAI}}': position_name,
            '{{DIVISI_PEGAWAI}}': division_name,

            # Overtime details
            '{{TANGGAL_LEMBUR}}': tanggal_lembur,
            '{{JAM_LEMBUR}}': jam_lembur,
            '{{DESKRIPSI_PEKERJAAN}}': deskripsi,
            '{{JUMLAH_GAJI_LEMBUR}}': jumlah,

            # Approval info
            '{{LEVEL1_APPROVER}}': lvl1_name,
            '{{LEVEL1_APPROVER_NIP}}': lvl1_nip,
            '{{LEVEL1_APPROVAL_DATE}}': lvl1_date,
            '{{FINAL_APPROVER}}': final_name,
            '{{FINAL_APPROVER_NIP}}': final_nip,
            '{{FINAL_APPROVAL_DATE}}': final_date,
        }

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old, new in replacements.items():
                if old in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        inline[i].text = inline[i].text.replace(old, str(new))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old, new in replacements.items():
                            if old in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    inline[i].text = inline[i].text.replace(old, str(new))

        # Serialize to bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)
        try:
            with open(tmp_path, 'rb') as f:
                content = f.read()
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        base_filename = f"Surat_Perintah_Kerja_Lembur_{employee_nip or 'pegawai'}_{overtime_request.date}.docx"
        return content, base_filename

    @action(detail=True, methods=['post'], permission_classes=[IsSupervisor])
    def approve(self, request, pk=None):
        """Approve an overtime request (Level 1 or Final)"""
        try:
            overtime_request = self.get_object()
        except OvertimeRequest.DoesNotExist:
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        supervisor = request.user
        approval_level = request.data.get('approval_level')

        if approval_level == 1:
            if overtime_request.status != 'pending':
                return Response({"detail": "Request is not pending approval."}, status=status.HTTP_400_BAD_REQUEST)
            
            overtime_request.status = 'level1_approved'
            overtime_request.level1_approved_by = supervisor
            overtime_request.level1_approved_at = timezone.now()
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response(serializer.data)

        elif approval_level == 2:
            if overtime_request.status not in ['pending', 'level1_approved']:
                return Response({"detail": "Request is not awaiting final approval."}, status=status.HTTP_400_BAD_REQUEST)
            
            overtime_request.status = 'approved'
            # If level 1 was skipped, fill it in too
            if not overtime_request.level1_approved_by:
                overtime_request.level1_approved_by = supervisor
                overtime_request.level1_approved_at = timezone.now()
            
            overtime_request.final_approved_by = supervisor
            overtime_request.final_approved_at = timezone.now()
            overtime_request.save()
            
            serializer = self.get_serializer(overtime_request)
            return Response(serializer.data)

        else:
            return Response({"detail": "Invalid or missing approval_level."}, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['post'], permission_classes=[IsSupervisor])
    def reject(self, request, pk=None):
        """Reject an overtime request"""
        try:
            overtime_request = self.get_object()
        except OvertimeRequest.DoesNotExist:
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        supervisor = request.user
        rejection_reason = request.data.get('rejection_reason', '')

        # Check if request can be rejected
        if overtime_request.status not in ['pending', 'level1_approved']:
            return Response({"detail": "Request cannot be rejected in current status."}, status=status.HTTP_400_BAD_REQUEST)
        
        # Set rejection status and reason
        overtime_request.status = 'rejected'
        overtime_request.rejection_reason = rejection_reason
        overtime_request.save()
        
        serializer = self.get_serializer(overtime_request)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def pending(self, request):
        """Get pending overtime requests for approval"""
        if not (request.user.is_superuser or 
                request.user.groups.filter(name__in=['admin', 'supervisor']).exists()):
            return Response(
                {"error": "Only admins and supervisors can view pending overtime requests"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        queryset = self.get_queryset().filter(status='pending')
        serializer = OvertimeRequestListSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def my_overtime(self, request):
        """Get current user's overtime requests"""
        queryset = OvertimeRequest.objects.filter(user=request.user)
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def by_status(self, request):
        """Get overtime requests filtered by status"""
        status_filter = request.query_params.get('status', '')
        if status_filter:
            queryset = self.get_queryset().filter(status=status_filter)
        else:
            queryset = self.get_queryset()
        
        serializer = OvertimeRequestListSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def by_type(self, request):
        """Get overtime requests filtered by type"""
        request_type = request.query_params.get('type', '')
        if request_type:
            queryset = self.get_queryset().filter(request_type=request_type)
        else:
            queryset = self.get_queryset()
        
        serializer = OvertimeRequestListSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def potential_overtime(self, request):
        """Get attendance records that could be submitted as overtime requests"""
        user = self.request.user
        
        # Only employees can view their potential overtime
        if not user.groups.filter(name='pegawai').exists():
            return Response(
                {"detail": "Hanya pegawai yang dapat melihat potensi lembur"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            employee = user.employee_profile
        except:
            return Response(
                {"detail": "User tidak memiliki profil employee"}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get date range (default: last 30 days)
        from datetime import date, timedelta
        from apps.attendance.models import Attendance
        from apps.settings.models import WorkSettings
        
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        
        if start_date:
            try:
                start_date = date.fromisoformat(start_date)
            except ValueError:
                return Response(
                    {"detail": "Format start_date tidak valid. Gunakan YYYY-MM-DD"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            start_date = date.today() - timedelta(days=30)
        
        if end_date:
            try:
                end_date = date.fromisoformat(end_date)
            except ValueError:
                return Response(
                    {"detail": "Format end_date tidak valid. Gunakan YYYY-MM-DD"}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            end_date = date.today()
        
        # Get work settings
        try:
            ws = WorkSettings.objects.first()
        except:
            ws = None
        
        # Get overtime threshold (default: 30 minutes)
        overtime_threshold = int(ws.overtime_threshold_minutes or 30) if ws else 30
        
        # Get attendance records for the date range
        # Exclude records that already have overtime requests
        attendance_records = Attendance.objects.filter(
            user=user,
            date_local__gte=start_date,
            date_local__lte=end_date,
            check_in_at_utc__isnull=False,
            check_out_at_utc__isnull=False
        ).exclude(
            overtime_requests__isnull=False
        ).order_by('date_local')
        
        potential_records = []
        
        for att in attendance_records:
            # Determine required minutes based on day of week and settings
            if att.date_local.weekday() == 4:  # Friday
                required_minutes = int(ws.friday_required_minutes or 240) if ws else 240
            else:
                required_minutes = int(ws.required_minutes or 480) if ws else 480
            
            # Calculate potential overtime (total work - required - threshold)
            potential_overtime_minutes = att.total_work_minutes - required_minutes - overtime_threshold
            
            # Only include if there's potential overtime (worked more than required + threshold)
            if potential_overtime_minutes > 0:
                potential_overtime_hours = potential_overtime_minutes / 60
                
                # Calculate potential overtime amount
                potential_amount = 0
                if employee.gaji_pokok and ws:
                    monthly_hours = 22 * 8
                    hourly_wage = float(employee.gaji_pokok) / monthly_hours
                    
                    # Determine rate
                    if att.is_holiday:
                        rate = float(ws.overtime_rate_holiday or 0.75)
                    else:
                        rate = float(ws.overtime_rate_workday or 0.50)
                    
                    potential_amount = potential_overtime_hours * hourly_wage * rate
                
                # Format times
                check_in_time = None
                check_out_time = None
                
                if att.check_in_at_utc:
                    check_in_time = att.check_in_at_utc.strftime('%H:%M')
                if att.check_out_at_utc:
                    check_out_time = att.check_out_at_utc.strftime('%H:%M')
                
                potential_records.append({
                    'date_local': att.date_local.isoformat(),
                    'weekday': att.date_local.strftime('%A'),
                    'check_in_time': check_in_time,
                    'check_out_time': check_out_time,
                    'total_work_minutes': att.total_work_minutes,
                    'required_minutes': required_minutes,
                    'overtime_threshold_minutes': overtime_threshold,
                    'potential_overtime_minutes': potential_overtime_minutes,
                    'potential_overtime_hours': round(potential_overtime_hours, 2),
                    'potential_overtime_amount': round(potential_amount, 2),
                    'is_holiday': att.is_holiday,
                    'within_geofence': att.within_geofence,
                    'can_submit': True,  # All records here are eligible for submission
                })
        
        return Response({
            'start_date': start_date.isoformat(),
            'end_date': end_date.isoformat(),
            'overtime_threshold_minutes': overtime_threshold,
            'total_potential_records': len(potential_records),
            'total_potential_hours': sum(r['potential_overtime_hours'] for r in potential_records),
            'total_potential_amount': sum(r['potential_overtime_amount'] for r in potential_records),
            'potential_records': potential_records,
        })

    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get overtime summary for current user"""
        # Get date range from query params
        end_date = date.today()
        start_date = end_date.replace(day=1)

        start_date_str = request.query_params.get('start_date', start_date.isoformat())
        end_date_str = request.query_params.get('end_date', end_date.isoformat())

        try:
            start_date = date.fromisoformat(start_date_str)
            end_date = date.fromisoformat(end_date_str)
        except ValueError:
            return Response(
                {"error": "Invalid date format. Use YYYY-MM-DD"},
                status=status.HTTP_400_BAD_REQUEST
            )

        service = OvertimeService()
        result = service.get_overtime_summary(request.user, start_date, end_date)

        if result['success']:
            summary = result['summary']
            # Enrich with supervisor capability flags for frontend UI logic
            is_admin = request.user.is_superuser or request.user.groups.filter(name='admin').exists()
            can_org_wide = False
            try:
                emp = getattr(request.user, 'employee_profile', None)
                pos = getattr(emp, 'position', None) if emp else None
                if pos and getattr(pos, 'can_approve_overtime_org_wide', False):
                    can_org_wide = True
            except Exception:
                can_org_wide = False
            summary.update({
                'can_approve_overtime_org_wide': can_org_wide,
                'is_admin': is_admin,
            })
            return Response(summary)
        else:
            return Response(
                {"error": result['message']},
                status=status.HTTP_400_BAD_REQUEST
            )


# Monthly Summary Request Views
class MonthlySummaryRequestViewSet(viewsets.ModelViewSet):
    """Monthly summary request management ViewSet with role-based access"""
    serializer_class = MonthlySummaryRequestSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_serializer_class(self):
        """Return appropriate serializer based on user role"""
        if self.request.user.is_superuser or self.request.user.groups.filter(name='admin').exists():
            return MonthlySummaryRequestAdminSerializer
        elif self.request.user.groups.filter(name='supervisor').exists():
            return MonthlySummaryRequestSupervisorSerializer
        else:
            return MonthlySummaryRequestEmployeeSerializer
    
    def get_queryset(self):
        """Filter monthly summary requests based on user role"""
        if self.request.user.is_superuser or self.request.user.groups.filter(name='admin').exists():
            return MonthlySummaryRequest.objects.all()
        elif self.request.user.groups.filter(name='supervisor').exists():
            # Mirror overtime visibility: org-wide supervisors see all; else only their division
            if hasattr(self.request.user, 'employee_profile') and self.request.user.employee_profile.position:
                if self.request.user.employee_profile.position.can_approve_overtime_org_wide:
                    return MonthlySummaryRequest.objects.all()
                elif self.request.user.employee_profile.division:
                    return MonthlySummaryRequest.objects.filter(
                        employee__division=self.request.user.employee_profile.division
                    )
                else:
                    return MonthlySummaryRequest.objects.none()
            else:
                return MonthlySummaryRequest.objects.none()
        else:
            # Regular employees can only see their own monthly summary requests
            return MonthlySummaryRequest.objects.filter(user=self.request.user)
    
    def perform_create(self, serializer):
        """Set user when creating monthly summary request"""
        serializer.save(user=self.request.user)
    
    def _get_monthly_template_path(self):
        """Return best template path for monthly overtime summary document or None."""
        template_dir = os.path.join(settings.BASE_DIR, 'template')
        if not os.path.isdir(template_dir):
            return None

        priority_names = [
            'template_rekap_lembur.docx',
            'template_rekap_lembur_bulanan.docx',
        ]

        for name in priority_names:
            path = os.path.join(template_dir, name)
            if os.path.exists(path):
                return path

        # Fallback to any latest .docx in template dir
        docx_files = [
            os.path.join(template_dir, f)
            for f in os.listdir(template_dir)
            if f.endswith('.docx') and not f.startswith('~$')
        ]
        if not docx_files:
            return None
        docx_files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
        return docx_files[0]

    def _generate_monthly_summary_docx(self, monthly_summary):
        """Generate DOCX bytes for the monthly summary request. Returns (bytes, filename)."""
        try:
            from docx import Document
        except Exception:
            raise RuntimeError('python-docx is not installed')

        template_path = self._get_monthly_template_path()
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            doc.add_heading('Rekap Lembur Bulanan', level=1)

        employee = monthly_summary.employee
        employee_name = None
        if employee:
            try:
                employee_name = employee.display_name
            except Exception:
                employee_name = getattr(employee, 'fullname', None)
        if not employee_name:
            if employee and getattr(employee, 'user', None):
                employee_name = employee.user.get_full_name() or employee.user.username
            else:
                employee_name = '-'

        employee_nip = getattr(employee, 'nip', '-') if employee else '-'
        division_name = employee.get_division_name() if employee else '-'
        position_name = employee.get_position_name() if employee else '-'

        # Determine month range
        from calendar import monthrange
        year = monthly_summary.year
        month = monthly_summary.month
        last_day = monthrange(year, month)[1]
        from datetime import date as _date
        start_date = _date(year, month, 1)
        end_date = _date(year, month, last_day)

        # Fetch approved overtime requests for this employee in period
        requests_qs = OvertimeRequest.objects.filter(
            employee=employee,
            status='approved',
            date__gte=start_date,
            date__lte=end_date,
        ).order_by('date')

        total_hours = 0
        total_amount = 0
        rows = []
        for r in requests_qs:
            hrs = float(r.total_hours or 0)
            amt = float(r.total_amount or 0)
            total_hours += hrs
            total_amount += amt
            rows.append({
                'date': r.date,
                'hours': hrs,
                'amount': amt,
                'desc': r.work_description,
            })

        # Build replacements
        bulan_tahun = f"{month:02d}/{year}"
        
        # Calculate additional metrics
        total_days = len(rows)  # Number of days with overtime
        avg_per_day = total_hours / total_days if total_days > 0 else 0
        
        # Get current date for export
        from django.utils import timezone
        current_dt = timezone.now()
        tanggal_export = current_dt.strftime('%d %B %Y')
        
        # Get approval info for monthly summary
        def _approver_name(user):
            if not user:
                return '-'
            try:
                emp = getattr(user, 'employee_profile', None)
                if emp and getattr(emp, 'fullname', None):
                    return emp.fullname
            except Exception:
                pass
            return user.get_full_name() or user.username

        def _approver_nip(user):
            try:
                emp = getattr(user, 'employee_profile', None)
                return getattr(emp, 'nip', '-') if emp else '-'
            except Exception:
                return '-'

        lvl1_name = _approver_name(getattr(monthly_summary, 'level1_approved_by', None))
        lvl1_nip = _approver_nip(getattr(monthly_summary, 'level1_approved_by', None))
        lvl1_at = getattr(monthly_summary, 'level1_approved_at', None)
        lvl1_date = lvl1_at.strftime('%d %B %Y') if lvl1_at else '-'

        final_name = _approver_name(getattr(monthly_summary, 'final_approved_by', None))
        final_nip = _approver_nip(getattr(monthly_summary, 'final_approved_by', None))
        final_at = getattr(monthly_summary, 'final_approved_at', None)
        final_date = final_at.strftime('%d %B %Y') if final_at else '-'
        
        replacements = {
            '{{NAMA_PEGAWAI}}': employee_name,
            '{{NIP_PEGAWAI}}': employee_nip or '-',
            '{{JABATAN_PEGAWAI}}': position_name,
            '{{DIVISI_PEGAWAI}}': division_name,
            '{{PERIODE}}': bulan_tahun,
            '{{TOTAL_JAM_LEMBUR}}': f"{total_hours:.2f}",
            '{{TOTAL_GAJI_LEMBUR}}': f"{total_amount:.2f}",
            
            # New placeholders
            '{{PERIODE_EXPORT}}': f"{month:02d} {year}",
            '{{TOTAL_HARI_LEMBUR}}': f"{total_days} hari",
            '{{RATA_RATA_PER_HARI}}': f"{avg_per_day:.2f} jam",
            '{{TANGGAL_EXPORT}}': tanggal_export,
            
            # Approval placeholders
            '{{LEVEL1_APPROVER}}': lvl1_name,
            '{{LEVEL1_APPROVER_NIP}}': lvl1_nip,
            '{{LEVEL1_APPROVAL_DATE}}': lvl1_date,
            '{{FINAL_APPROVER}}': final_name,
            '{{FINAL_APPROVER_NIP}}': final_nip,
            '{{FINAL_APPROVAL_DATE}}': final_date,
        }

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for old, new in replacements.items():
                if old in paragraph.text:
                    runs = paragraph.runs
                    for i in range(len(runs)):
                        runs[i].text = runs[i].text.replace(old, str(new))

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old, new in replacements.items():
                            if old in paragraph.text:
                                runs = paragraph.runs
                                for i in range(len(runs)):
                                    runs[i].text = runs[i].text.replace(old, str(new))

        # Append overtime table at the end
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tanggal'
        hdr_cells[1].text = 'Jam Lembur'
        hdr_cells[2].text = 'Jumlah'
        hdr_cells[3].text = 'Deskripsi'
        for item in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = item['date'].strftime('%d %b %Y')
            row_cells[1].text = f"{item['hours']:.2f}j"
            row_cells[2].text = f"{item['amount']:.2f}"
            row_cells[3].text = item['desc'] or ''

        # Serialize to bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)
        try:
            with open(tmp_path, 'rb') as f:
                content = f.read()
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        filename = f"rekap_lembur_{employee_nip or 'pegawai'}_{year}-{month:02d}.docx"
        return content, filename

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """Approve or reject a monthly summary request (two-level approval)"""
        if not (request.user.is_superuser or 
                request.user.groups.filter(name__in=['admin', 'supervisor']).exists()):
            return Response(
                {"error": "Only admins and supervisors can approve monthly summary requests"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        monthly_summary = self.get_object()
        serializer = MonthlySummaryRequestApprovalSerializer(data=request.data)
        
        if serializer.is_valid():
            action = serializer.validated_data['action']
            approval_level = serializer.validated_data.get('approval_level')
            reason = serializer.validated_data.get('reason', '')
            
            if action == 'approve':
                # Determine approval level based on role if not provided
                if approval_level is None:
                    # Org-wide approvers (admin) default to final approval
                    if request.user.is_superuser or request.user.groups.filter(name='admin').exists():
                        approval_level = 2
                    else:
                        approval_level = 1

                if approval_level == 1:
                    if monthly_summary.status != 'pending':
                        return Response({"detail": "Request is not pending approval."}, status=status.HTTP_400_BAD_REQUEST)
                    monthly_summary.approve_level1(request.user)
                    return Response({"message": "Monthly summary request level 1 approved"}, status=status.HTTP_200_OK)
                elif approval_level == 2:
                    if monthly_summary.status not in ['pending', 'level1_approved']:
                        return Response({"detail": "Request is not awaiting final approval."}, status=status.HTTP_400_BAD_REQUEST)
                    monthly_summary.approve_final(request.user)
                    return Response({"message": "Monthly summary request final approved"}, status=status.HTTP_200_OK)
                else:
                    return Response({"detail": "Invalid approval_level"}, status=status.HTTP_400_BAD_REQUEST)
            elif action == 'reject':
                monthly_summary.reject(request.user, reason)
                return Response(
                    {"message": "Monthly summary request rejected successfully"}, 
                    status=status.HTTP_200_OK
                )
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=False, methods=['get'])
    def pending(self, request):
        """Get pending monthly summary requests for approval"""
        if not (request.user.is_superuser or 
                request.user.groups.filter(name__in=['admin', 'supervisor']).exists()):
            return Response(
                {"error": "Only admins and supervisors can view pending monthly summary requests"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        queryset = self.get_queryset().filter(status='pending')
        serializer = MonthlySummaryRequestListSerializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def export_docx(self, request, pk=None):
        """Export approved monthly summary to DOCX."""
        summary = self.get_object()
        if summary.status != 'approved':
            return Response({"detail": "Monthly summary must be approved to export."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            docx_bytes, filename = self._generate_monthly_summary_docx(summary)
            response = HttpResponse(
                docx_bytes,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            return Response({"detail": f"Gagal export DOCX: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Export approved monthly summary to PDF using converter service."""
        summary = self.get_object()
        if summary.status != 'approved':
            return Response({"detail": "Monthly summary must be approved to export PDF."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            # Generate DOCX first
            docx_bytes, base_filename = self._generate_monthly_summary_docx(summary)
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(docx_bytes)
                tmp_path = tmp.name
            try:
                with open(tmp_path, 'rb') as f:
                    files = {
                        'file': (
                            'document.docx',
                            f,
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        )
                    }
                    data = {'method': 'file'}
                    converter_url = 'http://docx_converter:5000/convert'
                    r = requests.post(converter_url, files=files, data=data, timeout=60)
                    if r.status_code != 200:
                        return Response({"detail": f"DOCX converter error: {r.status_code}"}, status=status.HTTP_502_BAD_GATEWAY)
                    pdf_content = r.content
            finally:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

            pdf_filename = base_filename.replace('.docx', '.pdf')
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
        except Exception as e:
            return Response({"detail": f"Gagal export PDF: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['get'])
    def my_summaries(self, request):
        """Get current user's monthly summary requests"""
        queryset = MonthlySummaryRequest.objects.filter(user=request.user)
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def by_month_year(self, request):
        """Get monthly summary requests filtered by month and year"""
        month = request.query_params.get('month')
        year = request.query_params.get('year')
        
        queryset = self.get_queryset()
        if month:
            queryset = queryset.filter(month=month)
        if year:
            queryset = queryset.filter(year=year)
        
        serializer = MonthlySummaryRequestListSerializer(queryset, many=True)
        return Response(serializer.data)


# Role-specific ViewSets for backward compatibility
class AdminOvertimeRequestViewSet(OvertimeRequestViewSet):
    """Admin-specific overtime request ViewSet"""
    permission_classes = [IsAdmin]
    
    def get_queryset(self):
        return OvertimeRequest.objects.all()


class SupervisorOvertimeRequestViewSet(OvertimeRequestViewSet):
    """Supervisor-specific overtime request ViewSet"""
    permission_classes = [IsSupervisor]
    
    def get_queryset(self):
        # Supervisors can see overtime requests of employees in their division
        # If they have org-wide approval, they can see all overtime requests
        if hasattr(self.request.user, 'employee_profile') and self.request.user.employee_profile.position:
            if self.request.user.employee_profile.position.can_approve_overtime_org_wide:
                # Org-wide approval: can see all overtime requests
                return OvertimeRequest.objects.all()
            elif self.request.user.employee_profile.division:
                # Division-only approval: can see only their division's requests
                return OvertimeRequest.objects.filter(
                    employee__division=self.request.user.employee_profile.division
                )
        return OvertimeRequest.objects.none()


class EmployeeOvertimeRequestViewSet(OvertimeRequestViewSet):
    """Employee-specific overtime request ViewSet"""
    permission_classes = [IsEmployee]
    
    def get_queryset(self):
        # Employees can only see their own overtime requests
        return OvertimeRequest.objects.filter(user=self.request.user)
