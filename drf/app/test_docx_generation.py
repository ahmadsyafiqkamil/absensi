#!/usr/bin/env python
"""
Test script untuk generate DOCX menggunakan template rekap lembur bulanan
"""
import os
import django
from datetime import datetime, date
from decimal import Decimal

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'core.settings')
django.setup()

from docx import Document
from django.conf import settings
from django.core.files.base import ContentFile
from api.models import OvertimeSummaryRequest, OvertimeSummaryDocument, Employee
from django.contrib.auth import get_user_model
from django.utils import timezone as dj_timezone

def test_template_loading():
    """Test loading template and checking placeholders"""
    print("=== Testing Template Loading ===")

    template_path = os.path.join(settings.BASE_DIR, 'template', 'template_rekap_lembur.docx')
    print(f"Template path: {template_path}")
    print(f"Template exists: {os.path.exists(template_path)}")

    if os.path.exists(template_path):
        try:
            doc = Document(template_path)
            print(f"✓ Template loaded successfully with {len(doc.paragraphs)} paragraphs")

            # Check all placeholders in document
            placeholders = []
            for para in doc.paragraphs:
                text = para.text
                if '{{' in text and '}}' in text:
                    import re
                    matches = re.findall(r'\{\{[^}]+\}\}', text)
                    placeholders.extend(matches)

            # Check in tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            if '{{' in text and '}}' in text:
                                import re
                                matches = re.findall(r'\{\{[^}]+\}\}', text)
                                placeholders.extend(matches)

            unique_placeholders = list(set(placeholders))
            print(f"✓ Found {len(unique_placeholders)} unique placeholders:")
            for placeholder in sorted(unique_placeholders):
                print(f"  - {placeholder}")

            return doc, unique_placeholders

        except Exception as e:
            print(f"✗ Error loading template: {str(e)}")
            return None, []
    else:
        print("✗ Template file not found")
        return None, []

def create_mock_employee():
    """Create mock employee for testing"""
    print("\n=== Creating Mock Employee ===")

    try:
        # Check if test user exists
        User = get_user_model()
        user, created = User.objects.get_or_create(
            username='test_employee',
            defaults={
                'email': 'test@example.com',
                'first_name': 'Test',
                'last_name': 'Employee'
            }
        )

        # Check if employee exists
        employee, created = Employee.objects.get_or_create(
            user=user,
            defaults={
                'nip': '123456789',
                'division_id': 1,  # Assuming division exists
                'position_id': 1,  # Assuming position exists
                'gaji_pokok': Decimal('5000.00'),
                'fullname': 'Test Employee'
            }
        )

        print(f"✓ Mock employee created: {employee.fullname} (NIP: {employee.nip})")
        return employee

    except Exception as e:
        print(f"✗ Error creating mock employee: {str(e)}")
        return None

def create_mock_overtime_request(employee, target_period='2025-08'):
    """Create mock overtime summary request for specific period"""
    print(f"\n=== Creating Mock Overtime Summary Request for {target_period} ===")

    try:
        # Delete existing request for the same period if exists
        try:
            existing = OvertimeSummaryRequest.objects.get(
                employee=employee,
                request_period=target_period
            )
            existing.delete()
            print(f"✓ Deleted existing request for period {target_period}")
        except OvertimeSummaryRequest.DoesNotExist:
            pass

        # Create request for August 2025
        request = OvertimeSummaryRequest.objects.create(
            employee=employee,
            user=employee.user,
            request_period=target_period,
            request_title=f"Rekap Lembur {datetime.strptime(target_period, '%Y-%m').strftime('%B %Y')}",
            request_description="Test request untuk generate DOCX dengan tabel overtime",
            include_overtime_details=True,
            include_overtime_summary=True,
            include_approver_info=True,
            status='approved'  # Set to approved for testing
        )

        print(f"✓ Mock request created: {request.request_title}")
        print(f"  Period: {request.request_period}")
        print(f"  Status: {request.status}")

        return request

    except Exception as e:
        print(f"✗ Error creating mock request: {str(e)}")
        return None

def create_mock_overtime_data(employee, request, period='2025-08'):
    """Create mock overtime data for August 2025"""
    print(f"\n=== Creating Mock Overtime Data for {period} ===")

    from api.models import OvertimeRequest

    try:
        # Delete existing overtime requests for this period
        existing_overtime = OvertimeRequest.objects.filter(
            employee=employee,
            date_requested__year=int(period.split('-')[0]),
            date_requested__month=int(period.split('-')[1])
        )
        count_deleted = existing_overtime.count()
        existing_overtime.delete()
        if count_deleted > 0:
            print(f"✓ Deleted {count_deleted} existing overtime records")

        # Create realistic overtime data for August 2025
        overtime_data = [
            # Week 1 (1-7 Agustus)
            {'date': '2025-08-01', 'hours': '2.5', 'description': 'Penyelesaian laporan bulanan keuangan'},
            {'date': '2025-08-02', 'hours': '1.5', 'description': 'Meeting dengan stakeholder internasional'},
            {'date': '2025-08-05', 'hours': '3.0', 'description': 'Persiapan presentasi diplomatik'},

            # Week 2 (8-14 Agustus)
            {'date': '2025-08-08', 'hours': '2.0', 'description': 'Review dokumen visa dan imigrasi'},
            {'date': '2025-08-12', 'hours': '2.5', 'description': 'Koordinasi dengan kedutaan asing'},
            {'date': '2025-08-13', 'hours': '1.5', 'description': 'Update sistem database konsuler'},

            # Week 3 (15-21 Agustus)
            {'date': '2025-08-15', 'hours': '3.5', 'description': 'Persiapan kunjungan delegasi'},
            {'date': '2025-08-18', 'hours': '2.0', 'description': 'Dokumentasi acara protokoler'},
            {'date': '2025-08-19', 'hours': '2.5', 'description': 'Terjemahan dokumen resmi'},

            # Week 4 (22-28 Agustus)
            {'date': '2025-08-22', 'hours': '1.5', 'description': 'Verifikasi data pelaporan'},
            {'date': '2025-08-25', 'hours': '2.0', 'description': 'Koordinasi dengan instansi terkait'},
            {'date': '2025-08-26', 'hours': '3.0', 'description': 'Penyelesaian laporan akhir bulan'},
        ]

        overtime_requests = []

        for i, data in enumerate(overtime_data):
            overtime_req = OvertimeRequest.objects.create(
                employee=employee,
                user=employee.user,
                date_requested=date.fromisoformat(data['date']),
                overtime_hours=Decimal(data['hours']),
                work_description=data['description'],
                status='approved'
            )
            overtime_requests.append(overtime_req)

        print(f"✓ Created {len(overtime_requests)} mock overtime requests for August 2025")

        # Convert to queryset for compatibility with the view method
        from django.db.models import QuerySet
        if isinstance(overtime_requests, list):
            # Get queryset from the same model
            queryset = OvertimeRequest.objects.filter(id__in=[req.id for req in overtime_requests])
            return queryset

        return overtime_requests

    except Exception as e:
        print(f"✗ Error creating mock overtime data: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def test_docx_generation():
    """Main test function"""
    print("=== Testing DOCX Generation ===\n")

    # Step 1: Test template loading
    doc, placeholders = test_template_loading()
    if not doc:
        return

    # Step 2: Create mock data
    employee = create_mock_employee()
    if not employee:
        return

    request = create_mock_overtime_request(employee)
    if not request:
        return

    overtime_data = create_mock_overtime_data(employee, request)
    if not overtime_data:
        return

    # Step 3: Test DOCX generation using the actual method
    print("\n=== Testing DOCX Generation ===")

    try:
        # Import the ViewSet to use its method
        from api.views import OvertimeSummaryRequestViewSet

        viewset = OvertimeSummaryRequestViewSet()

        # Generate DOCX
        docx_path = viewset.generate_monthly_summary_docx(request, overtime_data, request.request_period)

        if docx_path and os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"✓ DOCX generated successfully: {docx_path}")
            print(f"  File size: {file_size} bytes")

            # Check if document was created in database
            try:
                docx_document = OvertimeSummaryDocument.objects.get(
                    overtime_summary_request=request,
                    document_type='monthly_summary'
                )
                print(f"✓ Document record created in database")
                print(f"  Document ID: {docx_document.id}")
                print(f"  Status: {docx_document.status}")

                if docx_document.docx_file:
                    print(f"  DOCX file URL: {docx_document.docx_file.url}")

            except OvertimeSummaryDocument.DoesNotExist:
                print("⚠ Document record not found in database")

        else:
            print("✗ DOCX generation failed")

    except Exception as e:
        print(f"✗ Error during DOCX generation: {str(e)}")
        import traceback
        traceback.print_exc()

def test_template_replacement():
    """Test template placeholder replacement"""
    print("\n=== Testing Template Replacement ===")

    try:
        # Load template
        template_path = os.path.join(settings.BASE_DIR, 'template', 'template_rekap_lembur.docx')
        doc = Document(template_path)

        # Sample replacement data
        replacements = {
            '{{PERIODE_EXPORT}}': 'Januari 2025',
            '{{TANGGAL_EXPORT}}': '15 Januari 2025',
            '{{NAMA_PEGAWAI}}': 'Ahmad Syafiq Kamil',
            '{{NIP_PEGAWAI}}': '123456789',
            '{{TOTAL_HARI_LEMBUR}}': '5 hari',
            '{{TOTAL_JAM_LEMBUR}}': '12.5 jam',
            '{{TABEL_OVERTIME}}': '[TABEL OVERTIME AKAN DIGANTI]'
        }

        # Apply replacements
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        # Save test file
        output_path = '/tmp/test_replacement.docx'
        doc.save(output_path)

        if os.path.exists(output_path):
            print(f"✓ Template replacement successful: {output_path}")
            print(f"  File size: {os.path.getsize(output_path)} bytes")
        else:
            print("✗ Template replacement failed")

    except Exception as e:
        print(f"✗ Error during template replacement: {str(e)}")
        import traceback
        traceback.print_exc()

def test_full_docx_generation(target_period='2025-08'):
    """Test full DOCX generation with mock data for August"""
    print(f"\n=== Testing Full DOCX Generation for {target_period} ===\n")

    # Step 1: Create mock data
    employee = create_mock_employee()
    if not employee:
        return

    request = create_mock_overtime_request(employee, target_period)
    if not request:
        return

    overtime_data = create_mock_overtime_data(employee, request, target_period)
    if not overtime_data:
        return

    # Step 2: Test DOCX generation
    try:
        # Import the ViewSet to use its method
        from api.views import OvertimeSummaryRequestViewSet

        viewset = OvertimeSummaryRequestViewSet()

        # Generate DOCX
        print("Generating DOCX...")
        docx_path = viewset.generate_monthly_summary_docx(request, overtime_data, request.request_period)

        if docx_path and os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"✓ DOCX generated successfully: {docx_path}")
            print(f"  File size: {file_size} bytes")

            # Check if document was created in database
            try:
                docx_document = OvertimeSummaryDocument.objects.get(
                    overtime_summary_request=request,
                    document_type='monthly_summary'
                )
                print(f"✓ Document record created in database")
                print(f"  Document ID: {docx_document.id}")
                print(f"  Status: {docx_document.status}")

                if docx_document.docx_file:
                    print(f"  DOCX file URL: {docx_document.docx_file.url}")

                # Try to load generated document to verify
                try:
                    generated_doc = Document(docx_path)
                    print(f"✓ Generated document can be loaded with {len(generated_doc.paragraphs)} paragraphs")

                    # Check if placeholders were replaced
                    sample_text = generated_doc.paragraphs[0].text if generated_doc.paragraphs else ""
                    if employee.fullname in sample_text:
                        print("✓ Employee name found in generated document")
                    else:
                        print("⚠ Employee name not found in generated document")

                except Exception as e:
                    print(f"⚠ Could not verify generated document: {str(e)}")

            except OvertimeSummaryDocument.DoesNotExist:
                print("✗ Document record not found in database")

        else:
            print("✗ DOCX generation failed")
            if docx_path:
                print(f"Expected path: {docx_path}")

    except Exception as e:
        print(f"✗ Error during DOCX generation: {str(e)}")
        import traceback
        traceback.print_exc()

def test_api_endpoints():
    """Test API endpoints for overtime summary request"""
    print("\n=== Testing API Endpoints ===\n")

    try:
        from django.test import Client
        from django.contrib.auth import get_user_model

        client = Client()

        # Try to get existing user or create test user
        User = get_user_model()
        try:
            user = User.objects.get(username='test_employee')
            print(f"✓ Using existing test user: {user.username}")
            # Set password if not set
            if not user.has_usable_password():
                user.set_password('testpass123')
                user.save()
                print("✓ Password set for test user")
        except User.DoesNotExist:
            print("⚠ Test user not found, creating...")
            user = User.objects.create_user(
                username='test_employee',
                email='test@example.com',
                password='testpass123',
                first_name='Test',
                last_name='Employee'
            )
            print(f"✓ Created test user: {user.username}")

        # Login
        login_success = client.login(username='test_employee', password='testpass123')
        if login_success:
            print("✓ User login successful")
        else:
            print("✗ User login failed")
            return

        # Test list endpoint
        response = client.get('/api/overtime-summary-requests/')
        print(f"✓ List endpoint response: {response.status_code}")

        if response.status_code == 200:
            data = response.json()
            print(f"  Found {len(data)} requests")

        # Test create endpoint
        create_data = {
            'request_period': '2025-01',
            'request_title': 'Test Rekap Lembur Januari 2025',
            'request_description': 'Testing API endpoint'
        }

        response = client.post('/api/overtime-summary-requests/', create_data, content_type='application/json')
        print(f"✓ Create endpoint response: {response.status_code}")

        if response.status_code == 201:
            created_data = response.json()
            request_id = created_data.get('id')
            print(f"  Created request ID: {request_id}")

            # Test detail endpoint
            response = client.get(f'/api/overtime-summary-requests/{request_id}/')
            print(f"✓ Detail endpoint response: {response.status_code}")

            # Test generate report endpoint
            response = client.get(f'/api/overtime-summary-requests/{request_id}/generate_report/')
            print(f"✓ Generate report endpoint response: {response.status_code}")

            if response.status_code == 200:
                report_data = response.json()
                print("✓ Report data generated successfully")
                # Print some sample data
                if 'employee_info' in report_data:
                    employee_info = report_data['employee_info']
                    print(f"  Employee: {employee_info.get('fullname', 'N/A')}")

    except Exception as e:
        print(f"✗ Error testing API endpoints: {str(e)}")
        import traceback
        traceback.print_exc()

def test_tabel_overtime_placeholder(target_period='2025-08'):
    """Test specifically the TABEL_OVERTIME placeholder for August data"""
    print(f"\n=== Testing TABEL_OVERTIME Placeholder for {target_period} ===\n")

    # Step 1: Create mock data
    employee = create_mock_employee()
    if not employee:
        return

    request = create_mock_overtime_request(employee, target_period)
    if not request:
        return

    overtime_data = create_mock_overtime_data(employee, request, target_period)
    if not overtime_data:
        return

    # Step 2: Show overtime data that will be included in table
    print("Overtime data to be included in table:")
    for i, record in enumerate(overtime_data, 1):
        print(f"{i:2d}. {record.date_requested.strftime('%d/%m/%Y')} - {float(record.overtime_hours):.1f} jam - {record.work_description}")

    # Step 3: Test DOCX generation
    try:
        # Import the ViewSet to use its method
        from api.views import OvertimeSummaryRequestViewSet

        viewset = OvertimeSummaryRequestViewSet()

        # Generate DOCX
        print(f"\nGenerating DOCX for {target_period}...")
        docx_path = viewset.generate_monthly_summary_docx(request, overtime_data, target_period)

        if docx_path and os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"✓ DOCX generated successfully: {docx_path}")
            print(f"  File size: {file_size} bytes")

            # Analyze the generated document specifically for table content
            print("\n=== Analyzing Generated Table Content ===")
            from docx import Document
            doc = Document(docx_path)

            # Find tables in the document
            for table_idx, table in enumerate(doc.tables):
                print(f"\nTable {table_idx + 1}:")
                print("Headers:", [cell.text.strip() for cell in table.rows[0].cells])

                # Show first few data rows
                for row_idx in range(1, min(6, len(table.rows))):  # Show up to 5 data rows
                    row_data = [cell.text.strip() for cell in table.rows[row_idx].cells]
                    print(f"Row {row_idx}: {row_data}")

                if len(table.rows) > 6:
                    print(f"... and {len(table.rows) - 6} more rows")

            print(f"\n✓ Generated document contains {len(doc.tables)} table(s)")

        else:
            print("✗ DOCX generation failed")

    except Exception as e:
        print(f"✗ Error during DOCX generation: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    print("Starting DOCX Generation Test for TABEL_OVERTIME Placeholder...\n")

    # Test basic template loading
    test_template_loading()

    # Test specifically the TABEL_OVERTIME placeholder for August
    test_tabel_overtime_placeholder('2025-08')

    print("\n=== TABEL_OVERTIME Test Complete ===")
