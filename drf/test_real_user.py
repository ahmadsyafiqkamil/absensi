#!/usr/bin/env python
"""
Test script untuk generate DOCX dengan akun staff.konsuler1
"""
import os
import django
from datetime import date
from decimal import Decimal

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'core.settings')
django.setup()

from django.contrib.auth import get_user_model
from api.models import OvertimeSummaryRequest, OvertimeRequest
from api.views import OvertimeSummaryRequestViewSet
from django.utils import timezone as dj_timezone

def test_staff_konsuler1_august_2025():
    """Test DOCX generation with staff.konsuler1 for August 2025"""
    print("=== Testing DOCX Generation with staff.konsuler1 for August 2025 ===\n")

    try:
        # Get the user
        User = get_user_model()
        user = User.objects.get(username='staff.konsuler1')
        print(f"✓ Found user: {user.username}")

        # Get employee
        employee = user.employee
        print(f"✓ Employee: {employee.fullname} (NIP: {employee.nip})")

        # Create overtime summary request for August 2025
        print("\n--- Creating Overtime Summary Request ---")

        # Delete existing request if any
        try:
            existing = OvertimeSummaryRequest.objects.get(
                employee=employee,
                request_period='2025-08'
            )
            existing.delete()
            print("✓ Deleted existing request")
        except OvertimeSummaryRequest.DoesNotExist:
            pass

        # Create new request
        request = OvertimeSummaryRequest.objects.create(
            employee=employee,
            user=user,
            request_period='2025-08',
            request_title="Rekap Lembur Agustus 2025",
            request_description="Rekapitulasi lembur bulan Agustus 2025 dengan akun staff.konsuler1",
            include_overtime_details=True,
            include_overtime_summary=True,
            include_approver_info=True,
            status='approved'
        )
        print(f"✓ Created overtime summary request: {request.request_title}")

        # Create sample overtime data for August 2025
        print("\n--- Creating Overtime Data ---")

        # Delete existing overtime for August 2025
        existing_overtime = OvertimeRequest.objects.filter(
            employee=employee,
            date_requested__year=2025,
            date_requested__month=8
        )
        count_deleted = existing_overtime.count()
        existing_overtime.delete()
        if count_deleted > 0:
            print(f"✓ Deleted {count_deleted} existing overtime records")

        # Create sample overtime data for August 2025
        overtime_data = [
            (date(2025, 8, 1), Decimal('2.5'), 'Penyelesaian laporan bulanan keuangan'),
            (date(2025, 8, 2), Decimal('1.5'), 'Meeting dengan stakeholder internasional'),
            (date(2025, 8, 5), Decimal('3.0'), 'Persiapan presentasi diplomatik'),
            (date(2025, 8, 8), Decimal('2.0'), 'Review dokumen visa dan imigrasi'),
            (date(2025, 8, 12), Decimal('2.5'), 'Koordinasi dengan kedutaan asing'),
            (date(2025, 8, 13), Decimal('1.5'), 'Update sistem database konsuler'),
            (date(2025, 8, 15), Decimal('3.5'), 'Persiapan kunjungan delegasi'),
            (date(2025, 8, 18), Decimal('2.0'), 'Dokumentasi acara protokoler'),
            (date(2025, 8, 19), Decimal('2.5'), 'Terjemahan dokumen resmi'),
            (date(2025, 8, 22), Decimal('1.5'), 'Verifikasi data pelaporan'),
            (date(2025, 8, 25), Decimal('2.0'), 'Koordinasi dengan instansi terkait'),
            (date(2025, 8, 26), Decimal('3.0'), 'Penyelesaian laporan akhir bulan'),
        ]

        overtime_requests = []
        for req_date, hours, description in overtime_data:
            overtime_req = OvertimeRequest.objects.create(
                employee=employee,
                user=user,
                date_requested=req_date,
                overtime_hours=hours,
                work_description=description,
                status='approved'
            )
            overtime_requests.append(overtime_req)

        print(f"✓ Created {len(overtime_requests)} overtime requests for August 2025")

        # Show overtime data summary
        print("\n--- Overtime Data Summary ---")
        total_hours = sum(float(req.overtime_hours) for req in overtime_requests)
        total_amount = sum(float(req.overtime_amount) for req in overtime_requests)
        print(f"Total overtime hours: {total_hours:.1f} jam")
        print(f"Total overtime amount: {total_amount:.2f} AED")
        print("Overtime details:")
        for i, req in enumerate(overtime_requests, 1):
            print(f"{i:2d}. {req.date_requested.strftime('%d/%m/%Y')} - {float(req.overtime_hours):.1f} jam - {req.work_description}")

        # Generate DOCX
        print("\n--- Generating DOCX Report ---")
        viewset = OvertimeSummaryRequestViewSet()

        print("Generating DOCX report...")
        docx_path = viewset.generate_monthly_summary_docx(request, overtime_requests, '2025-08')

        if docx_path and os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"✓ DOCX generated successfully: {docx_path}")
            print(f"  File size: {file_size} bytes")

            # Check document record
            try:
                docx_document = request.docx_document
                if docx_document:
                    print(f"✓ Document record created in database")
                    print(f"  Document ID: {docx_document.id}")
                    print(f"  Status: {docx_document.status}")
                    print(f"  DOCX URL: {docx_document.docx_file.url}")
                else:
                    print("⚠ Document record not linked")
            except:
                print("⚠ Could not check document record")

            # Analyze the generated document
            print("\n--- Analyzing Generated Document ---")
            from docx import Document
            doc = Document(docx_path)

            print(f"✓ Document loaded with {len(doc.paragraphs)} paragraphs")

            # Find overtime table
            overtime_table_found = False
            for table_idx, table in enumerate(doc.tables):
                if len(table.rows) > 1:  # Has data rows
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    if 'Tanggal' in headers and 'Jam Lembur' in headers:
                        overtime_table_found = True
                        print(f"✓ Found overtime table (Table {table_idx + 1})")
                        print(f"  Headers: {headers}")
                        print(f"  Data rows: {len(table.rows) - 2}")  # -2 for header and total

                        # Show first few rows
                        print("  Sample data rows:")
                        for row_idx in range(1, min(4, len(table.rows) - 1)):  # Show first 3 data rows
                            row_data = [cell.text.strip() for cell in table.rows[row_idx].cells]
                            print(f"    Row {row_idx}: {row_data}")
                        break

            if not overtime_table_found:
                print("⚠ Overtime table not found in document")

            print("\n✓ Staff.konsuler1 test completed successfully!")
            print(f"✓ Generated DOCX file: {docx_path}")

        else:
            print("✗ DOCX generation failed")
            if docx_path:
                print(f"Expected path: {docx_path}")

    except Exception as e:
        print(f"✗ Error testing with staff.konsuler1: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    test_staff_konsuler1_august_2025()
